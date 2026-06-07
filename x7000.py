import os
from docx import Document
from docx.shared import Pt
from docx2pdf import convert

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
doc = Document(path)

t = '''关于数据时效性的补充说明。本报告核心数据来自36氪2017年行业分析文章。虽然该文章发布距今已有近九年时间，但它仍是公开可查的最为详细和系统的高空幕墙清洗行业分析。文章中的市场规模、政府政策、技术难点等核心内容具有较长时间的参考价值。竞品信息通过搜狗微信公众号搜索获取，可以复现搜索结果进行核实。建议在使用本报告时结合最新的行业动态进行更新。'''

for line in t.strip().split('\n'):
    line = line.strip()
    if line:
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

doc.save(path)

text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text
hz = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
t = len(text)
print(f'Hanzi: {hz}')
print(f'Total: {t}')

img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_count += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Images: {img_count}')

pdf_path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
