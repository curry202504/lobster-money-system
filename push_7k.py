import os
from docx import Document
from docx.shared import Pt
from docx2pdf import convert

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
doc = Document(path)

endtext = '''
关于行业展望的补充思考。高空清洗机器人行业的发展路径可以参考其他行业的自动化替代过程。一般来说，替代过程会经历几个阶段。第一阶段是技术验证期，少量先行者开始使用新产品。第二阶段是成本拐点期，随着技术成熟和规模效应，新方案的成本降至与旧方案相当或更低，替代加速。第三阶段是政策推动期，政府出台法规强制或鼓励采用新方案。第四阶段是全面普及期，新方案成为行业标配。目前高空清洗机器人行业可能正处于第一阶段向第二阶段过渡的时期。如果安全监管政策进一步收紧，可能直接推动行业进入第三阶段，加速普及进程。这意味着市场爆发的时间点可能比预期更早到来。'''

for line in endtext.strip().split('\n'):
    line = line.strip()
    if line:
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

doc.save(path)

text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text
hz = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
t = len(text)
print(f'Hanzi: {hz}')
print(f'Total: {t}')

img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_count += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Images: {img_count}')

pdf_path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
