# -*- coding: utf-8 -*-
"""MASTER REPORT - Following original document structure with images"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image
from docx2pdf import convert

OUT = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace')
IMG = os.path.join(OUT, 'report-images')

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

def H(t, lv=1):
    hs = doc.add_heading(t, level=lv)
    for r in hs.runs: r.font.name = 'Calibri'
    return hs

def P(t, b=False, s=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t)
    r.font.name = 'Calibri'
    if b: r.bold = True
    if s: r.font.size = Pt(s)
    return p

def TB(hd, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hd))
    t.style = 'Light Shading Accent 1'
    for i, h in enumerate(hd):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(v)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph(); return t

def add_img(img_path, caption):
    if os.path.exists(img_path):
        try:
            img = Image.open(img_path)
            w, h = img.size
            ratio = min(1.0, 380 / w) if w > 0 else 1
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(img_path, width=Cm(ratio * w * 0.035))
        except: pass
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(8)
    run.italic = True

def BR(): doc.add_page_break()

S = 'https://36kr.com/p/1721639649281'

# ====== COVER ======
for _ in range(4): doc.add_paragraph()
P('高空清洗机器人市场快速调研报告', b=True, s=26)
P('哈工鹏泽（深圳）机器人技术有限公司', b=True, s=16)
P('2026年5月', s=12)
doc.add_paragraph()
P('调研目标', b=True)
P('1. 市场蛋糕有多大，增长有多快？')
P('2. 主要对手是谁，他们怎么切蛋糕？')
P('3. 客户最愿意为什么买单，我们怎么把蛋糕抢过来？')
doc.add_paragraph()
P('数据来源：36氪行业分析、搜狗微信公众号', s=10)
BR()

# ====== MODULE 1 ======
H('模块一：市场与竞争')
P('对应调研文档问题：市场蛋糕有多大？主要对手是谁？他们怎么切蛋糕？')

H('1.1 真实市场规模')
P('根据36氪行业分析文章，中国目前建筑面积约500亿平方米，幕墙占比约5%，幕墙总面积约30亿平方米。每年新建幕墙约9000万平方米。以玻璃幕墙清洗服务计算，按每平方米3元的低端收费标准、每年清洗一次，市场规模约75亿元以上。若将高空墙体粉刷等延伸服务纳入，规模可达百亿元以上。')
P('多城市已出台强制性外墙清洗规定。北京每年至少1次，上海每年2至4次，广州重点区4次一般区2次，深圳、天津、杭州、南京等均有明确要求。36氪记者直接咨询了北京泰康金融大厦38层物业，确认了收费标准和清洗频率。')
P('目前高空清洗主要依赖蜘蛛人人工清洗，从业者属5类6类高风险职业，保费高且多无保险。媒体报道每年数以万计高空作业事故发生，死亡率80%以上。升降平台成本高，轨道系统需前置设计。58同城数据显示高危职业用工缺口逐年扩大。')

H('1.2 技术难点')
P('高空幕墙清洗机器人研发有四大技术难点。吸附功能需牢固吸附垂直墙面。移动功能需适应玻璃、金属、粉墙等多种材质和曲面壁面。越障功能要求跨越窗框等障碍物，被36氪评价为目前比较难解决的问题。清洗功能要求在移动中完成有效清洁。')
P('国际方面，德国Skywash、日本BE等曾推出方案但未大规模推广。国内自1988年上海大学研制首款机器人以来，哈尔滨工业大学在863计划支持下开展研究，北京航空航天大学研制了多款样机，但均未能商业化。现阶段核心零部件技术成熟、用工成本上升、政策管控加强，行业具备了发展土壤。')

H('1.3 核心对手画像（对应调研文档要求）')

# ---- 竞品1: 哈工鹏泽 ----
P('竞品1：哈工鹏泽（深圳）机器人技术有限公司（本公司）', b=True)
add_img(os.path.join(IMG, '哈工鹏泽_ge02_1.jpg'), '哈工鹏泽GE-02产品图')
add_img(os.path.join(IMG, '哈工鹏泽_ge02_0.jpg'), '哈工鹏泽产品图')
TB(['维度', '详情'],
    [['公司', '哈工鹏泽（深圳）机器人技术有限公司'],
     ['核心产品', 'GE-02滚轮式越障高空清洗机器人\nGE-05轻型吸附式横向清洗机器人'],
     ['近期活动', '2026年5月走进深圳光明区凤凰街道智能清洗推介\n出席2026场景创新大会\n参展深圳绿色饭店发展促进会'],
     ['微信公众号', '哈工鹏泽机器人技术']])
P('来源：搜狗微信搜索哈工鹏泽', s=9)

# ---- 竞品2: 埃欧珞 ----
P('竞品2：埃欧珞（杭州）科技有限公司', b=True)
add_img(os.path.join(IMG, '埃欧珞_0.png'), '埃欧珞产品图')
add_img(os.path.join(IMG, '埃欧珞_rs_0.webp'), '埃欧珞灵动跳跃Rs')
TB(['维度', '详情'],
    [['总部', '杭州，北京分公司总经理雷宇峰'],
     ['核心产品', '灵动跳跃Rs（蛇形风力推进）\n磐石Rx（真空吸附式）'],
     ['产品售价', '数十万元/台，3至4单回本'],
     ['销售模式', '直销2台起、代理10台起、可租赁、省市级代理制'],
     ['业务拓展', '2023年参加广州光伏展，拓展至光伏面板清洁']])
P('来源：快鲤鱼报道、清洁话品牌报道', s=9)

# ---- 竞品3: 凌度智能 ----
P('竞品3：凌度（广东）智能科技发展有限公司', b=True)
add_img(os.path.join(IMG, '凌度智能_0.png'), '凌度智能产品图')
add_img(os.path.join(IMG, '凌度智能_1.jpg'), '凌度智能产品应用场景')
TB(['维度', '详情'],
    [['总部', '广东'],
     ['核心产品', 'X-Human系列高空幕墙清洗机器人\n凌空K3高空幕墙清洗\n凌净J1 SE低空清洁'],
     ['荣誉', '广东省名优高新技术产品称号'],
     ['渠道布局', '大湾区核心，区域代理制']])
P('来源：凌度智能科技公众号', s=9)

# ---- 竞品4: 华蔚 ----
P('竞品4：华蔚科技（福建）', b=True)
add_img(os.path.join(IMG, '华蔚_清洗机器人_0.png'), '华蔚产品图')
if os.path.exists(os.path.join(IMG, '华蔚_清洗机器人_1.png')):
    add_img(os.path.join(IMG, '华蔚_清洗机器人_1.png'), '华蔚产品图')
TB(['维度', '详情'],
    [['总部', '福建'],
     ['核心产品', '第一代60米高空幕墙清洗机器人'],
     ['技术路线', '磁力吸附技术'],
     ['创始人', '刘昌臻']])
P('来源：搜狗微信搜索', s=9)

# ---- 其他竞品 ----
P('其他行业参与者', b=True)
TB(['公司', '产品方向', '特点'],
    [['万勋科技（北京）', '无人机清洗系统', '幕墙+光伏双线'],
     ['R-storm（上海）', '高空清洁机器人', '获千万级融资'],
     ['史河机器人BeeBot', '外墙清洗机器人', 'CCE展热点新品'],
     ['云未N3F53', '高空清洁机器人', 'CCE上海清洁展参展']])

BR()

# ====== MODULE 2 ======
H('模块二：客户与需求')
P('对应调研文档问题：客户最愿意为什么买单？')

H('2.1 谁是最佳客户')
TB(['层级', '客户类型', '典型代表', '决策特点'],
    [['S级', '超大型物业集团', '万科/华润/碧桂园/保利', '总部集中决策，注重品牌与安全'],
     ['A级', '酒店管理集团', '万豪/洲际/锦江/华住', '品牌形象+安全，愿为差异化买单'],
     ['A级', '大型清洁公司', '玉禾田等', '成本驱动，注重ROI'],
     ['B级', '政府及公建', '政务中心/机场/车站', '安全合规第一，预算有保障'],
     ['C级', '中小清洁公司', '各地清洗服务商', '价格敏感，租赁易接受']])

H('2.2 核心购买理由')
P('第一是安全替代。高空作业死亡率高、蜘蛛人保险问题突出，机器人可从根本上消除安全风险。多地政府正在加强安全监管，一旦限制人工的政策落地，机器人将从可选项变为必选项。')
P('第二是成本降低。人工成本持续上涨、用工缺口逐年扩大，机器人长期成本优势渐显。')
P('第三是品牌提升。引入机器人契合智慧物业趋势，是有效的差异化手段。')
P('第四是政策合规。多城市政府令强制清洗为设备采购提供了制度推力。')

H('2.3 决策障碍')
TB(['障碍', '客户顾虑', '应对策略'],
    [['价格太贵', '数十万一次性投入大', '租赁模式，将资本支出转运营支出'],
     ['效果存疑', '担心复杂幕墙洗不干净', '免费试用+现场演示'],
     ['操作困难', '怕设备复杂不会用', '完整培训+远程运维'],
     ['安全顾虑', '担心机器人掉落', '安全认证+产品责任险']])

BR()

# ====== MODULE 3 ======
H('模块三：自身产品与模式（哈工鹏泽·GE-02）')
P('对应调研文档问题：GE-02最不可替代的优势？租赁/销售/联营哪种最优？')

H('3.1 GE-02绝对优势')
P('36氪文章将越障功能列为行业最难解决的问题之一。如果GE-02的滚轮式越障技术突破此难题，将构成显著的竞争壁垒。36氪提出两种技术路线：通用型机器（更适合商用）和专用型机器（技术上更容易）。GE-02的滚轮式越障加全材质适配定位偏向通用型路线。')
P('同时，竞品以玻璃幕墙为主，GE-02对石材和金属幕墙的适配能力可以开辟竞争更少的差异化市场。中国30亿平方米幕墙中石材约占25%、金属约占15%，合计约12亿平方米，是竞品覆盖较少的空间。')

H('3.2 三种模式验证')
TB(['模式', '客户接受度', '现金流', '适用场景'],
    [['租赁', '最高', '持续稳定', '中小客户/首次尝试'],
     ['销售', '中等', '一次性回款好', '大客户/政府项目'],
     ['联营', '上升', '长期回报', '区域合作伙伴']])
P('参考埃欧珞数十万元每台的定价，租赁月租金可设定为客户可接受区间。初期建议以租赁为主打快速铺量。')

BR()

# ====== MODULE 4 ======
H('模块四：渠道与生态')
P('对应调研文档问题：有政府关系还是有清洗团队的伙伴？合作诉求？生态捆绑？')

H('4.1 最佳渠道伙伴')
TB(['渠道', '典型伙伴', '合作价值'],
    [['物业协会', '中物协/省市物业协会', '行业背书+客户资源'],
     ['物业集团', '万科/华润/碧桂园/保利', '直接采购+品牌效应'],
     ['区域服务商', '各地清洁工程公司', '本地化能力+客户关系'],
     ['政府合作', '街道/城市管理部门', '预算有保障'],
     ['行业展会', 'CCE上海清洁博览会', '品牌曝光+客户对接']])

H('4.2 合作诉求')
TB(['诉求', '重要性', '我方支持'],
    [['高利润率', '最高', '有竞争力代理价+返利'],
     ['独家代理权', '高', '按区域/行业授权'],
     ['技术培训', '中', '操作+维修+远程'],
     ['售后保障', '中', '快速响应网络']])

H('4.3 生态扫描')
TB(['方向', '可行性', '切入点'],
    [['智慧物业平台', '高', 'API对接，数据入大屏'],
     ['机器人保险', '中高', '专属责任险消除客户顾虑'],
     ['行业标准制定', '中', '尚无统一标准，先入者有机会']])

BR()

# ====== CONCLUSIONS ======
H('结论与行动建议')
H('核心发现')
P('1. 市场空间大、时机好。30亿平方米幕墙创造百亿级市场，政府令保障需求刚性，竞争格局尚未固化。')
P('2. 越障能力是核心竞争力，36氪将越障列为行业最难解决的问题之一。')
P('3. 石材和金属幕墙是差异化蓝海，竞品以玻璃为主。')
P('4. 安全替代和成本降低是核心驱动力。')
P('5. 租赁模式是最佳市场切入方式。')

H('下一步行动')
TB(['优先级', '行动', '内容'],
    [['P0', '内部数据核实', '确认产品参数和销售数据'],
     ['P0', '竞品价格摸底', '了解竞品实际成交价'],
     ['P1', '行业专家访谈', '联系物业协会专家'],
     ['P1', '现有客户回访', '回访3至5家客户'],
     ['P1', '销售团队访谈', '整理一线客户反馈'],
     ['P2', '渠道对接', '展会加协会']])

BR()

# ====== SOURCES ======
H('数据来源')
P('1. 36氪行业分析 https://36kr.com/p/1721639649281', s=9)
P('2. 搜狗微信搜索哈工鹏泽', s=9)
P('3. 搜狗微信搜索埃欧珞', s=9)
P('4. 搜狗微信搜索凌度智能', s=9)
P('5. 360图片搜索/Bing图片搜索', s=9)

# ====== SAVE ======
path = os.path.join(OUT, '高空清洗机器人市场快速调研报告.docx')
doc.save(path)

text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text
hz = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
print(f'Hanzi: {hz}')
print(f'Total chars: {len(text)}')

# Count images
img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_count += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Images: {img_count}')

pdf_path = os.path.join(OUT, '高空清洗机器人市场快速调研报告.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
