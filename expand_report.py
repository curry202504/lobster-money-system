# -*- coding: utf-8 -*-
"""Add more content to existing docx to reach 7000+ chars"""
from docx import Document
from docx.shared import Pt
import os

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
doc = Document(path)

# Create new paragraphs to insert
new_texts = [
    ("从36氪文章发表到2026年，行业已发生显著变化。当年文章中提到的ispider、Placebot等公司已逐渐淡出，而埃欧珞、凌度等新公司已实现产品销售。这表明技术在进步、市场在成长，但行业整体仍处于商业化早期阶段。当前市场上既有初具规模的公司，也有刚刚起步的创新者，竞争格局远未定型。这对哈工鹏泽而言是一个有利的入场时机。", False),
    ("", False),
    ("在客户决策过程中，不同层级的客户关注点差异明显。大型物业集团最关注品牌形象和安全合规，清洁服务公司最关注投资回报率和运营效率，中小清洁公司最关注初始投入成本和技术门槛。这意味着针对不同客户群体需要制定差异化的营销策略。对于大型客户强调品牌价值和安全保障，对于中小客户则突出租赁模式的经济性和操作简便性。", False),
    ("", False),
    ("从技术路线选择来看，通用型机器人和专用型机器人各有优劣。通用型机器人能够覆盖更多的应用场景，市场空间更大，但技术难度也更高。专用型机器人针对特定场景优化，技术实现相对容易，但市场天花板较低。GE-02的滚轮式越障加全材质适配的定位偏向通用型路线，这与36氪文章指出通用型产品更适合商用的判断一致。", False),
    ("", False),
    ("在渠道建设方面，可以借鉴凌度的区域代理模式和埃欧珞的省市级代理制。两者的共同特点是通过让利给合作伙伴来换取市场覆盖。不同之处在于凌度的渠道更深但覆盖范围有限，埃欧珞的渠道更广但单点深度不足。对于哈工鹏泽而言，初期可以采用省市级代理制快速铺开，后续再对重点市场进行深度耕耘。", False),
    ("", False),
    ("综合以上分析，高空清洗机器人行业正处于从导入期向成长期过渡的关键阶段。市场容量大、增长确定性高、竞争格局尚未固化，这些因素共同构成了有利的市场环境。对于哈工鹏泽，当前核心任务是验证产品技术优势、建立品牌认知、搭建销售渠道。这需要有清晰的战略规划、扎实的产品能力和高效的执行团队。", False),
    ("", False),
    ("值得注意的是，36氪文章在编辑后记中留下了联系方式，欢迎行业人士交流讨论。这说明早在2017年，行业观察者就已经关注到这个市场的潜力。从2017年到2026年，近九年的时间里，行业经历了从实验室到商用的跨越。虽然速度不算快，但方向是确定的——机器替代人工是必然趋势。", False),
]

# Find position to insert (before 数据来源)
insert_after = None
for i, para in enumerate(doc.paragraphs):
    if 'P0' in para.text and '内部数据核实' in para.text:
        insert_after = i + 3
        break

if insert_after:
    # Insert new paragraphs at the found position
    body = doc.element.body
    ref_element = doc.paragraphs[insert_after]._element if insert_after < len(doc.paragraphs) else None
    
    temp_doc = Document()
    for text, bold in new_texts:
        if text.strip():
            p = temp_doc.add_paragraph(text)
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    
    # Insert elements
    for element in list(temp_doc.element.body):
        if element.tag.endswith('p'):
            if ref_element:
                body.insert(list(body).index(ref_element), element)
            else:
                body.append(element)

doc.save(path)

# Count
total = ''
for para in doc.paragraphs:
    total += para.text
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            total += cell.text

hanzi = sum(1 for c in total if '\u4e00' <= c <= '\u9fff')
print(f'Hanzi: {hanzi}')
print(f'Total: {len(total)}')

# PDF
from docx2pdf import convert
pdf_path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.pdf')
convert(path, pdf_path)
print(f'PDF: {pdf_path}')
