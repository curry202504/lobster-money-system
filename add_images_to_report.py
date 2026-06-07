# -*- coding: utf-8 -*-
"""Insert competitor images into existing report"""
import os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from docx2pdf import convert

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
img_dir = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', 'report-images')

doc = Document(path)

# Find the competitor sections and insert images after them
# We'll insert after the paragraphs mentioning each competitor

insertions = []

for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    # Insert 哈工鹏泽 image after the paragraph mentioning 哈工鹏泽 product info
    if '哈工鹏泽（深圳）机器人技术有限公司（本公司）' in text:
        insertions.append((i + 2, '哈工鹏泽'))  # insert 2 paragraphs after
    
    # Insert 埃欧珞 image
    if '埃欧珞（杭州）科技有限公司' in text:
        insertions.append((i + 2, '埃欧珞'))
    
    # Insert 凌度 image
    if '凌度（广东）智能科技发展有限公司' in text:
        insertions.append((i + 2, '凌度'))

# Find available images
image_files = {
    '哈工鹏泽': None,
    '埃欧珞': None,
    '凌度': None
}

# Search for best images
for f in os.listdir(img_dir):
    fpath = os.path.join(img_dir, f)
    if '哈工鹏泽' in f:
        sz = os.path.getsize(fpath)
        if image_files['哈工鹏泽'] is None or sz > os.path.getsize(image_files['哈工鹏泽']):
            image_files['哈工鹏泽'] = fpath
    elif '埃欧珞' in f:
        sz = os.path.getsize(fpath)
        if image_files['埃欧珞'] is None or sz > os.path.getsize(image_files['埃欧珞']):
            image_files['埃欧珞'] = fpath
    elif '凌度' in f:
        sz = os.path.getsize(fpath)
        if image_files['凌度'] is None or sz > os.path.getsize(image_files['凌度']):
            image_files['凌度'] = fpath

print("Images to insert:")
for k, v in image_files.items():
    if v:
        print(f"  {k}: {v} ({os.path.getsize(v)} bytes)")
    else:
        print(f"  {k}: NOT FOUND")

# Insert images in reverse order to preserve indices
for insert_idx, company in sorted(insertions, key=lambda x: -x[0]):
    img_path = image_files.get(company)
    if not img_path or not os.path.exists(img_path):
        print(f"  Skipping {company}: no image")
        continue
    
    # Get the paragraph to insert after
    if insert_idx >= len(doc.paragraphs):
        insert_idx = len(doc.paragraphs) - 1
    
    ref_para = doc.paragraphs[insert_idx]
    ref_element = ref_para._element
    
    try:
        img = Image.open(img_path)
        w, h = img.size
        max_w = 380
        ratio = min(1.0, max_w / w) if w > 0 else 1
        img_width = Cm(ratio * w * 0.035)
        
        # Create a new paragraph with the image
        new_para = doc.add_paragraph()
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = new_para.add_run()
        run.add_picture(img_path, width=img_width)
        
        # Move the new paragraph after ref_para
        ref_element.addnext(new_para._element)
        
        # Add caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = cap.add_run(company + '产品图')
        run2.font.size = Pt(8)
        
        ref_element.addnext(cap._element)
        
        print(f"  Inserted {company} image after paragraph {insert_idx}")
    except Exception as e:
        print(f"  Error inserting {company} image: {e}")

doc.save(path)

# Count chars
text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text
hz = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
print(f"\nHanzi: {hz}")
print(f"Total chars: {len(text)}")

pdf_path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.pdf')
convert(path, pdf_path)
print(f"PDF: {os.path.getsize(pdf_path)} bytes")
