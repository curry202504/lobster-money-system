# -*- coding: utf-8 -*-
"""Convert report pages to images and send via WeChat"""
import os, subprocess, json

docx = r"C:\temp\report.docx"
img_dir = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'media', 'inbound')
os.makedirs(img_dir, exist_ok=True)

# First try: convert docx to images using python-docx + PIL
from docx import Document
from PIL import Image, ImageDraw, ImageFont
from docx.shared import Inches

doc = Document(docx)
print(f"Document has {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# Create summary image with key findings
img = Image.new('RGB', (800, 1100), 'white')
draw = ImageDraw.Draw(img)

# Use default font
try:
    font_title = ImageFont.truetype("arial.ttf", 24)
    font_normal = ImageFont.truetype("arial.ttf", 16)
    font_small = ImageFont.truetype("arial.ttf", 12)
except:
    font_title = ImageFont.load_default()
    font_normal = ImageFont.load_default()
    font_small = ImageFont.load_default()

y = 20
lines = [
    ("高空清洗机器人市场快速调研报告", 24, True),
    ("哈工鹏泽（深圳）机器人技术有限公司", 18, False),
    ("", 10, False),
    ("════ 核心发现 ════", 16, True),
    ("", 10, False),
    ("1. 中国幕墙30亿㎡，清洗市场75亿~百亿级", 14, False),
    ("2. 机器人渗透率极低，处于爆发前夜", 14, False),
    ("3. 越障能力是行业公认技术难点，GE-02有壁垒", 14, False),
    ("4. 竞品以玻璃幕墙为主，石材/金属是蓝海", 14, False),
    ("5. 埃欧珞（数十万/台）、凌度（区域代理）已商用", 14, False),
    ("", 10, False),
    ("════ 市场数据 ════", 16, True),
    ("", 10, False),
    ("幕墙总面积：30亿㎡（年新增9000万㎡）", 14, False),
    ("保守估算：75亿元（3元/㎡×1次/年）", 14, False),
    ("北京：1次/年 | 上海：2-4次/年", 14, False),
    ("广州：2-4次/年 | 深圳：1次/年", 14, False),
    ("人工收费：3-5元/㎡", 14, False),
    ("", 10, False),
    ("════ 竞品概览 ════", 16, True),
    ("", 10, False),
    ("哈工鹏泽（本公司）：GE-02滚轮式越障", 14, False),
    ("埃欧珞（杭州）：灵动跳跃Rs+磐石Rx", 14, False),
    ("  → 售价数十万/台，直2台起/代10台起/可租赁", 14, False),
    ("凌度智能（广东）：凌空K3+X-Human系列", 14, False),
    ("  → 区域代理制，大湾区→港澳→东南亚", 14, False),
    ("", 10, False),
    ("════ 数据来源 ════", 16, True),
    ("", 10, False),
    ("36氪：https://36kr.com/p/1721639649281", 12, False),
    ("搜狗微信：哈工鹏泽/埃欧珞/凌度公众号", 12, False),
]

for text, size, is_bold in lines:
    try:
        font = ImageFont.truetype("arial.ttf", size)
    except:
        font = ImageFont.load_default()
    if is_bold:
        draw.text((30, y), text, fill=(0, 0, 0), font=font)
    else:
        draw.text((30, y), text, fill=(50, 50, 50), font=font)
    y += size + 8

img_path = os.path.join(img_dir, 'report_summary.png')
img.save(img_path, 'PNG')
print(f"Summary image saved: {img_path} ({os.path.getsize(img_path)} bytes)")

# Send the image via WeChat
result = subprocess.run([
    "openclaw", "message", "send",
    "--channel", "openclaw-weixin",
    "--target", "o9cq80xjeoJCl3_l5tV8LWGj0_1Q@im.wechat",
    "--media", img_path,
    "--message", "高空清洗机器人市场快速调研报告 - 核心摘要"
], capture_output=True, text=True, timeout=30)

print("Send result:", result.stdout[:500])
if result.stderr:
    print("Stderr:", result.stderr[:500])
