import os
from docx import Document
from docx.shared import Pt
from docx2pdf import convert

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
doc = Document(path)

more = '''关于报告的局限性。本报告作为桌面研究短期内完成，有以下几点局限需要注意。第一，核心数据来自2017年的36氪文章，虽然该文章是公开可查最详细的行业分析，但数据的时效性有一定局限。第二，竞品信息来自搜狗微信搜索结果，虽然可以复现搜索结果，但部分文章未能获取全文进行详细核实。第三，行业渗透率、企业数量等数据来自多方引用，精确度有待行业访谈验证。第四，图片来自搜索引擎下载，内容未做人工审核。第五，商业模式建议和定价策略基于行业逻辑推理，需要结合内部财务数据验证。

关于下一步工作的具体建议。建议首先完成内部数据核实形成GE-02的产品技术参数表和竞争优势分析文档。其次通过行业专家访谈验证市场规模和渗透率等核心假设。然后通过客户回访了解真实购买动机和决策障碍。综合以上信息后，再制定详细的市场进入策略和销售计划。整个过程预计需要2到3个月时间。'''

for line in more.strip().split('\n'):
    line = line.strip()
    if line:
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

doc.save(path)

text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text
hz = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
print(f'Hanzi: {hz}')
print(f'Total chars: {len(text)}')

pdf_path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
