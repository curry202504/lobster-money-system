import os
from docx import Document
from docx.shared import Pt

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
doc = Document(path)

p = doc.add_paragraph('希望本报告能为哈工鹏泽的市场决策提供有价值的参考。后续如有需要，可以针对特定模块做更深入的研究和分析。')
for r in p.runs:
    r.font.name = 'Calibri'
    r.font.size = Pt(11)

doc.save(path)

text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text
hz = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
print(f'Hanzi: {hz}')
print(f'Total chars: {len(text)}')
