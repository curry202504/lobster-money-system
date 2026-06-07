# -*- coding: utf-8 -*-
"""Report with timestamps and newest data"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
from docx2pdf import convert

OUT = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace')
IMG = os.path.join(OUT, 'report-images-safe')
S1 = 'https://36kr.com/p/1721639649281'  # 2017-06-23
S2 = 'https://news.qq.com/rain/a/20260423A0850O00'  # 2026-04-23

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

def H(t, lv=1):
    hs = doc.add_heading(t, level=lv)
    for r in hs.runs: r.font.name = 'Calibri'
    return hs

def P(t, b=False, s=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t)
    r.font.name = 'Calibri'
    if b: r.bold = True
    if s: r.font.size = Pt(s)
    return p

def TB(hd, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hd))
    t.style = 'Light Shading Accent 1'
    for i, h in enumerate(hd):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(v)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph(); return t

def add_img(filename, caption):
    path = os.path.join(IMG, filename)
    if os.path.exists(path):
        try:
            img = PILImage.open(path); w, h = img.size
            ratio = min(1.0, 400 / w) if w > 0 else 1
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(); run.add_picture(path, width=Cm(ratio * w * 0.035))
        except: pass
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.font.size = Pt(8); r.italic = True

def BR(): doc.add_page_break()

# ====== COVER ======
for _ in range(3): doc.add_paragraph()
P('高空清洗机器人市场快速调研报告', b=True, s=26)
P('哈工鹏泽（深圳）机器人技术有限公司', b=True, s=16)
P('2026年5月', s=12)
P('')
P('每条数据均标注来源时间 ⏱️', s=10)
BR()

# ====== CH1 ======
H('一、市场与竞争')

H('1.1 市场规模（含时间标注）')
P('以下数据按时间排序，越新的越优先参考：')

TB(['数据', '数值', '时间', '来源'],
    [['全国幕墙清洗市场规模预测（2025）', '520亿元', '2024年8月报告', '豆丁网行业报告'],
     ['全国幕墙清洗市场规模预测（2025）', '150亿元，同比+25%', '2024年10月报告', '博研咨询/豆丁网'],
     ['全国幕墙清洗市场规模预测（2026）', '180亿元', '2024年10月报告', '博研咨询/豆丁网'],
     ['全国幕墙工程产值', '1516亿元', '2025年2月', '搜狐报道'],
     ['中国幕墙总面积', '约30亿㎡', '2017年6月', '36氪：500亿㎡×5%'],
     ['每年新建幕墙', '约9000万㎡', '2017年6月', '36氪'],
     ['清洗市场保守估算', '约75亿元+', '2017年6月', '36氪：3元/㎡×1次/年']])

P('说明：2017年36氪的75亿元是保守估算基准（仅玻璃幕墙、最低价3元/㎡、1次/年）。实际按一线城市2-4次/年、3-5元/㎡计算，目前行业预测已达150-180亿元（2025-2026年）。', s=9)

TB(['城市清洗政策', '频率', '时间', '来源'],
    [['北京', '至少1次/年，违者处罚', '2017年6月', '36氪引用北京市政府令'],
     ['上海', '2-4次/年', '2017年6月', '36氪引用上海市规定'],
     ['广州', '重点区4次/年，一般区2次/年', '2017年6月', '36氪引用广州市政府文件'],
     ['深圳', '至少1次/年', '2017年6月', '36氪引用深圳市政府令'],
     ['天津', '每年2次', '2017年6月', '36氪原文'],
     ['杭州', '每年1次', '2017年6月', '36氪原文'],
     ['南京', '每年1次，特殊活动统一清洗', '2017年6月', '36氪原文']])

H('1.2 行业痛点（2026年最新报道验证）')
P('2026年4月23日，腾讯新闻报道哈工鹏泽发布会（' + S2 + '），指出行业现状：')
P('1. 行业仍大量依赖人工吊板作业，工人坐木板上靠两根绳子悬在百米高空')
P('2. 每年坠亡、被风吹撞、绳索磨损断裂事故不罕见')
P('3. 一线城市日薪500-800元但招人越来越难')
P('4. 幕墙越来越复杂，传统机器人越障能力不够')
P('5. 报道确认哈工鹏泽产品已在深圳、苏州、武汉地标建筑投入使用')
P('6. 罗湖区政府提供产业资金和科技创新政策支持')
P('这些痛点与36氪2017年描述一致，说明行业长期存在但未解决，而哈工鹏泽的技术突破正在改变这一局面。')

H('1.3 核心竞品')
add_img('哈工鹏泽_GE02_滚轮越障机器人.jpg', '哈工鹏泽GE-02滚轮式越障高空清洗机器人')
P('哈工鹏泽（深圳）机器人技术有限公司（本公司）。2026年4月23日发布4款产品。滚轮式越障机器人跨越450mm障碍物。总设计师袁立鹏教授。产品已在深圳、苏州、武汉地标建筑使用。发布会获罗湖区政府支持。来源：腾讯新闻2026年4月。')

add_img('埃欧珞_高空幕墙清洁机器人_产品图.png', '埃欧珞高空幕墙清洁机器人')
P('埃欧珞（杭州）。灵动跳跃Rs + 磐石Rx。售价数十万元/台。直销2台起/代理10台起/可租赁。2023年拓展至光伏清洁。来源：快鲤鱼报道。')

P('凌度（广东）。X-Human系列获广东省名优高新技术产品称号。来源：凌度智能公众号。')

BR()

# ====== CH2 ======
H('二、客户需求（按时间标注）')

TB(['客户/场景', '具体需求', '时间', '来源'],
    [['深圳/苏州/武汉地标建筑', '哈工鹏泽GE-02已实际投入使用', '2026年4月', '腾讯新闻发布报道'],
     ['深圳罗湖区政府', '智能机器人为重点产业，释放100+AI场景，支持哈工鹏泽发布会', '2026年4月', '腾讯新闻'],
     ['深圳光明区街道', '与哈工鹏泽合作开展智能高空清洗推介', '2026年5月', '搜狗微信文章'],
     ['深圳物业公司（政策驱动）', '深圳政府令要求至少每年清洗1次外墙', '2017年6月（政府令有效）', '36氪引用深圳政府令'],
     ['全国物业公司（通用）', '人工蜘蛛人事故率高、招工难、日薪500-800元招不到人', '2026年4月验证', '腾讯新闻'],
     ['万科物业', '王石表态：未来十年30%岗位或被机器人取代', '2021年8月', '知乎引用'],
     ['全国幕墙清洗行业', '传统机器人越障能力不足是行业瓶颈', '2026年4月验证', '腾讯新闻']])

P('注：表中"地标建筑"具体名称未在公开报道中披露，如需确认可联系哈工鹏泽销售团队。', s=9, italic=True)
P('注：深圳政府令2017年已被36氪引用，作为地方性法规持续有效。', s=9, italic=True)

BR()

# ====== CH3 ======
H('三、产品与模式')
P('GE-02滚轮式越障450mm，解决36氪（2017年）和腾讯新闻（2026年）双认定的行业技术瓶颈。动力系统源自航天器/船舶装备技术。商业模式参考埃欧珞：数十万元/台，支持租赁和代理。')

BR()

# ====== CH4 ======
H('四、渠道与生态')
TB(['渠道', '验证状态', '时间'],
    [['深圳罗湖区政府', '已支持哈工鹏泽发布会', '2026年4月'],
     ['深圳光明区街道', '已合作开展推介活动', '2026年5月'],
     ['CCE上海清洁博览会', '2026年拟邀大疆等行业公司参会', '2026年'],
     ['政策利好', '多城市政府令强制清洗，持续有效', '持续']])

BR()

# ====== CH5 ======
H('五、结论')
P('1. 市场百亿级且持续增长（2025年预测150-180亿元，2024-2025年数据）')
P('2. 越障450mm是核心竞争力，36氪（2017）和腾讯新闻（2026）均确认此为行业技术瓶颈')
P('3. 已有实际客户验证：深圳/苏州/武汉地标建筑投入使用（腾讯新闻2026年4月证实）')
P('4. 渠道已验证：罗湖区政府产业支持、光明区街道政企合作')
P('5. 安全替代+成本降低是核心驱动力')

path = os.path.join(OUT, '高空清洗机器人市场快速调研报告.docx')
doc.save(path)

t = ''
for p in doc.paragraphs:
    t += p.text
for tb in doc.tables:
    for r in tb.rows:
        for c in r.cells:
            t += c.text
hz = sum(1 for ch in t if '\u4e00' <= ch <= '\u9fff')
print(f'Hanzi: {hz}')
print(f'Total: {len(t)}')

img_c = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_c += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Images: {img_c}')

pdf_path = path.replace('.docx', '.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
