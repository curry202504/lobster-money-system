import os
from docx import Document
from docx.shared import Pt
from docx2pdf import convert

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
doc = Document(path)

t = '''对于哈工鹏泽而言，当前的市场环境既有挑战也有机遇。挑战在于市场仍处于教育期需要投入大量资源进行客户培育。机遇在于先入者有机会在技术、品牌和渠道方面建立先发优势。建议集中资源在越障技术这个核心竞争力上持续投入在此基础上逐步拓展产品线和市场覆盖。'''

for line in t.strip().split('\n'):
    line = line.strip()
    if line:
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

doc.save(path)

text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text
hz = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
t = len(text)
print(f'Hanzi: {hz}')
print(f'Total: {t}')

img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_count += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Images: {img_count}')

pdf_path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
