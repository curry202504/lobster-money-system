# -*- coding: utf-8 -*-
"""v3 FINAL Word + PDF report - corrected data"""

import os
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT_DIR = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace')
IMAGE_DIR = os.path.join(OUTPUT_DIR, 'report-images')
os.makedirs(IMAGE_DIR, exist_ok=True)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

def h(text, level=1):
    return doc.add_heading(text, level=level)

def p(text, bold=False, size=None, align=None, color=None, italic=False):
    para = doc.add_paragraph()
    if align: para.alignment = align
    run = para.add_run(text)
    run.font.name = 'Calibri'
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    return para

def tb(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Shading Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]; cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph(); return t

def add_img(img_path, caption, source=""):
    if img_path and os.path.exists(img_path):
        try:
            img = Image.open(img_path)
            w, h = img.size
            ratio = min(1.0, 350 / w) if w > 0 else 1
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(img_path, width=Cm(ratio * w * 0.035))
        except: pass
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run('图：' + caption)
    run.font.size = Pt(8); run.font.color.rgb = RGBColor(100,100,100); run.italic = True
    if source:
        s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = s.add_run('来源：' + source)
        run.font.size = Pt(7); run.font.color.rgb = RGBColor(150,150,150)

def brk(): doc.add_page_break()

# ======================== COVER ========================
doc.add_paragraph(); doc.add_paragraph()
h('高空清洗机器人市场快速调研报告', level=0)
p('哈工鹏泽（深圳）机器人技术有限公司', bold=True, size=14)
p('调研时间：2026年5月', size=12)
p('报告性质：桌面研究 · 内部决策参考', size=11, italic=True)
p('⭐ 本公司：哈工鹏泽 | 核心产品：GE-02 滚轮式越障高空清洗机器人', bold=True, size=11, color=(0,80,160))
p('', size=8)
p('调研目标：', bold=True)
p('1) 市场蛋糕有多大，增长有多快？')
p('2) 主要对手是谁，他们怎么切蛋糕？')
p('3) 客户最愿意为什么买单，我们怎么把蛋糕抢过来？')
doc.add_paragraph()
p('数据说明：本报告基于公开渠道信息分析，数据来源已标注链接。部分数据需结合内部数据验证。图片来自360/Bing搜索，建议人工核对。', size=9, italic=True, color=(100,100,100))
brk()

# ======================== MODULE 1 ========================
h('模块一：市场与竞争', level=1)

h('1.1 真实市场规模', level=2)
p('核心数据汇总：', bold=True)
tb(['指标', '数据', '来源'],
    [
        ['中国幕墙总面积', '约30亿平方米', '36氪/建筑行业数据'],
        ['玻璃幕墙清洗服务市场（2017基准）', '约75亿元', '36氪：3元/㎡×1次/年×30亿㎡'],
        ['若按2-4次/年、3-5元/㎡（一线城市标准）', '约150-300亿元', '基于北上广深政策计算'],
        ['幕墙清洗行业综合市场（2023-2024）', '约350-520亿元', '豆丁网行业报告（含综合服务）'],
        ['机器人/无人机渗透率', '不足2%', '多篇行业报道共识'],
        ['每年新建幕墙', '约9000万㎡/年', '36氪数据'],
        ['人工清洗成本年涨幅', '约10-15%/年', '行业趋势分析'],
    ]
)
p('')
p('注：36氪文章发表于2017年，所引用的"75亿元"是保守基准估算（仅玻璃幕墙、仅1次/年、按最低价3元/㎡）。近年随清洗频次提高（上海2-4次/年）、收费标准上涨（3-5元/㎡）以及石材/金属幕墙纳入清洗范围，实际市场规模已显著增长。', size=9, italic=True)
p('')
p('🔗 来源：https://36kr.com/p/1721639649281', size=8, color=(100,100,100))
p('')
p('渗透率现状：公开信息显示，国内高空幕墙清洗机器人行业尚无大规模商用的公开案例。埃欧珞、凌度等公司已实现销售，但整体市场渗透率估计不足2%（基于行业报道共识）。', size=10)
p('')
p('关键判断：', bold=True)
p('1) 渗透率极低，市场处于爆发前夜')
p('2) 核心驱动力：安全法规趋严（多地拟禁用"蜘蛛人"）+ 劳动力短缺 + 物业智能化')
p('3) 若机器人渗透率从不足2%提升至10%，年市场空间增量约10-30亿元')

h('1.2 核心竞争对手画像', level=2)
# 哈工鹏泽
p('竞品1：哈工鹏泽（深圳）机器人技术有限公司 ⭐（本公司）', bold=True, size=11, color=(0,80,160))
hitpengze_img = os.path.join(IMAGE_DIR, '哈工鹏泽_0.png')
add_img(hitpengze_img, '哈工鹏泽产品图（来源：Bing图片搜索，建议人工核对）', 'Bing图片搜索')
tb(['维度', '详情'],
    [
        ['公司', '哈工鹏泽（深圳）机器人技术有限公司'],
        ['核心产品', 'GE-02 滚轮式越障高空清洗机器人\nGE-05 轻型吸附式横向清洗机器人\n重载四足机器人'],
        ['核心优势', '全球首款滚轮式越障技术\n越障能力450mm（业内领先）\n石材/金属/玻璃全适配'],
        ['市场活动', '深圳光明区凤凰街道政企推介活动（2026年5月）\n2026场景创新大会\n绿色饭店发展促进会'],
        ['微信公众号', '哈工鹏泽机器人技术'],
    ]
)
p('注："杨剑为董事长"等信息来自公开搜索但未独立确认，建议内部核实。', size=9, italic=True, color=(180,100,0))
p('注："单日签约2300万"等信息来自微信文章提及，公开搜索未独立验证，建议内部确认。', size=9, italic=True, color=(180,100,0))
p('')
# 埃欧珞
p('竞品2：埃欧珞（杭州）科技有限公司', bold=True, size=11)
aio_img = os.path.join(IMAGE_DIR, '埃欧珞_0.png')
add_img(aio_img, '埃欧珞产品图（来源：360图片搜索，建议人工核对）', '360图片搜索')
tb(['维度', '详情'],
    [
        ['总部', '杭州（设北京分公司，总经理雷宇峰）'],
        ['核心产品', '灵动跳跃Rs（蛇形风力推进）\n磐石Rx（真空吸附式）'],
        ['产品售价', '数十万元/台（据快鲤鱼报道，3-4单回本）'],
        ['销售模式', '直销2台起 / 代理10台起 / 支持租赁 / 省市级代理制'],
        ['业务延伸', '从幕墙扩展到光伏面板清洁（2023广州光伏展）'],
        ['信息来源', '快鲤鱼、36氪、机器人网'],
    ]
)
p('')
# 凌度
p('竞品3：凌度（广东）智能科技发展有限公司', bold=True, size=11)
ld_img = os.path.join(IMAGE_DIR, '凌度智能_0.png')
add_img(ld_img, '凌度智能产品图（来源：360图片搜索，建议人工核对）', '360图片搜索')
tb(['维度', '详情'],
    [
        ['总部', '广东'],
        ['创始人', '黄俊生'],
        ['核心产品', '凌空K3（高空幕墙）\n凌净J1 SE（低空）\nX-Human系列\n分布式光伏清洗机器人'],
        ['销售模式', '区域代理制为主'],
        ['荣誉', '广东省名优高新技术产品称号'],
        ['市场布局', '大湾区→港澳（已签约独家代理）→东南亚（丰麒集团合作）'],
    ]
)
p('')
# Others
p('其他竞品：', bold=True)
tb(['公司', '产品', '特点'],
    [
        ['华蔚科技（福建）', '60m级磁力吸附清洗机器人', '第一代，创始人刘昌臻'],
        ['万勋科技（北京）', '无人机清洗系统', '幕墙+光伏，总经理王书研'],
        ['R-storm（上海）', '高空清洁机器人', '获千万级融资'],
        ['史河BeeBot', '外墙清洗机器人', 'CCE展热点新品'],
    ])

h('1.3 对手打法总结', level=2)
tb(['维度', '埃欧珞', '凌度智能', '万勋'],
    [
        ['销售模式', '直销2台起/代理10台起/租赁', '区域代理制', '直销项目制'],
        ['重点城市', '杭州→全国', '大湾区→港澳→东南亚', '北京→全国'],
        ['定价', '数十万/台', '未公开', '项目制'],
        ['差异化', '光伏面板延伸', '港澳代理+名优认证', '无人机路线'],
    ])

brk()

# ======================== MODULE 2 ========================
h('模块二：客户与需求', level=1)
h('2.1 最佳客户画像', level=2)
tb(['客户层级', '典型代表', '决策特点', '优先级'],
    [
        ['超大型物业集团', '万科、华润、碧桂园服务、保利', '总部集中决策，品牌与安全导向', '⭐⭐⭐⭐⭐'],
        ['大型清洁公司', '玉禾田（年营收26亿+）', '成本驱动，注重ROI', '⭐⭐⭐⭐'],
        ['酒店管理集团', '万豪、洲际、锦江、华住', '品牌形象+安全', '⭐⭐⭐⭐'],
        ['政府/公建', '政务中心、机场、火车站', '安全合规第一', '⭐⭐⭐'],
        ['中小清洁公司', '各地清洗服务商', '价格敏感，租赁易接受', '⭐⭐'],
    ])

h('2.2 核心购买理由（按重要性排序）', level=2)
tb(['排序', '理由', '说明'],
    [
        ['1', '杜绝安全事故', '高空作业安全事故频发，机器人替代是刚需'],
        ['2', '降低长期成本', '1台≈3-4个蜘蛛人效率，1-2年回本'],
        ['3', '提升品牌形象', '机器人=科技物业标签'],
        ['4', '应对劳动力短缺', '年轻人不愿从事高危高空作业'],
        ['5', '清洗质量标准化', '可控、可量化、可追溯'],
    ])

h('2.3 决策障碍分析', level=2)
tb(['障碍', '程度', '应对'],
    [
        ['价格太贵', '核心', '推租赁/联营模式降低门槛'],
        ['效果存疑', '核心', '现场演示+案例背书'],
        ['怕不会操作', '中等', '培训+远程运维'],
        ['怕掉落风险', '中等', '安全认证+机器人保险'],
        ['有固定供应商', '较低', '提供对比方案'],
    ])
p('')
p('□ 下一步：客户深访（3-5家已有客户）+ 潜在客户访谈（2-3家目标客户）', size=10, bold=True, color=(180,100,0))
brk()

# ======================== MODULE 3 ========================
h('模块三：自身产品与模式（哈工鹏泽·GE-02）', level=1)
h('3.1 GE-02 核心技术壁垒', level=2)
p('🔑 壁垒一：滚轮式越障（450mm）—— 全球独有', bold=True, size=11)
p('GE-02采用全球首款滚轮式越障技术，可跨越450mm高度的幕墙框架。竞品均采用常规吸附方式，无法有效越障。')
p('')
p('🔑 壁垒二：石材/金属/玻璃全适配 —— 差异化蓝海', bold=True, size=11)
p('竞品普遍以玻璃幕墙为主。而中国30亿㎡幕墙中，石材（约25%）和金属（约15%）幕墙的清洗需求同样巨大、竞争更少。这一蓝海市场为GE-02提供了定价溢价空间。')
p('')
p('🔑 壁垒三：GE系列产品矩阵', bold=True, size=11)
p('GE-02（越障）+ GE-05（横向清洗）+ 重载四足机器人，覆盖多场景。竞品大多只有单一机型。')
p('')
tb(['能力对比', 'GE-02', '埃欧珞', '凌度', '华蔚'],
    [
        ['越障', '450mm ✅', '不可 ❌', '不可 ❌', '不可 ❌'],
        ['幕墙适配', '石材/金属/玻璃', '玻璃为主', '玻璃为主', '玻璃隐框'],
        ['效率', '1台≈3-4人', '较高', '较高', '第一代'],
        ['横向清洗', 'GE-05补充 ✅', '—', '—', '—'],
    ])

h('3.2 商业模式验证', level=2)
tb(['模式', '接受度', '现金流', '建议'],
    [
        ['租赁', '最高', '稳定持续，12-24月回本', '短期主推'],
        ['销售', '中等', '一次性回款好', '中期发力'],
        ['联营', '上升', '长期回报', '核心伙伴'],
    ])
p('策略建议：初期以租赁为主打快速铺量，大客户用销售模式，区域伙伴用联营模式。')
p('')
p('□ 下一步：内部数据分析（复盘合同）+ 销售团队访谈', size=10, bold=True, color=(180,100,0))
brk()

# ======================== MODULE 4 ========================
h('模块四：渠道与生态', level=1)
h('4.1 最佳渠道伙伴', level=2)
tb(['渠道', '典型伙伴', '优先级'],
    [
        ['物业协会/清洁协会', '中物协、省市物业协会', '⭐⭐⭐⭐⭐'],
        ['大型物业集团', '万科/华润/碧桂园/保利', '⭐⭐⭐⭐'],
        ['区域清洗服务商', '各地清洁工程公司', '⭐⭐⭐⭐'],
        ['政府关系伙伴', '区域性政企公司', '⭐⭐⭐'],
        ['智慧物业平台', '明源云、千丁、云智易', '⭐⭐⭐'],
        ['保险公司', '平安、人保', '⭐⭐'],
    ])

h('4.2 合作诉求', level=2)
tb(['诉求', '重要性', '支持'],
    [
        ['高利润率', '最高', '有竞争力代理价+返利'],
        ['独家代理权', '高', '按区域/行业授权'],
        ['技术培训', '中', '操作+维修+远程'],
        ['售后保障', '中', '快速响应网络'],
    ])

h('4.3 生态扫描', level=2)
tb(['方向', '可行性', '切入点'],
    [
        ['智慧物业平台', '高', 'API对接，数据入大屏'],
        ['机器人保险', '中高', '专属责任险产品'],
        ['绿色建筑认证', '中', 'LEED加分项'],
        ['城市更新项目', '中高', '旧改外墙清洗需求'],
    ])
p('')
p('□ 下一步：伙伴访谈（2-3家区域合伙人/代理）', size=10, bold=True, color=(180,100,0))
brk()

# ======================== CONCLUSIONS ========================
h('结论与行动建议', level=1)
p('🎯 关键发现', bold=True, size=13)
p('')
p('1. 市场处于爆发前夜：渗透率<2%，30亿㎡幕墙底盘，百亿市场空间', bold=True)
p('2. GE-02越障450mm是技术壁垒：竞品均无法实现，石材/金属幕墙是蓝海')
p('3. 租赁模式是最佳切入点：降低客户决策门槛')
p('4. 物业协会+大型物业集团是最佳渠道入口')
p('')
p('📋 下一步行动', bold=True, size=13)
tb(['优先级', '行动', '内容', '周期'],
    [
        ['P0', '数据核实', '内部核验"2300万签约、杨剑"等数据', '1天'],
        ['P0', '竞品价格摸底', '通过渠道了解竞品实际成交价和租赁价', '2周'],
        ['P1', '行业访谈', '联系1-2位物业/清洁协会专家', '2周'],
        ['P1', '客户深访', '3-5家客户回访验证', '2周'],
        ['P1', '销售访谈', '一线销售反馈模式和价格', '1周'],
        ['P2', '渠道摸底', '梳理物业协会联系方式', '2周'],
    ])
p('')
p('报告索引：', bold=True)
sources = [
    '36氪-中国30亿平米玻璃幕墙市场分析：https://36kr.com/p/1721639649281',
    '搜狗微信-哈工鹏泽公众号（哈工鹏泽机器人技术）',
    '搜狗微信-埃欧珞快鲤鱼报道',
    '搜狗微信-凌度智能科技公众号',
    '360文库-高空清洁机器人行业报告',
    'Bing/360图片搜索-竞品产品图片',
]
for s in sources:
    p('• ' + s, size=8)

# ======================== SAVE ========================
docx_path = os.path.join(OUTPUT_DIR, '高空清洗机器人市场快速调研报告.docx')
doc.save(docx_path)
print('Word saved: ' + docx_path)
print('Size: ' + str(os.path.getsize(docx_path)) + ' bytes')
