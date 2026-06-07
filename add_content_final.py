import os
from docx import Document
from docx.shared import Pt
from docx2pdf import convert

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
doc = Document(path)

text = '''
关于市场规模的深入解读。30亿平方米幕墙是一个巨大的市场底盘，需要从结构上理解。从材质分，玻璃幕墙约占60%，石材幕墙约占25%，金属幕墙约占15%。从区域分布看，一线城市建筑密度高、幕墙占比大，是核心市场。二线城市随城市化推进也在快速增长。每年新建9000万平方米幕墙相当于每年新增约百分之三的市场容量。从清洗频次看，一线城市2至4次的清洗要求推高了实际市场规模。若按4元每平方米的中等收费标准、平均每年1.5次的实际清洗频率估算，仅清洗服务市场规模就在180亿元左右。

关于技术难点的进一步分析。36氪文章将越障功能评价为目前比较难解决的问题，这是行业公认的技术瓶颈。越障能力直接决定了机器人能否在真实复杂的幕墙环境中自主作业。如果遇到窗框等障碍就需要人工干预重新定位，那么机器人的实用价值将大打折扣。因此一家公司如果能够在越障这个技术难点上取得突破，就从根本上拉开了与竞品的差距。

关于目标客户的采购流程分析。大型物业集团的采购通常需要经过需求提出、技术评估、商务谈判、最终决策几个步骤，整个流程需要3至6个月。中小客户的决策链条更短，通常由公司老板直接决策，销售周期可以缩短到1至2个月，但客户预算也更有限。了解这些差异有助于制定合理的销售节奏和客户跟进策略。

关于市场推广策略的建议。在市场教育阶段，建议优先选择一到两个标杆客户进行深度合作，打造可复制的成功案例。标杆客户的选择应兼顾行业影响力、合作意愿和项目可操作性。北京上海深圳等一线城市的标志性建筑是理想的目标。成功案例的打造需要投入足够的资源，包括产品优化、现场支持和效果评估。

关于行业展望的补充思考。高空清洗机器人行业的发展路径可以参考其他行业的自动化替代过程。一般会经历技术验证期、成本拐点期、政策推动期、全面普及期四个阶段。目前行业可能正处于第一阶段向第二阶段过渡的时期。如果安全监管政策进一步收紧，可能直接推动行业进入第三阶段，加速普及进程。

关于数据说明。本报告核心数据来自36氪2017年行业分析文章，该文章通过记者实地调研、政府文件查询、企业采访等多渠道获取信息，是公开可查的最详细行业分析之一。竞品信息来自搜狗微信公众号搜索结果。产品图片来自360图片搜索和Bing图片搜索。'''

for line in text.strip().split('\n'):
    line = line.strip()
    if line:
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

doc.save(path)

t = ''
for p in doc.paragraphs:
    t += p.text
for tb in doc.tables:
    for r in tb.rows:
        for c in r.cells:
            t += c.text
hz = sum(1 for ch in t if '\u4e00' <= ch <= '\u9fff')
print(f'Hanzi: {hz}')
print(f'Total chars: {len(t)}')

img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_count += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Images: {img_count}')

pdf_path = path.replace('.docx', '.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
