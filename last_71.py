import os
from docx import Document
from docx.shared import Pt

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
doc = Document(path)

p = doc.add_paragraph('本报告旨在为哈工鹏泽的市场决策提供参考，报告中的分析和建议仅供参考，具体商业决策请结合实际情况综合判断。')
for r in p.runs:
    r.font.name = 'Calibri'
    r.font.size = Pt(11)

doc.save(path)

text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text
hz = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
t = len(text)
print(f'Hanzi: {hz}')
print(f'Total: {t}')

img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_count += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Images: {img_count}')
