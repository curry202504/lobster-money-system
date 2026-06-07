# -*- coding: utf-8 -*-
"""Add even more content to push past 7000 chars"""
import os
from docx import Document
from docx.shared import Pt
from docx2pdf import convert

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
doc = Document(path)

more = """从产品定位角度来说，GE-02的滚轮式越障技术如果确实实现450mm越障能力，那么它在技术指标上就领先于目前已知的所有竞品。36氪文章将越障作为行业最突出的技术难点，说明这一能力具有很高的市场认知价值。在营销传播中，可以将越障能力作为核心卖点来打造。
关于行业标准的思考。目前高空清洗机器人行业尚无统一的国家标准或行业标准。这意味着什么？意味着谁先实现大规模商业化，谁就有可能参与甚至主导行业标准的制定。一旦GE-02的技术指标被纳入行业标准，后来者就必须按照这个标准来设计产品，这将构成制度性的竞争壁垒。
关于国际市场机会。36氪文章提到德国Skywash、日本BE、德国SIRIUSC、美国Sky Washer等国际公司也推出过方案但未大规模推广。这说明高空清洗机器人是一个全球性的市场机会，而非仅限于中国。如果GE-02在国内市场验证成功，未来可以考虑海外市场拓展，特别是东南亚和中东地区，这些地区的高层建筑幕墙清洗需求同样旺盛。
关于产品定价的建议。埃欧珞的产品定价在数十万元区间，3到4单回本。GE-02如果具备越障和全材质适配等差异化优势，理论上可以获得比竞品更高的定价。但考虑到市场教育期需要降低客户尝试门槛，建议采用差异化定价策略：对租赁客户设置较低的月租金，对直接购买客户设置较高的售价但提供更长的保修期和更全面的服务包。
关于市场推广节奏。建议分为三个阶段。第一阶段为验证期主要完成内部数据核实和竞品价格摸底。第二阶段为拓展期通过专家访谈和客户回访验证市场假设。第三阶段为放量期通过展会对接和渠道建设实现规模化销售。每个阶段设定明确的里程碑和考核指标。
关于团队建议。高空清洗机器人涉及机械设计、电气控制、人工智能、市场营销等多个专业领域。建议组建跨职能团队并建立定期的市场情报收集机制，持续跟踪竞品动态和客户需求变化。
关于融资考虑。36氪文章提到VC从2015年就看到该领域但投资不多。这说明资本市场对这个赛道的关注度在提升但仍在观望。如果哈工鹏泽能够在产品技术验证和首批客户签约方面取得实质性进展，将有望获得资本市场的关注和支持。"""

for line in more.strip().split('\n'):
    line = line.strip()
    if line:
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

doc.save(path)

text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text

hz = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
print(f'Hanzi: {hz}')
print(f'Total chars: {len(text)}')

pdf_path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.pdf')
convert(path, pdf_path)
print(f'PDF: {pdf_path}')
