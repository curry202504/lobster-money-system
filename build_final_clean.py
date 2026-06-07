# -*- coding: utf-8 -*-
"""FINAL CLEAN REPORT - Only data I'm confident about, no labels"""

import os
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace')
IMG = os.path.join(OUT, 'report-images')
os.makedirs(IMG, exist_ok=True)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

def h(text, level=1): return doc.add_heading(text, level=level)

def p(text, bold=False, size=None, color=None, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Calibri'
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    return para

def tb(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Shading Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for para in c.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for para in c.paragraphs:
                for r in para.runs: r.font.size = Pt(9)
    doc.add_paragraph()
    return t

def add_img(path, caption):
    if path and os.path.exists(path):
        try:
            img = Image.open(path); w, h = img.size
            ratio = min(1.0, 350 / w) if w > 0 else 1
            para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(); run.add_picture(path, width=Cm(ratio * w * 0.035))
        except: pass
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption); run.font.size = Pt(8); run.font.color.rgb = RGBColor(100,100,100); run.italic = True

def brk(): doc.add_page_break()

# ======================== COVER ========================
for _ in range(3): doc.add_paragraph()
p('高空清洗机器人市场快速调研报告', bold=True, size=22)
p('哈工鹏泽（深圳）机器人技术有限公司', bold=True, size=14, color=(0,80,160))
p('2026年5月', size=11, color=(100,100,100))
brk()

# ======================== MODULE 1 ========================
h('一、市场与竞争', level=1)

h('1.1 市场规模', level=2)
p('数据来源：36氪行业分析文章（2017年）。原文通过对建筑面积、政策法规、企业报价等多方面调研整理得出以下数据。')
p('来源：https://36kr.com/p/1721639649281', size=8, color=(100,100,100))

tb(['指标', '数据', '说明'],
    [
        ['中国建筑幕墙总面积', '约30亿㎡', '全国建筑面积500亿㎡，幕墙占比约5%'],
        ['每年新建幕墙', '约9,000万㎡', '持续增长的增量市场'],
        ['市场规模（保守估算）', '约75亿元以上', '3元/㎡ × 每年1次清洗 × 30亿㎡幕墙'],
        ['（若含粉墙等延伸服务）', '百亿元以上', '36氪原文：含可迁移的高空墙体粉刷'],
        ['北京清洗规定', '至少每年1次', '已列入政府令'],
        ['上海清洗规定', '每年2-4次', '高于全国平均水平'],
        ['广州清洗规定', '每年2-4次', '重点维护区4次，一般维护区2次'],
        ['深圳清洗规定', '至少每年1次', '已列入政府令'],
        ['人工清洗收费标准', '3-5元/㎡', '北京物业报价，依层高不同'],
    ]
)

p('')
p('行业现状：36氪调研时指出，大多数相关公司"处于产品研发或者产品研发刚刚结束的阶段，均未商用"。目前埃欧珞、凌度等公司已实现商用，行业已有明显进展。', size=10)
p('')
p('行业痛点（来源：36氪同文）：', bold=True)
p('1. 人工高空作业危险性高——"蜘蛛人"属高风险职业，媒体报道死亡率较高，部分工人无保险')
p('2. 招工难——高危职业用工缺口逐年扩大，年轻人不愿从事')
p('3. 现有清洁方式效率低——蜘蛛人效率低、升降平台成本高、轨道系统需预先安装')
p('4. 政策日趋严格——多城市以政府令强制规范清洗周期，物业需按规定执行')

h('1.2 核心竞品', level=2)

p('哈工鹏泽（深圳）机器人技术有限公司', bold=True, size=12, color=(0,80,160))
p('本公司。近期活动信息可通过微信公众号"哈工鹏泽机器人技术"查询。')
p('公开活动：2026年5月走进深圳光明区凤凰街道开展智能高空清洗推介活动；出席2026场景创新大会（深港同心·罗湖创景）；参展深圳绿色饭店发展促进会。')
p('')
p('埃欧珞（杭州）科技有限公司', bold=True, size=11)
p('总部杭州，设北京分公司。核心产品为灵动跳跃Rs（蛇形风力推进）和磐石Rx（真空吸附式），售价数十万元/台，3-4单回本。销售模式为直销2台起售、代理10台起售，支持租赁。2023年起从幕墙清洗拓展至光伏面板清洁。')
p('来源：快鲤鱼报道、清洁话品牌报道', size=8, color=(100,100,100))
p('')
p('凌度（广东）智能科技发展有限公司', bold=True, size=11)
p('总部广东。产品包括凌空K3（高空幕墙清洗）等系列，X-Human系列高空幕墙清洗机器人获广东省名优高新技术产品称号。市场布局覆盖大湾区并向东南亚拓展。')
p('来源：凌度智能科技公众号', size=8, color=(100,100,100))
p('')
add_img(os.path.join(IMG, '凌度智能_0.png'), '凌度智能产品')

# 哈工鹏泽产品图
hgpz_img = os.path.join(IMG, '哈工鹏泽_ge02_0.jpg')
if not os.path.exists(hgpz_img):
    hgpz_img = os.path.join(IMG, '哈工鹏泽_0.png')
add_img(hgpz_img, '哈工鹏泽产品')

# 埃欧珞产品图
aio_img = os.path.join(IMG, '埃欧珞_rs_0.webp')
if not os.path.exists(aio_img):
    aio_img = os.path.join(IMG, '埃欧珞_0.png')
add_img(aio_img, '埃欧珞产品')

brk()

# ======================== MODULE 2 ========================
h('二、客户与需求', level=1)

h('2.1 客户类型', level=2)
tb(['客户类型', '特征', '需求驱动'],
    [
        ['物业集团/管理公司', '直接负责建筑外墙清洗的采购决策', '安全合规 + 成本控制 + 品牌形象'],
        ['专业清洁服务公司', '为客户提供外包清洗服务', '效率提升 + 用工替代 + 差异化服务'],
        ['政府/公共设施', '政务中心、车站、机场等公建', '安全合规第一，预算有保障'],
    ])

h('2.2 行业痛点与机器人价值', level=2)
tb(['痛点', '技术难点（36氪原文）', '说明'],
    [
        ['吸附功能', '机器人需牢固吸附在垂直墙面', '各竞品基本解决'],
        ['移动功能', '适应玻璃/金属/粉墙等多种材质\n曲面壁面行走', '不同材质需不同方案'],
        ['越障功能', '跨越窗框等障碍物\n被认为"目前比较难解决的问题"', '当前行业技术难点'],
        ['清洗功能', '有效清洁幕墙表面', '各竞品基本解决'],
    ])

h('2.3 购买驱动力', level=2)
tb(['驱动力', '说明'],
    [
        ['安全替代人工', '高空作业事故频发，机器人可从根本上消除人员安全风险'],
        ['降低成本', '人工成本年涨幅10-15%，机器人综合成本优势逐步显现'],
        ['劳动力替代', '年轻人不愿从事高危高空作业，用工缺口逐年扩大'],
        ['满足政策合规', '多城市已出台政府令强制清洗周期'],
    ])

brk()

# ======================== MODULE 3 ========================
h('三、产品与商业模式', level=1)

h('3.1 GE-02优势分析', level=2)
p('从行业公开信息分析：')
p('')
p('1. 越障能力是行业公认难点', bold=True)
p('36氪文章专门指出：越障功能"要求机器在移动的过程中，能够跨越窗框等障碍物……依然是目前比较难解决的问题。"如果GE-02的滚轮式越障技术确实突破了这个难题，将是显著的差异化优势。')
p('')
p('2. 石材/金属幕墙是潜在蓝海', bold=True)
p('公开信息显示埃欧珞、凌度等竞品以玻璃幕墙为主要目标。而全国30亿㎡幕墙中，除玻璃外还有石材和金属材质，这部分市场竞品覆盖较少。')
p('')
p('3. 竞品定价参考', bold=True)
p('埃欧珞产品售价数十万元/台，3-4单回本。可参考此区间制定定价策略。')

h('3.2 商业模式参考', level=2)
tb(['模式', '参考依据'],
    [
        ['销售', '埃欧珞提供直销/代理销售方案，说明市场存在直接购买需求'],
        ['租赁', '埃欧珞支持租赁模式，说明租赁是可行的市场切入点'],
        ['代理', '埃欧珞代理10台起售，凌度采用区域代理制，说明代理渠道有效'],
    ])

brk()

# ======================== MODULE 4 ========================
h('四、渠道与生态', level=1)

h('4.1 渠道方向', level=2)
tb(['渠道', '依据', '说明'],
    [
        ['物业集团', '36氪文章直接采访了物业公司，物业是采购决策方', '总部级集中采购，单品订单大'],
        ['清洁服务公司', '全国约1200家高空清洗服务企业（36氪提及）', '批量采购潜力大，需教育市场'],
        ['行业协会', '中国物业管理协会等组织', '行业背书+客户资源对接'],
        ['行业展会', 'CCE上海清洁博览会等', '接触客户和渠道商的有效平台'],
    ])

h('4.2 生态机会', level=2)
tb(['方向', '分析'],
    [
        ['政策驱动', '多城市政府令强制清洗周期，为行业提供稳定需求基础'],
        ['行业标准制定', '当前高空清洗机器人行业标准尚在空白期，先入场者有参与标准制定的机会'],
    ])

brk()

# ======================== CONCLUSIONS ========================
h('结论', level=1)

tb(['序号', '结论'],
    [
        ['1', '幕墙清洗市场有百亿级规模（2017年保守估算75亿元），渗透率极低，增长空间大'],
        ['2', '越障能力是行业公认的技术难点，GE-02若突破此难题即构成核心壁垒'],
        ['3', '竞品以玻璃幕墙为主，石材/金属幕墙是差异化机遇'],
        ['4', '高安全性要求+劳动力短缺是推动机器人替代的核心驱动力'],
        ['5', '竞品已实现商用（埃欧珞/凌度），但行业整体仍处早期，先发优势窗口仍在'],
    ])

h('下一步建议', level=2)
tb(['优先级', '行动', '目的'],
    [
        ['P0', '核实内部数据', '确认产品技术参数、过往销售数据、高管信息等'],
        ['P0', '竞品实际价格摸底', '通过渠道了解竞品实际成交价'],
        ['P1', '行业专家访谈', '验证市场数据假设'],
        ['P1', '客户回访', '验证购买理由和障碍的排序'],
        ['P1', '销售团队访谈', '获取一线客户反馈'],
        ['P2', '展会/协会对接', '建立渠道网络'],
    ])

# ======================== SAVE ========================
docx_path = os.path.join(OUT, '高空清洗机器人市场快速调研报告.docx')
doc.save(docx_path)
print('Word:', docx_path)
print('Size:', os.path.getsize(docx_path), 'bytes')

from docx2pdf import convert
pdf_path = os.path.join(OUT, '高空清洗机器人市场快速调研报告.pdf')
print('Converting to PDF...')
convert(docx_path, pdf_path)
print('PDF:', pdf_path)
print('PDF size:', os.path.getsize(pdf_path), 'bytes')
