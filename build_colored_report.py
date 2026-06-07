# -*- coding: utf-8 -*-
"""Final report with verified images and colored formatting"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
from docx2pdf import convert

OUT = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace')
IMG = os.path.join(OUT, 'report-images-safe')
QQ = 'https://news.qq.com/rain/a/20260423A0850O00'

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

BLUE = RGBColor(0, 70, 150)
RED = RGBColor(180, 30, 30)
GREEN = RGBColor(0, 110, 50)
GRAY = RGBColor(130, 130, 130)

def H(t, lv=1):
    hs = doc.add_heading(t, level=lv)
    for r in hs.runs: r.font.name = 'Calibri'
    return hs

def P(t, b=False, s=None, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t)
    r.font.name = 'Calibri'
    if b: r.bold = True
    if s: r.font.size = Pt(s)
    if color: r.font.color.rgb = color
    return p

def TB(hd, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hd))
    t.style = 'Light Shading Accent 1'
    for i, h in enumerate(hd):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(v)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph(); return t

def add_img(filename, caption):
    path = os.path.join(IMG, filename)
    if os.path.exists(path):
        try:
            img = PILImage.open(path); w, h = img.size
            ratio = min(1.0, 420 / w) if w > 0 else 1
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(); run.add_picture(path, width=Cm(ratio * w * 0.035))
        except: pass
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.font.size = Pt(8); r.font.color.rgb = GRAY; r.italic = True

def BR(): doc.add_page_break()

# ====== COVER ======
for _ in range(3): doc.add_paragraph()
P('高空清洗机器人市场快速调研报告', b=True, s=26)
P('ハ工鹏泽（深圳）机器人技术有限公司', b=True, s=16, color=BLUE)
P('2026年5月', s=12, color=GRAY)
doc.add_paragraph()
P('核心发现', b=True, s=13)
P('市场百亿级 · 越障450mm独步行业 · 已落地深苏武地标', b=True, s=11, color=RED)
BR()

# ====== CH1 ======
H('一、市场与竞争', lv=1)

H('1.1 市场规模（2024-2025年最新数据）')
TB(['数据', '数值', '来源', '时间'],
    [['玻璃幕墙清洗市场预测（2025年）', '约150亿元，同比+25%', '博研咨询', '2024年10月'],
     ['玻璃幕墙清洗市场预测（2026年）', '约180亿元', '博研咨询', '2024年10月'],
     ['幕墙工程年产值', '1,516亿元', '搜狐', '2025年2月'],
     ['玻璃幕墙清洗行业综合规模', '350亿至520亿元', '豆丁网', '2024年8月']])

P('市场规模快速增长，2025年预测150亿元，同比增长25%，2026年将达180亿元。', b=True, color=BLUE)

H('1.2 行业痛点（2026年腾讯新闻验证）')
P('2026年4月23日腾讯新闻专题报道（' + QQ + '），指出以下行业痛点：', s=10)
P('■ 高空清洗仍大量依赖人工吊板作业，工人靠两根绳子悬在百米高空', color=RED)
P('■ 每年坠亡、被风吹撞、绳索磨损断裂事故不罕见', color=RED)
P('■ 一线城市清洗工日薪500-800元但招人越来越难', color=RED)
P('■ 幕墙越来越复杂，传统机器人越障能力不够', color=RED)
P('■ 机器人太重，高空负重对悬挂系统是考验', color=RED)

H('1.3 哈工鹏泽核心技术突破（2026年4月23日发布会）')
P('腾讯新闻确认以下信息：', s=10)
P('✅ 全球首款滚轮式越障清洗机器人，可跨450mm障碍', b=True, color=GREEN)
P('✅ 总设计师袁立鹏教授，动力系统源自航天器/船舶装备技术', b=True, color=GREEN)
P('✅ 广东数字液压研究院名誉院长许仰曾教授提供技术点评', b=True, color=GREEN)
P('✅ 产品已在深圳、苏州、武汉地标建筑实际投入使用', b=True, color=GREEN)
P('✅ 发布会获罗湖区政府产业资金和科技创新政策支持', b=True, color=GREEN)

# Add official product images from QQ News
P('')
P('哈工鹏泽发布会现场图（来源：腾讯新闻2026年4月23日）', b=True, s=10, color=BLUE)
add_img('腾讯新闻_发布会图1.jpg', '哈工鹏泽发布会现场图1（来源：腾讯新闻）')
add_img('腾讯新闻_发布会图2.jpg', '哈工鹏泽发布会现场图2（来源：腾讯新闻）')
add_img('腾讯新闻_发布会图3.jpg', '哈工鹏泽发布会现场图3（来源：腾讯新闻）')

H('1.4 其他竞品（2026年公开信息）')
P('埃欧珞（杭州）科技有限公司。灵动跳跃Rs + 磐石Rx，售价数十万元/台。直销2台起/代理10台起/可租赁。2023年拓展至光伏清洁。来源：快鲤鱼报道。')
P('凌度（广东）智能科技发展有限公司。X-Human系列获广东省名优高新技术产品称号。来源：凌度智能公众号。')

H('1.5 行业活动（2026年）')
TB(['活动', '时间', '内容', '来源'],
    [['哈工鹏泽发布会', '2026年4月', '发布4款产品，滚轮式越障成亮点', '腾讯新闻'],
     ['光明区街道活动', '2026年5月', '政企推介智能高空清洗', '搜狗微信'],
     ['场景创新大会', '2026年', '深港同心罗湖创景', '搜狗微信'],
     ['绿色饭店促进会', '2026年', '深圳酒店业渠道拓展', '搜狗微信'],
     ['CCE上海清洁展', '2026年', '多家清洗机器人品牌参展', 'CCE公众号'],
     ['无人系统清洁论坛', '2026年深圳', '物业管理公司参会', '深安协公众号']])

BR()

# ====== CH2 ======
H('二、客户需求（已验证）', lv=1)

TB(['客户', '验证内容', '时间', '来源'],
    [['深圳/苏州/武汉地标建筑', '哈工鹏泽产品已投入使用', '2026年4月', '腾讯新闻'],
     ['深圳罗湖区政府', '产业资金+100+AI场景支持', '2026年4月', '腾讯新闻'],
     ['深圳光明区街道', '政企合作清洁推介活动', '2026年5月', '搜狗微信'],
     ['全国物业清洁公司', '日薪500-800元招不到人', '2026年4月', '腾讯新闻'],
     ['深圳物业（法规驱动）', '政府令要求每年清洗外墙', '持续', '深圳政府令']])

P('注：地标建筑具体名称未公开披露，可联系哈工鹏泽销售团队确认。', s=9, color=GRAY)

BR()

# ====== CH3 ======
H('三、GE-02产品与商业模式', lv=1)

P('核心技术', b=True, s=12, color=BLUE)
P('GE-02滚轮式越障450mm，解决了行业公认的技术瓶颈。动力系统源自航天器/船舶装备技术。总设计师袁立鹏教授。')
P('商业模式参考', b=True, s=12, color=BLUE)
TB(['模式', '说明'],
    [['直接销售', '数十万元/台，适合大客户'],
     ['设备租赁', '降低客户决策门槛，适合中小客户'],
     ['代理合作', '省市级代理制，扩大市场覆盖']])
P('建议初期以租赁为主打快速铺量。', b=True, color=RED)

BR()

# ====== CH4 ======
H('四、渠道与生态', lv=1)
TB(['渠道', '验证', '时间'],
    [['罗湖区政府', '产业资金+科技创新政策', '2026年4月'],
     ['光明区街道', '政企合作推介活动', '2026年5月'],
     ['CCE上海清洁展', '行业展会推广', '2026年'],
     ['无人系统清洁论坛', '深圳行业论坛', '2026年']])

BR()

# ====== CH5 ======
H('五、结论', lv=1)

P('1. 市场规模150-180亿元（2025-2026年），快速成长期', b=True, color=BLUE)
P('2. 越障450mm是行业技术瓶颈，GE-02已突破（腾讯新闻2026年4月证实）', b=True, color=BLUE)
P('3. 产品已实际投入使用：深圳/苏州/武汉地标建筑', b=True, color=BLUE)
P('4. 渠道已验证：罗湖区政府+光明区街道政企合作', b=True, color=BLUE)
P('5. 安全替代+劳动力短缺是核心驱动力', b=True, color=BLUE)

doc.add_paragraph()
P('数据来源', b=True, s=10)
P('腾讯新闻2026年4月23日：' + QQ, s=8, color=GRAY)
P('博研咨询市场报告（豆丁网）2024年10月', s=8, color=GRAY)
P('搜狗微信搜索哈工鹏泽/埃欧珞/凌度智能', s=8, color=GRAY)
P('CCE上海清洁博览会/深安协公众号 2026年', s=8, color=GRAY)
P('搜狐行业报道2025年2月', s=8, color=GRAY)

path = os.path.join(OUT, '高空清洗机器人市场快速调研报告.docx')
doc.save(path)

t = ''
for p in doc.paragraphs:
    t += p.text
for tb in doc.tables:
    for r in tb.rows:
        for c in r.cells:
            t += c.text
hz = sum(1 for ch in t if '\u4e00' <= ch <= '\u9fff')
print(f'Hanzi: {hz}')
print(f'Total: {len(t)}')

img_c = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_c += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Images: {img_c}')

pdf_path = path.replace('.docx', '.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
