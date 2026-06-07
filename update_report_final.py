# -*- coding: utf-8 -*-
"""Replace generic personas with local Shenzhen ones + update text"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
from docx2pdf import convert

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
safe = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', 'report-images-safe')

doc = Document(path)

# First, remove old generic persona images
removed = 0
for para in doc.paragraphs:
    for run in para.runs:
        drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        if drawings:
            for d in drawings:
                d.getparent().remove(d)
                removed += 1

print(f'Removed {removed} old images')

# Insert new 7 Shenzhen persona cards after 客户画像 heading
local_files = [
    '深圳本地画像_S级-万科物业（深圳总部）.png',
    '深圳本地画像_S级-华润物业（深圳总部）.png',
    '深圳本地画像_A级-平安金融中心.png',
    '深圳本地画像_A级-京基100-瑞吉酒店.png',
    '深圳本地画像_A级-深圳湾1号-莱佛士酒店.png',
    '深圳本地画像_B级-深圳市民中心-福田区政府.png',
    '深圳本地画像_C级-深圳中小清洁公司.png',
]

# Also insert back the product images
product_files = {
    '凌度智能_XHuman_高空幕墙清洗机器人.png': '凌度智能X-Human高空幕墙清洗机器人',
    '凌度智能_产品应用场景.jpg': '凌度智能产品应用场景',
    '埃欧珞_高空幕墙清洁机器人_产品图.png': '埃欧珞高空幕墙清洁机器人',
    '华蔚_60米磁力吸附机器人_产品图.png': '华蔚科技60米磁力吸附高空幕墙清洗机器人',
    '哈工鹏泽_GE02_滚轮越障机器人.jpg': '哈工鹏泽GE-02滚轮式越障高空清洗机器人',
}

# Find competitor sections and customer section
insertions = []

for i, para in enumerate(doc.paragraphs):
    text = para.text
    # Product images after competitor names
    if '哈工鹏泽（深圳）机器人技术有限公司（本公司）' in text and para.style.name.startswith('Heading') == False:
        insertions.append((i + 2, 'product_hgpz'))
    elif '埃欧珞（杭州）科技有限公司' in text:
        insertions.append((i + 2, 'product_aio'))
    elif '凌度（广东）智能科技发展有限公司' in text:
        insertions.append((i + 2, 'product_ld'))
    elif '竞品4：华蔚科技（福建）' in text:
        insertions.append((i + 2, 'product_hw'))
    # Customer persona images after 客户画像 heading
    if '最佳客户画像' in text and para.style.name.startswith('Heading'):
        insertions.append((i + 1, 'personas'))

# Sort by index descending
insertions.sort(key=lambda x: -x[0])

for idx, action in insertions:
    if idx >= len(doc.paragraphs):
        idx = len(doc.paragraphs) - 1
    ref = doc.paragraphs[idx]._element
    
    if action == 'product_hgpz':
        img_path = os.path.join(safe, '哈工鹏泽_GE02_滚轮越障机器人.jpg')
        if os.path.exists(img_path):
            img = PILImage.open(img_path)
            w, h = img.size
            ratio = min(1.0, 400 / w)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Cm(ratio * w * 0.035))
            ref.addnext(p._element)
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = cap.add_run('哈工鹏泽GE-02滚轮式越障高空清洗机器人')
            run2.font.size = Pt(8); run2.italic = True
            ref.addnext(cap._element)
            print('+ 哈工鹏泽 product image')
    
    elif action == 'product_aio':
        img_path = os.path.join(safe, '埃欧珞_高空幕墙清洁机器人_产品图.png')
        if os.path.exists(img_path):
            img = PILImage.open(img_path)
            w, h = img.size
            ratio = min(1.0, 400 / w)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Cm(ratio * w * 0.035))
            ref.addnext(p._element)
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = cap.add_run('埃欧珞高空幕墙清洁机器人')
            run2.font.size = Pt(8); run2.italic = True
            ref.addnext(cap._element)
            print('+ 埃欧珞 product image')
    
    elif action == 'product_ld':
        for pf in ['凌度智能_XHuman_高空幕墙清洗机器人.png', '凌度智能_产品应用场景.jpg']:
            img_path = os.path.join(safe, pf)
            if os.path.exists(img_path):
                img = PILImage.open(img_path)
                w, h = img.size
                ratio = min(1.0, 400 / w)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Cm(ratio * w * 0.035))
                ref.addnext(p._element)
                print(f'+ {pf}')
    
    elif action == 'product_hw':
        img_path = os.path.join(safe, '华蔚_60米磁力吸附机器人_产品图.png')
        if os.path.exists(img_path):
            img = PILImage.open(img_path)
            w, h = img.size
            ratio = min(1.0, 400 / w)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Cm(ratio * w * 0.035))
            ref.addnext(p._element)
            print('+ 华蔚 product image')
    
    elif action == 'personas':
        # Insert 7 persona cards
        for pf in reversed(local_files):
            img_path = os.path.join(safe, pf)
            if os.path.exists(img_path):
                img = PILImage.open(img_path)
                w, h = img.size
                ratio = min(1.0, 480 / w)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Cm(ratio * w * 0.035))
                ref.addnext(p._element)
                space = doc.add_paragraph()
                ref.addnext(space._element)
                print(f'+ {pf}')

doc.save(path)

# Count
img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_count += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Total images: {img_count}')

text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text
print(f'Total chars: {len(text)}')

pdf_path = path.replace('.docx', '.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
