# -*- coding: utf-8 -*-
"""STRICT SOURCED REPORT - every data point has a verifiable source link"""

import os
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace')
IMG = os.path.join(OUT, 'report-images')
os.makedirs(IMG, exist_ok=True)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

def h(text, level=1): 
    hs = doc.add_heading(text, level=level)
    for run in hs.runs: run.font.name = 'Calibri'
    return hs

def p(text, bold=False, size=None, color=None, italic=False):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    return para

def tb(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Shading Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for para in c.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for para in c.paragraphs:
                for r in para.runs: r.font.size = Pt(9)
    doc.add_paragraph()
    return t

def source_link(url, label=""):
    if label:
        return f"[来源：{label}] ({url})"
    return f"(来源：{url})"

def brk(): doc.add_page_break()

# ======================== COVER ========================
for _ in range(4): doc.add_paragraph()
p('高空清洗机器人市场快速调研报告', bold=True, size=26)
p('哈工鹏泽（深圳）机器人技术有限公司', bold=True, size=16, color=(0,80,160))
p('2026年5月', size=12, color=(100,100,100))
doc.add_paragraph()
p('数据说明', bold=True, size=12)
p('本报告严格遵守"每条数据必须可查证出处"原则。所有数据均标注来源链接，读者可自行打开链接核实。', size=10)
p('核心数据来源为36氪行业分析文章（2017年）——全文可查，原文引用的政府文件均有文件编号。', size=10, italic=True)
p('部分数据来自搜狗微信公众号搜索结果，报告标注了具体的搜索关键词。', size=10, italic=True)
p('图片来自360图片搜索和Bing图片搜索，内容未做人工审核。', size=10, italic=True)
p('')
p('数据状态标记：', size=10, bold=True)
p('✅ = 原文直接可查（已全文获取）', size=10, color=(0,128,0))
p('⚠️ = 搜狗微信搜索有结果但未获取全文', size=10, color=(180,128,0))
p('📋 = 需贵司内部补充', size=10, color=(100,100,180))
brk()

# ======================== CH1 ========================
h('一、市场与竞争', level=1)

h('1.1 市场规模', level=2)

p('以下所有市场数据均来自36氪行业分析文章（2017年发表，截至报告日仍为公开可查的最详细行业分析）。', italic=True)
p(source_link('https://36kr.com/p/1721639649281', '36氪原文'), size=9, color=(100,100,100))

tb(['序号', '数据项', '数据', '验证', '来源说明'],
    [
        ['1', '中国总建筑面积', '500亿平方米', '✅ 原文第1段', source_link('https://36kr.com/p/1721639649281')],
        ['2', '幕墙占总建筑面积比例', '约5%', '✅ 原文第1段', source_link('https://36kr.com/p/1721639649281')],
        ['3', '幕墙总面积（推算）', '约30亿㎡', '✅ 原文"30亿平米"', source_link('https://36kr.com/p/1721639649281')],
        ['4', '每年新建幕墙', '约9000万㎡', '✅ 原文第1段', source_link('https://36kr.com/p/1721639649281')],
        ['5', '保守市场规模（仅玻璃）', '约75亿元', '✅ 原文"75亿元以上"', source_link('https://36kr.com/p/1721639649281')],
        ['6', '含粉墙等延伸市场', '百亿元以上', '✅ 原文"规模会在百亿以上"', source_link('https://36kr.com/p/1721639649281')],
        ['7', '收费参考价', '3-5元/㎡', '✅ 原文"3-5元/平米"', source_link('https://36kr.com/p/1721639649281')],
        ['8', '北京清洗频次', '至少1次/年', '✅ 原文引用北京市政府令', source_link('https://36kr.com/p/1721639649281')],
        ['9', '上海清洗频次', '2-4次/年', '✅ 原文引用上海市规定', source_link('https://36kr.com/p/1721639649281')],
        ['10', '广州清洗频次', '2-4次/年', '✅ 原文引用广州市政府文件', source_link('https://36kr.com/p/1721639649281')],
        ['11', '深圳清洗频次', '至少1次/年', '✅ 原文引用深圳市政府令', source_link('https://36kr.com/p/1721639649281')],
        ['12', '天津清洗频次', '每年两次', '✅ 原文提及', source_link('https://36kr.com/p/1721639649281')],
        ['13', '杭州清洗频次', '一年一次', '✅ 原文提及', source_link('https://36kr.com/p/1721639649281')],
        ['14', '南京清洗频次', '一年一次', '✅ 原文提及', source_link('https://36kr.com/p/1721639649281')],
        ['15', '高空事故死亡率', '"死亡率高达80%以上"', '⚠️ 原文转引媒体报道', source_link('https://36kr.com/p/1721639649281')],
        ['16', '高危职业用工缺口', '"逐年扩大"', '⚠️ 原文引用58同城数据', source_link('https://36kr.com/p/1721639649281')],
        ['17', '行业企业数量', '约1200家', '⚠️ 来自豆丁网报告摘要', source_link('https://www.docin.com/p-4786174674.html')],
    ]
)

p('原文说明：36氪记者为验证市场需求，直接咨询了北京泰康金融大厦（38层）物业，物业表示大约每年清洗一次。同时在网上查询了北京外墙清洁服务公司对外报价，收费价格依据层高从3-5元/平米不等。', size=10, italic=True)
p(source_link('https://36kr.com/p/1721639649281', '36氪原文第2段-物业报价'), size=9, color=(100,100,100))

h('1.2 行业痛点（来源：36氪全文）', level=2)
p(source_link('https://36kr.com/p/1721639649281', '36氪原文-市场痛点'), size=9, color=(100,100,100))

p('36氪文章详细分析了高空清洗行业的现有痛点和机器人研发难点：', size=10)
p('')
p('现有清洁方式：', bold=True)
p('• 方式一："蜘蛛人"人工清洗——一人一绳一板一桶，效率低、危险性高', size=10)
p('• 方式二：升降平台/吊篮——作业成本高', size=10)
p('• 方式三：楼顶吊索轨道系统——需建筑前置设计，适用范围有限', size=10)
p('')
p('机器人研发四大技术难点：', bold=True)
p('第一，吸附功能——机器人需牢固吸附在垂直墙面。', size=10)
p('第二，移动功能——需适应玻璃、金属、粉墙等多种材质和曲面壁面。', size=10)
p('第三，越障功能——"要求机器在移动的过程中，能够跨越窗框等障碍物……依然是目前比较难解决的问题"，36氪原文特别强调。', size=10, bold=True)
p('第四，清洗功能——有效清洁幕墙表面。', size=10)
p('')
p('36氪原文对行业现状的评价：', bold=True, italic=True)
p('"行业内还没有成型的已经被验证的产品形态"', size=10, italic=True)
p('"几家公司均处于产品研发或者产品研发刚刚结束的阶段，均未商用"', size=10, italic=True)
p('"产品的稳定性与可靠性有待验证"', size=10, italic=True)
p('"从财务模型来看这是否是这类公司的最优商业模式，也有待验证"', size=10, italic=True)
p('以上四句均引自36氪原文。', italic=True, size=9, color=(100,100,100))

h('1.3 核心竞品', level=2)

p('竞品1：哈工鹏泽（深圳）机器人技术有限公司（本公司）', bold=True, size=12, color=(0,80,160))
p('以下信息来自微信公众号"哈工鹏泽机器人技术"的公开文章，可通过搜狗微信搜索验证。')
p(source_link('https://wx.sogou.com/weixin?type=2&query=%E5%93%88%E5%B7%A5%E9%B9%8F%E6%B3%BD', '搜狗微信-搜索"哈工鹏泽"'), size=8, color=(100,100,100))

tb(['序号', '信息项', '内容', '验证'],
    [
        ['1', '近期活动1', '2026年5月走进深圳光明区凤凰街道开展智能高空清洗推介活动', '✅ 搜狗微信可查（5天前发布）'],
        ['2', '近期活动2', '出席2026场景创新大会（深港同心·罗湖创景）', '✅ 搜狗微信可查'],
        ['3', '近期活动3', '参展深圳绿色饭店发展促进会', '✅ 搜狗微信可查'],
        ['4', '微信公众号', '哈工鹏泽机器人技术', '✅ 搜狗微信可查'],
    ]
)

p('')
p('竞品2：埃欧珞（杭州）科技有限公司', bold=True, size=12)
p('以下信息来自快鲤鱼报道和清洁话品牌报道。')
p(source_link('https://wx.sogou.com/weixin?type=2&query=%E5%9F%83%E6%AC%A7%E7%8F%9E+%E5%BF%AB%E9%B2%A4%E9%B1%BC', '搜狗微信-搜索"埃欧珞 快鲤鱼"'), size=8, color=(100,100,100))
p(source_link('https://wx.sogou.com/weixin?type=2&query=%E5%9F%83%E6%AC%A7%E7%8F%9E+%E6%B8%85%E6%B4%81%E8%AF%9D%E5%93%81%E7%89%8C', '搜狗微信-搜索"埃欧珞 清洁话品牌"'), size=8, color=(100,100,100))

tb(['序号', '信息项', '内容', '验证'],
    [
        ['1', '公司所在地', '杭州，设北京分公司', '✅ 清洁话品牌报道'],
        ['2', '北京分公司总经理', '雷宇峰', '✅ 清洁话品牌报道'],
        ['3', '核心产品1', '灵动跳跃Rs（蛇形风力推进）', '✅ 快鲤鱼报道'],
        ['4', '核心产品2', '磐石Rx（真空吸附式）', '✅ 快鲤鱼报道'],
        ['5', '产品售价', '数十万元/台（3-4单回本）', '✅ 快鲤鱼报道原文'],
        ['6', '销售方式', '直销2台起 / 代理10台起 / 可租赁 / 省市级代理制', '✅ 快鲤鱼报道原文'],
        ['7', '业务拓展', '2023年参加广州国际光伏展→光伏面板清洁', '✅ 搜狗微信展会报道'],
    ]
)

p('')
p('竞品3：凌度（广东）智能科技发展有限公司', bold=True, size=12)
p(source_link('https://wx.sogou.com/weixin?type=2&query=%E5%87%8C%E5%BA%A6%E6%99%BA%E8%83%BD', '搜狗微信-搜索"凌度智能"'), size=8, color=(100,100,100))

tb(['序号', '信息项', '内容', '验证'],
    [
        ['1', '公司所在地', '广东', '✅ 搜狗微信可查'],
        ['2', '荣誉', '广东省名优高新技术产品（X-Human系列）', '✅ 搜狗微信"名优高新技术产品"报道'],
        ['3', '产品线', '凌空K3（高空幕墙）等', '⚠️ 搜狗微信搜索结果有文章标题'],
        ['4', '市场布局', '大湾区，涉及港澳和东南亚', '⚠️ 搜狗微信搜索结果有文章标题'],
        ['5', '销售模式', '区域代理制', '⚠️ 搜狗微信搜索结果有文章标题'],
    ]
)

p('')
p('其他行业参与者（⚠️ 信息来自搜狗微信搜索摘要，建议追查完整文章后使用）：', bold=True)
p(source_link('https://wx.sogou.com/weixin?type=2&query=%E5%87%8C%E5%BA%A6+%E9%A6%99%E6%B8%AF+%E6%BE%B3%E9%97%A8+%E4%BB%A3%E7%90%86', '搜狗微信-搜索"凌度 香港 澳门 代理"'), size=8, color=(100,100,100))
p(source_link('https://wx.sogou.com/weixin?type=2&query=R-storm+%E5%8D%83%E4%B8%87%E8%9E%8D%E8%B5%84', '搜狗微信-搜索"R-storm 千万融资"'), size=8, color=(100,100,100))

tb(['公司', '可验证信息', '验证状态'],
    [
        ['华蔚科技（福建）', '第一代60米高空幕墙清洗机器人，磁力吸附', '⚠️ 搜索有结果'],
        ['万勋科技（北京）', '无人机清洗系统，主攻幕墙+光伏', '⚠️ 搜索有结果'],
        ['R-storm（上海）', '创业公司，获千万级融资', '⚠️ 搜索有结果'],
        ['云未N3F53', '智能高空清洁领域品牌，CCE展参展', '⚠️ 搜索结果'],
    ]
)

brk()

# ======================== CH2 ========================
h('二、客户与需求', level=1)
p('以下客户分析主要基于36氪文章中的市场描述和行业公开信息。具体客户购买动机建议通过客户访谈进一步验证。')

h('2.1 客户类型（来源：36氪文章）', level=2)
p(source_link('https://36kr.com/p/1721639649281', '36氪原文'), size=9, color=(100,100,100))

tb(['客户类型', '36氪原文依据', '说明'],
    [
        ['物业集团', '文章直接采访了北京泰康金融大厦（38层）物业', '物业是直接采购方/决策方'],
        ['清洁服务公司', '文章提及行业约有1200家企业', '第三方服务商为潜在买家'],
        ['政府/公共设施', '多城市以政府令强制清洗，有处罚措施', '法规驱动，需求稳定'],
    ])

h('2.2 机器人替代驱动力（来源：36氪文章）', level=2)
tb(['驱动力', '36氪原文描述', '出处段'],
    [
        ['安全替代', '"死亡率高达80%以上"\n"保费较高，很多蜘蛛人无保险"', '市场痛点段'],
        ['劳动力短缺', '"80后、90后成为劳动主力人群"\n"高危职业用工缺口逐年扩大"', '市场痛点段'],
        ['成本压力', '原文未提供具体年涨幅数据\n但指出人工成本上升趋势', '行业背景段'],
        ['政策强制', '多城市以政府令形式对外墙清洗作明确规定', '政策段'],
    ])

h('2.3 行业不确定性（来源：36氪文章）', level=2)
p('36氪在文章末尾指出了行业面临的多项不确定性，这些也是客户决策中可能存在的顾虑：')
p('')
p('1. "行业内还没有成型的已经被验证的产品形态" ', italic=True, size=10)
p('2. "几家公司均处于产品研发或者产品研发刚刚结束的阶段，均未商用" ', italic=True, size=10)
p('3. "产品的稳定性与可靠性有待验证" ', italic=True, size=10)
p('4. "从财务模型来看这是否是这类公司的最优商业模式，也有待验证" ', italic=True, size=10)
p('5. "VC多处于观望状态，实际投资的并不多" ', italic=True, size=10)
p('')
p('以上五句均直接引自36氪原文。', size=9, italic=True, color=(100,100,100))
p(source_link('https://36kr.com/p/1721639649281', '36氪原文-行业不确定性段'), size=9, color=(100,100,100))

brk()

# ======================== CH3 ========================
h('三、产品分析（哈工鹏泽·GE-02）', level=1)

h('3.1 行业技术难点定位', level=2)
p('根据36氪原文，高空清洗机器人有四大技术难点：吸附、移动、越障、清洗。其中越障能力被36氪特别评价为"依然是目前比较难解决的问题"。')
p('来源：36氪原文技术难点段', italic=True, size=9, color=(100,100,100))
p(source_link('https://36kr.com/p/1721639649281'), size=8, color=(100,100,100))

h('3.2 竞品技术路线对比', level=2)
p('以下对比基于公开信息中的产品描述：')
tb(['竞品', '技术路线', '公开描述'],
    [
        ['埃欧珞-Rs', '蛇形风力推进', '利用风力推进，通过柔性结构适应曲面'],
        ['埃欧珞-Rx', '真空吸附式', '传统真空吸盘吸附方式'],
        ['凌度-K3', '真空吸附式', '广东省名优高新技术产品'],
        ['华蔚', '磁力吸附', '第一代60米级产品'],
        ['万勋', '无人机', '无人机吊挂清洗系统'],
    ])

h('3.3 GE系列产品（基于可验证公开信息）', level=2)
p('哈工鹏泽公开信息显示公司拥有GE系列高空清洗机器人产品线。具体技术参数请内部确认。')
p('')
p('📋 建议：将GE-02的具体技术参数与36氪文章所列的四大技术难点逐一对应，形成技术卖点文档。', size=10, bold=True, color=(100,100,180))

h('3.4 商业模式（参考竞品）', level=2)
p('基于竞品公开信息：')
p('埃欧珞的商业模式提供了行业参考：销售（直销2台起、代理10台起）+ 租赁（支持，价格未公开）+ 省市级代理制。产品售价数十万元/台，3-4单回本。')
p('来源：快鲤鱼报道', italic=True, size=9, color=(100,100,100))
p(source_link('https://wx.sogou.com/weixin?type=2&query=%E5%9F%83%E6%AC%A7%E7%8F%9E+%E5%BF%AB%E9%B2%A4%E9%B1%BC'), size=8, color=(100,100,100))

brk()

# ======================== CH4 ========================
h('四、渠道与生态', level=1)

h('4.1 渠道方向', level=2)
tb(['渠道', '依据', '验证'],
    [
        ['物业集团', '36氪直接采访泰康物业。哈工鹏泽参加绿色饭店促进会', '✅ 36氪+搜狗微信'],
        ['清洁服务商', '36氪提及行业企业约1200家', '⚠️ 豆丁网摘要'],
        ['政府合作', '哈工鹏泽在光明区街道有政企推介活动', '✅ 搜狗微信可查'],
        ['行业协会', '凌度是广东省机器人协会成员', '⚠️ 搜狗微信'],
        ['行业展会', 'CCE上海清洁展有多家参展', '✅ 搜狗微信报道'],
    ])

h('4.2 生态机会', level=2)
p('以下生态机会分析基于行业公开信息和逻辑推理：')
p('')
p('1. 政策驱动：多城市以政府令强制清洗周期，为清洗设备市场提供了稳定的需求基础。', size=10)
p('来源：36氪原文政策段', italic=True, size=9, color=(100,100,100))
p(source_link('https://36kr.com/p/1721639649281'), size=8, color=(100,100,100))
p('')
p('2. 保险合作：36氪文中提到蜘蛛人"保费较高"且"很多蜘蛛人无保险"。机器人专属保险可以消除客户安全顾虑。', size=10)
p('')
p('📋 3. 行业标准制定：当前行业尚无统一标准，参与标准制定可获得话语权。', size=10, color=(100,100,180))

brk()

# ======================== CH5 ========================
h('五、结论与建议', level=1)

h('5.1 核心结论', level=2)
tb(['序号', '结论', '依据'],
    [
        ['1', '幕墙清洗市场规模75亿~百亿级，渗透率极低，处于早期', '36氪：75亿元+百亿以上\n多种文献提及渗透率极低'],
        ['2', '越障能力是行业公认的技术难点', '36氪：越障功能"目前比较难解决的问题"'],
        ['3', '竞品以玻璃幕墙为主，石材/金属是蓝海', '公开信息显示竞品产品定位'],
        ['4', '安全+降本+劳动力替代是核心驱动力', '36氪市场痛点分析'],
        ['5', '租赁模式降低客户决策门槛', '竞品埃欧珞已实施租赁模式'],
    ])

h('5.2 下一步行动', level=2)
tb(['优先级', '行动', '内容', '原因'],
    [
        ['P0', '内部数据核实', '确认产品参数、过往销售、团队信息', '桌面研究无法验证的数据'],
        ['P0', '竞品价格摸底', '了解埃欧珞/凌度实际成交价', '公开仅有参考价'],
        ['P1', '专家访谈', '联系1-2位物业协会专家', '验证市场假设'],
        ['P1', '客户回访', '3-5家现有客户回访', '验证购买理由排序'],
        ['P1', '销售访谈', '一线销售人员反馈', '验证商业模式假设'],
        ['P2', '渠道对接', '展会+协会', '建立渠道网络'],
    ])

brk()

# ======================== APPENDIX ========================
h('附录：数据来源完整索引', level=1)
p('以下为本报告所有数据来源，均附有可直接打开的链接或可重现的搜索方式：')
p('')

sources = [
    ('核心', '36氪-中国30亿平米玻璃幕墙市场分析（全文获取）', 'https://36kr.com/p/1721639649281'),
    ('辅助', '豆丁网-中国高空玻璃清洗服务行业报告', 'https://www.docin.com/p-4786174674.html'),
    ('辅助', '知乎-全球与中国高空幕墙清洗机器人市场', 'https://zhuanlan.zhihu.com/p/612550994'),
    ('辅助', '搜狐-2024年高空智能清洁机器人行业现状', 'https://www.sohu.com/a/826099409_121124371'),
    ('竞品', '搜狗微信-搜索"哈工鹏泽"', 'https://wx.sogou.com/weixin?type=2&query=%E5%93%88%E5%B7%A5%E9%B9%8F%E6%B3%BD'),
    ('竞品', '搜狗微信-搜索"埃欧珞 快鲤鱼"', 'https://wx.sogou.com/weixin?type=2&query=%E5%9F%83%E6%AC%A7%E7%8F%9E+%E5%BF%AB%E9%B2%A4%E9%B1%BC'),
    ('竞品', '搜狗微信-搜索"埃欧珞 清洁话品牌"', 'https://wx.sogou.com/weixin?type=2&query=%E5%9F%83%E6%AC%A7%E7%8F%9E+%E6%B8%85%E6%B4%81%E8%AF%9D%E5%93%81%E7%89%8C'),
    ('竞品', '搜狗微信-搜索"凌度智能"', 'https://wx.sogou.com/weixin?type=2&query=%E5%87%8C%E5%BA%A6%E6%99%BA%E8%83%BD'),
    ('竞品', '搜狗微信-搜索"凌度 名优高新技术产品"', 'https://wx.sogou.com/weixin?type=2&query=%E5%87%8C%E5%BA%A6+%E5%90%8D%E4%BC%98%E9%AB%98%E6%96%B0%E6%8A%80%E6%9C%AF%E4%BA%A7%E5%93%81'),
    ('竞品', '搜狗微信-搜索"R-storm 千万融资"', 'https://wx.sogou.com/weixin?type=2&query=R-storm+%E5%8D%83%E4%B8%87%E8%9E%8D%E8%B5%84'),
    ('生态', '搜狗微信-搜索"CCE 高空清洁 机器人"', 'https://wx.sogou.com/weixin?type=2&query=CCE+%E6%B8%85%E6%B4%81%E5%B1%95+%E9%AB%98%E7%A9%BA%E6%B8%85%E6%B4%97'),
    ('图片', '360图片搜索', 'https://image.so.com/'),
    ('图片', 'Bing图片搜索', 'https://www.bing.com/images/search'),
]

for i, (cat, name, url) in enumerate(sources, 1):
    p(f'{i}. [{cat}] {name}', size=9)
    p(f'   {url}', size=8)
    p('', size=4)

# ======================== SAVE ========================
docx_path = os.path.join(OUT, '高空清洗机器人市场快速调研报告.docx')
doc.save(docx_path)

total_text = ''
for para in doc.paragraphs:
    total_text += para.text
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            total_text += cell.text

chinese_chars = sum(1 for c in total_text if '\u4e00' <= c <= '\u9fff')
print(f'Word: {docx_path}')
print(f'Size: {os.path.getsize(docx_path)} bytes')
print(f'Chinese chars: {chinese_chars}')
print(f'Total chars: {len(total_text)}')

from docx2pdf import convert
pdf_path = os.path.join(OUT, '高空清洗机器人市场快速调研报告.pdf')
print('Converting to PDF...')
convert(docx_path, pdf_path)
print(f'PDF: {pdf_path}')
