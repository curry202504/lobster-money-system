import os
from docx import Document
from docx.shared import Pt
from docx2pdf import convert

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
doc = Document(path)

end = '''
关于市场推广策略的进一步建议。在市场教育阶段，建议优先选择一到两个标杆客户进行深度合作，打造可复制的成功案例。标杆客户的选择应兼顾行业影响力、合作意愿和项目可操作性。北京上海深圳等一线城市的标志性建筑是理想的目标。成功案例的打造需要投入足够的资源，包括产品优化、现场支持和效果评估。案例成熟后可以通过行业媒体、协会活动、展会展示等多种渠道进行传播。

关于产品差异化策略的总结。GE-02的滚轮式越障能力如果确实达到450毫米，在目前已知的竞品中具有独特性。这个优势需要在产品宣传材料中清晰传达。同时全材质适配能力也是一个重要的差异化卖点。在营销传播中可以将这两个优势与36氪文章指出的行业技术难点进行对照，强化产品的技术领先形象。'''

for line in end.strip().split('\n'):
    line = line.strip()
    if line:
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

doc.save(path)

text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text
hz = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
print(f'Hanzi: {hz}')
print(f'Total: {len(text)}')

img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_count += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Images: {img_count}')

pdf_path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
