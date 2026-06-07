# -*- coding: utf-8 -*-
"""Verify downloaded images and rebuild the Word document with embedded images"""
import os
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUTPUT_DIR = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace')
IMAGE_DIR = os.path.join(OUTPUT_DIR, 'report-images')

# Verify images first
print("=== Verifying Images ===")
image_files = {
    '哈工鹏泽_0': os.path.join(IMAGE_DIR, '哈工鹏泽_0.png'),
    '埃欧珞_0': os.path.join(IMAGE_DIR, '埃欧珞_0.png'),
    '埃欧珞_1': os.path.join(IMAGE_DIR, '埃欧珞_1.jpg'),
    '凌度智能_0': os.path.join(IMAGE_DIR, '凌度智能_0.png'),
    '凌度智能_1': os.path.join(IMAGE_DIR, '凌度智能_1.jpg'),
}

verified = {}
for name, path in image_files.items():
    if os.path.exists(path):
        try:
            img = Image.open(path)
            w, h = img.size
            fmt = img.format
            verified[name] = (path, w, h, fmt)
            print(f'  {name}: {w}x{h} {fmt} - OK')
        except Exception as e:
            print(f'  {name}: Error opening - {e}')
    else:
        print(f'  {name}: File not found')

print(f'\nVerified {len(verified)} images')

# Now rebuild the Word document with images embedded
print("\n=== Building Word Document ===")

doc = Document()

# Style setup
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False, size=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    if bold: run.bold = True
    if size: run.font.size = Pt(size)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri+1].cells[ci].text = str(val)
    doc.add_paragraph()
    return table

def add_image_from_file(img_path, caption, source=""):
    """Add image to document with caption"""
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(4.5))
        # Caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.italic = True
        if source:
            src = doc.add_paragraph()
            src.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = src.add_run(f'来源：{source}')
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(150, 150, 150)
    else:
        doc.add_paragraph(f'[图片未加载: {caption}]')

# ========== COVER PAGE ==========
doc.add_paragraph()
doc.add_paragraph()
add_heading('高空清洗机器人市场快速调研报告', level=0)
add_para('调研时间：2026年5月 | 报告性质：桌面研究初稿', size=12)
doc.add_page_break()

# ========== MODULE 1 ==========
add_heading('模块一：市场与竞争', level=1)

add_heading('1.1 真实市场规模', level=2)
add_table(
    ['指标', '数据', '说明'],
    [
        ['高空清洁机器人市场规模（2030E）', '约223亿元', '无人机清洗渗透率将从<2%提升至10-20%'],
        ['当前机器人/无人机清洗渗透率', '不足2%', '绝大部分依赖人工蜘蛛人'],
        ['全国高空外墙清洗年总金额（估算）', '500亿元+', '含人工清洗的整体市场盘子'],
        ['清洁机器人市场复合增长率', '约30%+', '市场处于快速上升期'],
    ]
)
add_para('关键判断：', bold=True)
add_para('• 渗透率极低（<2%），市场爆发前夜', size=10)
add_para('• 驱动力：安全法规趋严、劳动力短缺、物业智能化', size=10)
add_para('• 500亿市场若渗透率提升至10%=50亿替代空间', size=10)

add_heading('1.2 核心竞争对手画像', level=2)

# 竞品1: 哈工鹏泽
add_para('竞品1：哈工鹏泽（深圳）机器人技术有限公司 ⭐（自身定位为标杆）', bold=True, size=11)
image_file = verified.get('哈工鹏泽_0', [None])[0]
add_image_from_file(image_file, '哈工鹏泽-重载四足机器人与滚轮式越障高空清洗机器人', 'Bing图片搜索')
add_table(
    ['维度', '详情'],
    [
        ['总部', '深圳'],
        ['董事长', '杨剑'],
        ['核心产品', 'GE-02 滚轮式越障高空清洗机器人（全球首款滚轮式越障，越障450mm）\nGE-05 轻型吸附式横向清洗机器人\n重载四足机器人'],
        ['市场表现', '单日签约订单超2300万元\n亮相2026场景创新大会、绿色饭店发展促进会'],
        ['来源', '微信公众号：哈工鹏泽机器人技术'],
    ]
)

# 竞品2: 埃欧珞
add_para('竞品2：埃欧珞（杭州）科技有限公司', bold=True, size=11)
image_file_aio1 = verified.get('埃欧珞_0', [None])[0]
image_file_aio2 = verified.get('埃欧珞_1', [None])[0]
add_image_from_file(image_file_aio1, '埃欧珞-高空幕墙清洁机器人产品图', '360图片搜索')
add_image_from_file(image_file_aio2, '埃欧珞-灵动跳跃Rs系列产品图', '360图片搜索')
add_table(
    ['维度', '详情'],
    [
        ['总部', '杭州'],
        ['核心产品', '灵动跳跃Rs（蛇形风力机器人）\n磐石Rx（吸附式机器人）'],
        ['销售模式', '直销2台起售，代理10台起售；支持租赁；产品售价数十万元/台'],
        ['业务拓展', '从高空擦窗扩展到光伏面板清洁'],
        ['来源', '快鲤鱼'],
    ]
)

# 竞品3: 凌度智能
add_para('竞品3：凌度（广东）智能科技发展有限公司', bold=True, size=11)
image_file_ld1 = verified.get('凌度智能_0', [None])[0]
image_file_ld2 = verified.get('凌度智能_1', [None])[0]
add_image_from_file(image_file_ld1, '凌度智能-高空幕墙清洗机器人产品图', '360图片搜索')
add_image_from_file(image_file_ld2, '凌度智能-产品应用场景图', '360图片搜索')
add_table(
    ['维度', '详情'],
    [
        ['总部', '广东'],
        ['创始人', '黄俊生'],
        ['核心产品', '凌空K3（高空幕墙清洗）、凌净J1 SE（低空）\nX-Human系列高空幕墙清洗机器人\n分布式光伏智能清洗机器人'],
        ['市场布局', '深圳、大湾区、厦门、上海、港澳、东南亚'],
        ['来源', '凌度智能科技微信公众号'],
    ]
)

# Other competitors
add_para('其他值得关注的竞品：', bold=True)
add_table(
    ['公司', '特点'],
    [
        ['华蔚科技（福建）', '第一代60米高空幕墙清洗机器人，磁力吸附，创始人刘昌臻'],
        ['万勋科技（北京）', '无人机清洗系统，主攻幕墙+光伏，总经理王书研'],
        ['R-storm（上海）', '创业公司，获千万级融资'],
        ['云未N3F53', '智能高空清洁领域先行者，CCE上海清洁展参展'],
        ['史河机器人BeeBot', '外墙清洗机器人，CCE展热点'],
    ]
)

add_heading('1.3 对手打法总结', level=2)
add_table(
    ['维度', '哈工鹏泽（自）', '埃欧珞', '凌度'],
    [
        ['销售模式', '直销+代理', '直2/代10起/可租赁', '区域代理为主'],
        ['差异化', '滚轮式越障450mm\n石材/金属/玻璃全适配', '蛇形风力+吸附\n已拓展至光伏', '港澳代理\nX-Human认证'],
        ['市场', '深圳→全国', '杭州→全国', '大湾区→港澳→东南亚'],
    ]
)

doc.add_page_break()

# ========== MODULE 2 ==========
add_heading('模块二：客户与需求', level=1)

add_heading('2.1 最佳客户画像', level=2)
add_table(
    ['客户层级', '代表', '特点', '优先级'],
    [
        ['超大型物业集团', '万科、华润、碧桂园服务、保利物业', '总部决策，注重品牌与安全', '最优先'],
        ['大型清洁公司', '玉禾田（26亿+营收）', '成本驱动，注重ROI', '优先'],
        ['酒店管理集团', '万豪、洲际、锦江、华住', '品牌+安全+差异化', '优先'],
        ['政府/公建', '政务中心、机场、火车站', '安全合规', '潜在'],
        ['中小清洁公司', '各地清洗服务商', '价格敏感，租赁易接受', '需教育'],
    ]
)

add_heading('2.2 核心购买理由', level=2)
add_para('① 杜绝安全事故——高空作业安全事故频发，机器人替代"蜘蛛人"是刚需')
add_para('② 降低长期成本——1台机器人≈3-4个蜘蛛人效率，1-2年回本')
add_para('③ 提升品牌形象——科技物业标签')
add_para('④ 应对劳动力短缺——年轻人不愿从事高危高空作业')
add_para('⑤ 清洗质量标准化——可量化、可追溯')

add_heading('2.3 决策障碍', level=2)
add_table(
    ['障碍', '程度', '应对'],
    [
        ['价格太贵', '核心', '推租赁/联营降低门槛'],
        ['效果存疑', '核心', '现场演示+案例背书'],
        ['怕不会操作', '中等', '操作培训+远程运维'],
        ['怕掉落风险', '中等', '安全认证+保险'],
        ['有固定供应商', '较低', '找到决策链关键人'],
    ]
)

doc.add_page_break()

# ========== MODULE 3 ==========
add_heading('模块三：自身产品与模式', level=1)

add_heading('3.1 GE-02 绝对优势分析', level=2)
add_table(
    ['维度', 'GE-02', '埃欧珞', '凌度', '华蔚'],
    [
        ['越障能力', '450mm滚轮式\n（全球首款）', '常规', '常规', '常规'],
        ['吸附方式', '滚轮式+负压', '风力/吸盘', '吸附式', '磁力'],
        ['适用幕墙', '石材/金属/玻璃', '玻璃为主', '玻璃为主', '玻璃隐框'],
        ['效率', '1台≈3-4人', '较高', '较高', '第一代'],
    ]
)
add_para('GE-02不可替代优势：', bold=True)
add_para('1）滚轮式越障450mm——竞品难以匹敌的核心壁垒', size=10)
add_para('2）石材/金属幕墙清洗效果——覆盖更多幕墙材质', size=10)
add_para('3）GE系列产品矩阵——GE-02越障+GE-05横向互补', size=10)

add_heading('3.2 商业模式验证', level=2)
add_table(
    ['模式', '接受度', '现金流', '适用'],
    [
        ['🏆 租赁', '最高', '持续稳定', '中小客户/首试'],
        ['💰 销售', '中等', '一次性回款好', '大客户/政府'],
        ['🤝 联营', '上升', '长期回报', '合作伙伴'],
    ]
)
add_para('建议：初期以租赁为主，大客户用销售，区域合作联营。')

doc.add_page_break()

# ========== MODULE 4 ==========
add_heading('模块四：渠道与生态', level=1)

add_heading('4.1 最佳渠道伙伴', level=2)
add_table(
    ['渠道', '典型伙伴', '价值', '优先级'],
    [
        ['物业协会/清洁协会', '全国/省市物业协会', '行业背书+客户资源', '⭐⭐⭐⭐⭐'],
        ['大型物业集团', '万科、华润、保利', '品牌效应', '⭐⭐⭐⭐'],
        ['清洗服务商', '各地清洁公司', '本地客户', '⭐⭐⭐⭐'],
        ['政府关系伙伴', '区域政企公司', '政府项目', '⭐⭐⭐'],
        ['智慧物业平台', '明源云、云智易等', '平台入口', '⭐⭐⭐'],
        ['保险公司', '平安、人保等', '保险捆绑', '⭐⭐'],
    ]
)

add_heading('4.2 合作伙伴核心诉求', level=2)
add_table(
    ['诉求', '重要性', '支持'],
    [
        ['高利润率', '最高', '有竞争力代理价+分润'],
        ['独家代理权', '高', '按区域/行业授权'],
        ['技术培训', '中', '操作+维修培训'],
        ['售后保障', '中', '远程监控+快反网络'],
    ]
)

add_heading('4.3 生态绑定机会', level=2)
add_table(
    ['方向', '可行性', '切入点'],
    [
        ['智慧物业平台', '高', 'API打通，数据入看板'],
        ['机器人保险', '中高', '专属责任险产品'],
        ['绿色建筑认证', '中', 'LEED绿建加分项'],
        ['城市更新项目', '中高', '旧改外墙清洗需求'],
    ]
)

doc.add_page_break()

# ========== CONCLUSIONS ==========
add_heading('结论与行动建议', level=1)

add_para('🎯 五大关键发现', bold=True, size=13)
add_para('1️⃣ 市场爆发前夜——渗透率<2%，天花板500亿+')
add_para('2️⃣ 越障能力是核心壁垒——GE-02越障450mm全球独有')
add_para('3️⃣ 租赁是最佳切入点——降低决策门槛')
add_para('4️⃣ 物业协会+大型物业集团是最佳渠道入口')
add_para('5️⃣ 石材/金属幕墙是隐形优势')

add_para('')
add_para('📋 下一步行动', bold=True, size=13)
add_table(
    ['优先级', '行动项', '内容'],
    [
        ['P0', '案头研究深化', '获取付费行业报告补充数据'],
        ['P0', '竞品价格摸底', '了解竞品实际成交价和租赁价'],
        ['P1', '行业访谈', '联系物业协会/清洁协会专家'],
        ['P1', '客户深访', '3-5家客户回访验证'],
        ['P1', '销售访谈', '一线销售反馈商业模式偏好'],
        ['P2', '渠道摸底', '梳理物业协会联系方式'],
    ]
)

doc.add_page_break()

# ========== SOURCES ==========
add_heading('数据来源索引', level=1)
sources = [
    ('360文库', 'https://wenku.so.com/s?q=高空清洗机器人+市场规模+行业分析'),
    ('搜狐行业报告', 'https://www.sohu.com/a/826099409_121124371'),
    ('搜狗搜索-哈工鹏泽', 'https://www.sogou.com/web?query=哈工鹏泽+滚轮式越障+清洗机器人'),
    ('搜狗微信-哈工鹏泽公众号', 'https://wx.sogou.com/weixin?type=2&query=哈工鹏泽+高空清洗机器人'),
    ('搜狗搜索-埃欧珞', 'https://www.sogou.com/web?query=埃欧珞+高空幕墙清洁机器人'),
    ('搜狗搜索-凌度智能', 'https://www.sogou.com/web?query=凌度智能+高空幕墙清洗机器人'),
    ('无人机清洗渗透率', 'https://wx.sogou.com/weixin?type=2&query=无人机清洗渗透率+不足+2'),
    ('2030年223亿市场', 'https://wx.sogou.com/weixin?type=2&query=2030年中国无人机清洗市场+223亿元'),
    ('凌度港澳代理', 'https://wx.sogou.com/weixin?type=2&query=凌度智能+成功签约香港及澳门代理'),
    ('埃欧珞快鲤鱼报道', 'https://wx.sogou.com/weixin?type=2&query=埃欧珞+以机器代替人力+高空幕墙清洁机器人'),
]

for name, url in sources:
    add_para(f'• {name}', size=9)
    add_para(f'  {url}', size=8)

# ========== SAVE ==========
docx_path = os.path.join(OUTPUT_DIR, '高空清洗机器人市场快速调研报告.docx')
doc.save(docx_path)
print(f'\nReport saved to: {docx_path}')
print(f'File size: {os.path.getsize(docx_path)} bytes')
print(f'Images embedded: {len(verified)}')
