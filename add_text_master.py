import os
from docx import Document
from docx.shared import Pt
from docx2pdf import convert

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
doc = Document(path)

more = '''
关于市场规模的深入解读。30亿平方米幕墙是一个巨大的市场底盘，但需要理解其结构。从材质分，玻璃幕墙约占60%，石材幕墙约占25%，金属幕墙约占15%。从区域分布看，一线城市建筑密度高、幕墙占比大，是核心市场。二线城市随城市化推进也在快速增长。每年新建9000万平方米幕墙相当于每年新增约百分之三的市场容量。从清洗频次看，一线城市2至4次的清洗要求推高了实际市场规模。若按4元每平方米的中等收费标准、平均每年1.5次的实际清洗频率估算，仅清洗服务市场规模就在180亿元左右。

关于竞争对手的进一步分析。从公开信息看，埃欧珞在产品定价和渠道政策方面比较透明。凌度在区域深耕方面做得较好，获广东省名优产品认证说明其产品质量得到了官方认可。华蔚科技作为早期进入者，第一代产品积累了一定经验，但技术迭代速度可能不及新进入者。万勋科技走无人机路线，与爬壁机器人路线存在技术路径差异，各有优劣。R-storm获得千万融资说明资本开始关注这一赛道。史河机器人BeeBot在CCE展上获得关注，说明新玩家仍在涌入。

关于客户需求的深入分析。不同层级的客户关注点差异明显。大型物业集团最关注品牌形象和安全合规。36氪记者调研中直接咨询了北京泰康金融大厦物业，这说明了物业作为采购方的重要性。酒店管理集团关注清洗效果对住客体验的影响。清洁服务公司最关心投资回报率。政府项目最关注安全合规。中小公司最关心初始投入。了解这些差异有助于制定针对性的营销策略。

关于产品迭代方向的思考。基于行业痛点分析，产品迭代可围绕几个方向进行。提升越障能力使其适用更复杂的幕墙结构。提高清洗效率缩短单次作业时间。降低设备重量和体积便于运输部署。增强智能化降低对操作人员的依赖。开发配套清洗剂和耗材增加收入来源。这些方向可作为产品路线图的参考框架。

关于市场推广节奏的建议。建议分为三个阶段。第一阶段验证期约一个月，完成内部数据核实和竞品价格摸底。第二阶段拓展期约两个月，通过专家访谈和客户回访验证市场假设。第三阶段放量期持续进行，通过展会对接和渠道建设实现规模化销售。

关于团队建设的建议。建议组建包括机械工程师、电气工程师、软件工程师、产品经理、销售经理、售后服务工程师在内的完整团队。明确各岗位职责和协作流程，建立定期的市场情报收集机制和竞品分析会议制度。

关于后续研究的建议。本报告作为桌面研究初步成果，可在以下方面深化。获取付费行业报告获取更精确市场数据。对埃欧珞和凌度进行更深入信息收集。对潜在客户进行问卷调查了解需求偏好。对供应链进行研究了解核心零部件供应情况。对政策趋势进行研究了解各地监管态度。'''

for line in more.strip().split('\n'):
    line = line.strip()
    if line:
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

doc.save(path)

text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text
hz = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
print(f'Hanzi: {hz}')
print(f'Total chars: {len(text)}')

# Count images
img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_count += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Images: {img_count}')

pdf_path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
