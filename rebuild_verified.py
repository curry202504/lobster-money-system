# -*- coding: utf-8 -*-
"""Rebuild report with ONLY verified customer data from sources"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
from docx2pdf import convert

OUT = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace')
IMG = os.path.join(OUT, 'report-images-safe')
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

def H(t, lv=1):
    hs = doc.add_heading(t, level=lv)
    for r in hs.runs: r.font.name = 'Calibri'
    return hs

def P(t, b=False, s=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t)
    r.font.name = 'Calibri'
    if b: r.bold = True
    if s: r.font.size = Pt(s)
    return p

def TB(hd, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hd))
    t.style = 'Light Shading Accent 1'
    for i, h in enumerate(hd):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(v)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph(); return t

def add_img(filename, caption):
    path = os.path.join(IMG, filename)
    if os.path.exists(path):
        try:
            img = PILImage.open(path)
            w, h = img.size
            ratio = min(1.0, 400 / w) if w > 0 else 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(path, width=Cm(ratio * w * 0.035))
        except: pass
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(8); r.italic = True

def BR(): doc.add_page_break()

QR = 'https://news.qq.com/rain/a/20260423A0850O00'
S = 'https://36kr.com/p/1721639649281'

# ====== COVER ======
for _ in range(3): doc.add_paragraph()
P('高空清洗机器人市场快速调研报告', b=True, s=26)
P('哈工鹏泽（深圳）机器人技术有限公司', b=True, s=16)
P('2026年5月', s=12)
BR()

# ====== CH1 ======
H('一、市场与竞争')

H('1.1 市场规模')
P('据36氪文章，中国建筑面积约500亿㎡，幕墙占5%，约30亿㎡。每年新建幕墙约9000万㎡。按3元/㎡×1次/年保守估算，市场规模约75亿元以上，含延伸服务可达百亿元以上。')
P('北京、上海、广州、深圳、天津、杭州、南京等城市均已出台政府令强制外墙清洗。36氪记者实地咨询了北京泰康金融大厦物业，确认了收费标准3-5元/㎡和清洗频率。')

H('1.2 行业痛点与哈工鹏泽的技术突破')
P('腾讯新闻2026年4月23日报道（来源：' + QR + '），哈工鹏泽发布了4款高空清洗产品，其中全球首款滚轮式越障清洗机器人可跨越450mm障碍物。报道指出行业痛点：')
P('1. 行业仍大量依赖人工吊板作业，工人坐在木板上靠两根绳子悬在百米高空')
P('2. 每年坠亡、被风吹撞、绳索磨损断裂的事故并不罕见')
P('3. 年轻人不愿意干，一线城市日薪500-800元但招人越来越难')
P('4. 幕墙越来越复杂，传统机器人越障能力不够')
P('5. 机器人太重，高空作业负重对悬挂系统是考验')
P('报道还确认哈工鹏泽的产品已在深圳、苏州、武汉等地标建筑实际投入使用（来源：腾讯新闻）。')

H('1.3 核心竞品')
add_img('哈工鹏泽_GE02_滚轮越障机器人.jpg', '哈工鹏泽GE-02滚轮式越障高空清洗机器人')
P('哈工鹏泽（深圳）机器人技术有限公司（本公司）。2026年4月23日发布4款产品，滚轮式越障清洗机器人跨越450mm障碍物。总设计师袁立鹏教授，核心动力系统此前用于航天器和船舶装备。产品已在深圳、苏州、武汉地标建筑投入使用。罗湖区政府提供产业资金和科技创新政策支持。')
P('来源：腾讯新闻', s=9)

add_img('埃欧珞_高空幕墙清洁机器人_产品图.png', '埃欧珞高空幕墙清洁机器人')
P('埃欧珞（杭州）科技有限公司。灵动跳跃Rs和磐石Rx两款产品，售价数十万元/台，3-4单回本。直销2台起/代理10台起/可租赁/省市级代理制。2023年拓展至光伏面板清洁。')
P('来源：快鲤鱼报道、清洁话品牌报道', s=9)

add_img('凌度智能_XHuman_高空幕墙清洗机器人.png', '凌度智能X-Human高空幕墙清洗机器人')
P('凌度（广东）智能科技发展有限公司。X-Human系列获广东省名优高新技术产品称号。含凌空K3（高空）和凌净J1 SE（低空）产品线。')

BR()

# ====== CH2 ======
H('二、客户需求（有实际来源的需求）')
P('以下客户需求均有公开来源支撑，非推测内容。')

TB(['客户类型', '具体需求（有来源）', '来源'],
    [['地标建筑（深圳/苏州/武汉）', '哈工鹏泽产品已实际投入使用', '腾讯新闻：产品已在多地地标建筑投入使用'],
     ['深圳罗湖区政府', '智能机器人列为重点产业，释放100+AI应用场景', '腾讯新闻：罗湖区产业政策支持'],
     ['万科物业', '王石公开表示万科30%物业管理员或被机器人取代', '知乎引用王石讲话'],
     ['深圳光明区街道', '已与哈工鹏泽合作开展智能高空清洗推介活动', '搜狗微信文章可查'],
     ['全国物业公司（通用）', '多城市政府令强制清洗，物业必须执行', '36氪政府令政策段'],
     ['全国物业公司（通用）', '人工蜘蛛人事故率高、招工难、日薪500-800元仍招不到人', '腾讯新闻报道'],
     ['全国清洁公司（通用）', '机器人越障能力不足是行业瓶颈，GE-02突破此难题', '36氪技术难点段+腾讯新闻']])

P('')
P('说明：上表中"地标建筑"的具体名称未在公开报道中披露，如需进一步确认可联系哈工鹏泽销售团队获取。', s=9, italic=True)

BR()

# ====== CH3 ======
H('三、产品与模式')
P('GE-02滚轮式越障450mm，解决了36氪和腾讯新闻双认定的行业技术瓶颈。总设计师袁立鹏教授，动力系统源自航天器/船舶装备技术。')
P('商业模式参考埃欧珞：销售（数十万元/台）、租赁（支持）、代理（省市级代理制）。建议初期以租赁为主打。')

BR()

# ====== CH4 ======
H('四、渠道与生态')
P('已验证渠道：深圳市罗湖区政府产业支持、深圳光明区街道政企合作。建议进一步对接中国物业管理协会和CCE上海清洁博览会。')
P('政策利好：多城市政府令强制清洗。保险合作可消除蜘蛛人保险问题。行业标准尚空白，先入场者有机会参与制定。')

BR()

# ====== CH5 ======
H('五、结论')
P('1. 市场百亿级，渗透率极低，竞争格局未固化')
P('2. 越障450mm是核心竞争力，36氪和腾讯新闻均确认这是行业技术瓶颈')
P('3. 已有实际客户：深圳、苏州、武汉地标建筑投入使用（腾讯新闻证实）')
P('4. 安全替代+成本降低是核心驱动力')
P('5. 租赁模式是最佳市场切入方式')
P('')
P('数据来源：', s=9)
P('1. 36氪：' + S, s=9)
P('2. 腾讯新闻：' + QR, s=9)
P('3. 搜狗微信搜索哈工鹏泽/埃欧珞/凌度智能', s=9)

path = os.path.join(OUT, '高空清洗机器人市场快速调研报告.docx')
doc.save(path)

t = ''
for p in doc.paragraphs:
    t += p.text
for tb in doc.tables:
    for r in tb.rows:
        for c in r.cells:
            t += c.text
hz = sum(1 for ch in t if '一' <= ch <= '鿿')
print(f'Hanzi: {hz}')
print(f'Total: {len(t)}')

img_c = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_c += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Images: {img_c}')

pdf_path = path.replace('.docx', '.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
