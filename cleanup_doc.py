# -*- coding: utf-8 -*-
"""Open existing 7000-char docx, remove note lines, keep clean"""
import os
from docx import Document
from docx.shared import Pt
from docx2pdf import convert

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
doc = Document(path)

# Find and remove note paragraphs (small italic footer notes)
# They are identified by: starting with "来源：" or "以上" + size 8
to_remove = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    # Remove note lines: source citations starting with specific phrases
    if text.startswith('来源：') or text.startswith('以上数据') or text.startswith('上述数据'):
        to_remove.append(i)
    # Remove very small text paragraphs (note style)
    if para.runs and para.runs[0].font.size and para.runs[0].font.size < Pt(9):
        to_remove.append(i)
    # Remove URL-only lines
    if text.startswith('http'):
        to_remove.append(i)

# Remove from bottom up to preserve indices
for i in sorted(set(to_remove), reverse=True):
    p = doc.paragraphs[i]._element
    p.getparent().remove(p)

doc.save(path)

# Count
text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text
hz = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
print(f'Hanzi: {hz}')
print(f'Total chars: {len(text)}')

pdf_path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
