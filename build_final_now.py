# -*- coding: utf-8 -*-
"""FINAL - Clean professional report, no markers, no notes"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx2pdf import convert

OUT = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace')

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

def H(t, lv=1):
    hs = doc.add_heading(t, level=lv)
    for r in hs.runs: r.font.name = 'Calibri'
    return hs

def P(t, b=False, s=None, i=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t)
    r.font.name = 'Calibri'
    if b: r.bold = True
    if i: r.italic = True
    if s: r.font.size = Pt(s)
    return p

def TB(hd, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hd))
    t.style = 'Light Shading Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(hd):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(v)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph(); return t

def BR(): doc.add_page_break()

TITLE = '高空清洗机器人市场快速调研报告'
SRC = 'https://36kr.com/p/1721639649281'

# ====== COVER ======
for _ in range(4): doc.add_paragraph()
P(TITLE, b=True, s=26)
P('哈工鹏泽（深圳）机器人技术有限公司', b=True, s=16)
P('2026年5月', s=12)
BR()

# ====== CH1 ======
H('一、市场与竞争')

H('1.1 市场规模')
P('根据36氪发布的行业分析文章，中国目前建筑面积约为500亿平方米，幕墙占比约5%，幕墙总面积约30亿平方米。每年新建幕墙约9000万平方米。仅以玻璃幕墙清洗服务计算，按每平方米3元的最低收费标准、每年一次的清洗频率，市场规模约75亿元以上。若将高空墙体粉刷等延伸服务纳入，规模可达百亿元以上。')

P('多城市已出台强制性的外墙清洗规定。北京要求至少每年清洗一次并处罚违规物业，上海要求每年2至4次，广州要求重点维护区4次、一般维护区2次，深圳、天津、杭州、南京等城市均有明确规定。这些政府令使清洗成为刚性需求。')

P('目前高空清洗主要依赖三种方式。蜘蛛人人工清洗效率低、危险性高。从业者属于5类6类高风险职业，保费较高，很多工人没有保险。媒体报道每年有数以万计的高空作业事故发生，死亡率高达80%以上。升降平台和吊篮清洗成本高、受场地限制大。楼顶吊索轨道系统需建筑前置设计，适用范围有限。随着劳动力结构变化，高危职业用工缺口持续扩大，机器替代人工是必然趋势。上述数据和判断均来自36氪行业分析文章。')

H('1.2 技术难点')
P('36氪文章梳理了高空幕墙清洗机器人研发的四大技术难点。吸附功能要求机器人牢固吸附垂直墙面。移动功能需适应玻璃、金属、粉墙等多种材质。越障功能要求跨越窗框等障碍物，被36氪评价为"目前比较难解决的问题"。清洗功能要求在移动中完成有效清洁。文章还指出，国际方面德国Skywash、日本BE等推出过方案但未大规模推广。国内自1988年以来，上海大学、哈尔滨工业大学、北京航空航天大学等高校先后研制了多款样机，但均未能商业化。')

P('现阶段随着核心零部件技术成熟、用工成本上升、政策管控加强，行业具备了发展土壤。36氪同时指出，行业尚无成熟的已验证产品形态，几家公司均处于产品研发或刚刚结束的阶段，VC多处于观望状态。这说明市场尚处早期，先入者有机会建立优势。')

H('1.3 核心竞品')
P('哈工鹏泽（深圳）机器人技术有限公司（本公司）。近期公开活动包括2026年5月走进深圳光明区凤凰街道开展智能高空清洗推介活动，出席深港同心罗湖创景场景创新大会，参展深圳绿色饭店发展促进会。以上信息来自微信公众号哈工鹏泽机器人技术。')
P('埃欧珞（杭州）科技有限公司。杭州总部设北京分公司，总经理雷宇峰。核心产品为灵动跳跃Rs（蛇形风力推进）和磐石Rx（真空吸附式），售价数十万元每台，3至4个订单可回本。销售模式为直销2台起、代理10台起，支持租赁，省市级代理制。2023年参加广州光伏展拓展至光伏面板清洁。信息来自快鲤鱼和清洁话品牌报道。')
P('凌度（广东）智能科技发展有限公司。总部广东，X-Human系列高空幕墙清洗机器人获广东省名优高新技术产品称号。信息来自凌度智能科技公众号。')

BR()

# ====== CH2 ======
H('二、客户与需求')

H('2.1 客户类型')
P('第一类是大型物业集团，如万科、华润、碧桂园服务、保利物业等。36氪记者在调研中直接咨询了北京泰康金融大厦物业，确认了物业作为采购方的角色。这类客户总部集中决策、注重品牌与安全、预算充足，率先引入机器人的物业将获得品牌差异化优势。')
P('第二类是酒店管理集团，如万豪、洲际、锦江、华住等，对外墙清洁要求高，愿意为差异化买单。第三类是清洁服务公司，36氪提到全国约1200家高空清洗企业，核心诉求是降本增效。第四类是政府及公共设施，项目通过招投标进行。第五类是中小清洁公司，租赁模式对这类客群有较强吸引力。')

H('2.2 购买驱动力')
P('安全替代是第一驱动力。高空作业死亡率高，多地政府加强安全监管，限制人工的政策若落地将使机器人从可选项变为必选项。成本降低是第二驱动力，人工成本持续上涨、用工缺口扩大。品牌提升是第三驱动力，引入机器人契合智慧物业趋势。以上分析基于36氪文章的市场痛点描述。')

BR()

# ====== CH3 ======
H('三、产品分析（哈工鹏泽·GE-02）')

H('3.1 技术优势定位')
P('36氪将越障功能列为行业最难解决的问题之一。如果GE-02的滚轮式越障技术取得突破，将构成显著壁垒。竞品以玻璃幕墙为主，GE-02对石材和金属幕墙的适配能力可以开辟差异化市场。在技术路线上，36氪提出了通用型（更适合商用）和专用型两种思路，GE-02的全材质适配定位更接近通用型路线。')

H('3.2 商业模式参考')
P('参考埃欧珞的公开信息，行业可行模式包括三种。直接销售适合大型客户，参考定价为数万元至数十万元每台。设备租赁可降低客户决策门槛。代理合作通过区域代理商扩大市场覆盖。初期建议以租赁为主打快速铺开。')

BR()

# ====== CH4 ======
H('四、渠道与生态')

P('物业协会渠道。中国物业管理协会及地方分会拥有大量物业会员资源，可提供行业背书和客户引荐。物业集团渠道方面，万科、华润等头部物业采购由总部集中决策，签约后可全国推广。区域服务商渠道方面，各地清洁工程公司对租赁和联营接受度高。政府合作方面，哈工鹏泽已在深圳光明区开展政企活动，可向其他城市复制。行业展会方面，CCE上海清洁博览会等可实现品牌曝光和客户对接。')

P('生态机会方面。多城市政府令强制清洗为市场提供稳定需求。推出机器人责任险可消除客户安全顾虑。当前行业尚无统一标准，先入场者可参与标准制定。')

BR()

# ====== CH5 ======
H('五、结论与建议')

H('5.1 核心结论')
P('第一，市场空间大、时机好。30亿平方米幕墙创造百亿级市场，政府令保障需求刚性。')
P('第二，越障能力是核心竞争力，36氪将越障列为行业最难解决的问题。')
P('第三，石材和金属幕墙是差异化蓝海。')
P('第四，安全替代和成本降低是核心驱动力。')
P('第五，租赁模式是最佳市场切入方式。')

H('5.2 下一步行动')
TB(['优先级', '行动', '内容'],
    [['P0', '内部数据核实', '确认产品参数和销售数据'],
     ['P0', '竞品价格摸底', '了解竞品实际成交价'],
     ['P1', '专家访谈', '联系物业协会专家'],
     ['P1', '客户回访', '回访3至5家客户'],
     ['P1', '销售访谈', '整理一线反馈'],
     ['P2', '渠道对接', '展会加协会']])

BR()

# ====== SOURCES ======
H('数据来源')
P('36氪行业分析：' + SRC, s=9)
P('哈工鹏泽信息：搜狗微信搜索哈工鹏泽', s=9)
P('埃欧珞信息：搜狗微信搜索埃欧珞', s=9)
P('凌度信息：搜狗微信搜索凌度智能', s=9)

# ====== SAVE ======
path = os.path.join(OUT, TITLE + '.docx')
doc.save(path)

text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text
hz = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
print(f'Hanzi: {hz}')
print(f'Total: {len(text)}')

pdf_path = os.path.join(OUT, TITLE + '.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
