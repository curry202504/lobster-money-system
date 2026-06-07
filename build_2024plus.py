# -*- coding: utf-8 -*-
"""Report using ONLY 2024+ data"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
from docx2pdf import convert

OUT = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace')
IMG = os.path.join(OUT, 'report-images-safe')
QQ = 'https://news.qq.com/rain/a/20260423A0850O00'

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

def H(t, lv=1):
    hs = doc.add_heading(t, level=lv)
    for r in hs.runs: r.font.name = 'Calibri'
    return hs

def P(t, b=False, s=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t)
    r.font.name = 'Calibri'
    if b: r.bold = True
    if s: r.font.size = Pt(s)
    return p

def TB(hd, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hd))
    t.style = 'Light Shading Accent 1'
    for i, h in enumerate(hd):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(v)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph(); return t

def add_img(filename, caption):
    path = os.path.join(IMG, filename)
    if os.path.exists(path):
        try:
            img = PILImage.open(path); w, h = img.size
            ratio = min(1.0, 400 / w) if w > 0 else 1
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(); run.add_picture(path, width=Cm(ratio * w * 0.035))
        except: pass
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.font.size = Pt(8); r.italic = True

def BR(): doc.add_page_break()

# ====== COVER ======
for _ in range(3): doc.add_paragraph()
P('高空清洗机器人市场快速调研报告', b=True, s=26)
P('哈工鹏泽（深圳）机器人技术有限公司', b=True, s=16)
P('2026年5月', s=12)
doc.add_paragraph()
P('数据范围：仅使用2024年及以后的最新公开信息', s=10)
P('核心来源：腾讯新闻2026年4月/博研咨询2024-2025/搜狐2025', s=10)
BR()

# ====== CH1 ======
H('一、市场与竞争')

H('1.1 市场规模（2024-2025年最新数据）')
P('以下数据均来自2024-2025年公开报告，时效性较好：')

TB(['数据', '数值', '发布时间', '来源'],
    [['幕墙工程年产值', '1516亿元', '2025年2月', '搜狐报道'],
     ['玻璃幕墙清洗市场预测（2025年）', '约150亿元，同比+25%', '2024年10月', '博研咨询/豆丁网'],
     ['玻璃幕墙清洗市场预测（2026年）', '约180亿元', '2024年10月', '博研咨询/豆丁网'],
     ['玻璃幕墙清洗行业综合规模', '约350-520亿元', '2024年8月', '豆丁网行业报告'],
     ['全国幕墙总面积', '约30亿㎡', '行业公开数据', '多源交叉验证']])

P('注：博研咨询预测2025年市场规模150亿元（同比增长25%），2026年180亿元。不同口径（纯清洗服务 vs 含综合服务）数据差异较大，建议结合具体需求选用。', s=9)

H('1.2 行业最新动态（2026年）')
P('2026年4月23日，腾讯新闻专题报道哈工鹏泽发布会（' + QQ + '），确认以下行业现状和市场验证：')
P('')
P('行业痛点（2026年仍存在）：', b=True)
P('- 高空清洗仍大量依赖人工吊板作业，工人靠两根绳子悬在百米高空')
P('- 每年坠亡、被风吹撞、绳索磨损断裂事故不罕见')
P('- 一线城市清洗工日薪500-800元但招人越来越难')
P('- 幕墙越来越复杂，传统机器人越障能力不够')
P('- 机器人太重，高空负重对悬挂系统是考验')
P('')
P('哈工鹏泽市场验证：', b=True)
P('- 2026年4月23日发布会，获罗湖区政府产业资金和科技创新政策支持')
P('- 全球首款滚轮式越障清洗机器人，可跨450mm障碍')
P('- 总设计师袁立鹏教授，动力系统源自航天器/船舶装备技术')
P('- 产品已实际投入使用于深圳、苏州、武汉地标建筑')
P('- 广东数字液压研究院名誉院长许仰曾教授提供技术点评')

TB(['行业展会/活动', '时间', '内容', '来源'],
    [['哈工鹏泽发布会', '2026年4月23日', '发布4款产品，滚轮式越障机器人成亮点', '腾讯新闻'],
     ['哈工鹏泽光明区活动', '2026年5月', '政企代表和物业负责人齐聚，推介智能高空清洗', '搜狗微信'],
     ['哈工鹏泽场景创新大会', '2026年', '出席深港同心罗湖创景大会', '搜狗微信'],
     ['哈工鹏泽绿色饭店促进会', '2026年', '参展深圳绿色饭店发展促进会', '搜狗微信'],
     ['CCE上海清洁展', '2026年', '多家高空清洁机器人品牌参展', '展会公众号'],
     ['无人系统清洁论坛', '2026年深圳', '高空幕墙清洗机器人企业及物业公司参会', '深安协公众号']])

H('1.3 核心竞品（2026年信息）')
add_img('哈工鹏泽_GE02_滚轮越障机器人.jpg', '哈工鹏泽GE-02滚轮式越障高空清洗机器人')
P('哈工鹏泽（深圳）机器人技术有限公司（本公司）。2026年4月23日腾讯新闻确认：全球首款滚轮式越障清洗机器人，可跨450mm障碍，动力源航天/船舶技术。已在深圳/苏州/武汉地标建筑使用。近期活动：光明区街道政企推介、绿色饭店促进会、场景创新大会。')

add_img('埃欧珞_高空幕墙清洁机器人_产品图.png', '埃欧珞高空幕墙清洁机器人')
P('埃欧珞（杭州）科技有限公司。灵动跳跃Rs + 磐石Rx。售价数十万元/台。直销2台起/代理10台起/可租赁/省市级代理制。2023年拓展至光伏面板清洁。来源：快鲤鱼报道。')

P('凌度（广东）智能科技发展有限公司。X-Human系列高空幕墙清洗机器人获广东省名优高新技术产品称号。来源：凌度智能公众号。')

BR()

# ====== CH2 ======
H('二、客户需求（2024-2026年验证）')

TB(['客户/场景', '时间', '验证内容', '来源'],
    [['深圳/苏州/武汉地标建筑', '2026年4月', '哈工鹏泽产品已实际投入使用', '腾讯新闻'],
     ['深圳罗湖区政府', '2026年4月', '智能机器人为重点产业，释放100+AI场景', '腾讯新闻'],
     ['深圳光明区街道', '2026年5月', '政企合作智能高空清洗推介活动', '搜狗微信'],
     ['全国物业/清洁公司', '2026年4月', '日薪500-800元招不到人，机器人替代必然', '腾讯新闻'],
     ['全国幕墙清洗行业', '2026年4月', '越障能力不足是行业技术瓶颈', '腾讯新闻'],
     ['CCE上海清洁博览会', '2026年', '多家清洗机器人品牌参展，行业关注度提升', 'CCE官方公众号'],
     ['深圳物业（法规驱动）', '持续', '深圳政府令要求外墙至少每年清洗一次', '深圳政府令']])

P('注：地标建筑具体名称未在报道中披露，如需确认可联系哈工鹏泽销售团队。', s=9)

BR()

# ====== CH3 ======
H('三、产品与模式')
P('GE-02滚轮式越障450mm，2026年4月腾讯新闻确认为解决行业技术瓶颈的产品。总设计师袁立鹏教授。商业模式参考竞品：销售（数十万元/台）+ 租赁 + 代理。建议初期以租赁模式为主打。')

BR()

# ====== CH4 ======
H('四、渠道与生态')
TB(['渠道', '验证', '时间'],
    [['罗湖区政府', '产业资金+科技创新政策支持', '2026年4月'],
     ['光明区街道', '政企合作推介活动', '2026年5月'],
     ['CCE上海清洁展', '行业展会推广', '2026年'],
     ['无人系统清洁论坛', '深圳行业论坛', '2026年']])

BR()

# ====== CH5 ======
H('五、结论')
P('1. 市场规模150-180亿元（2025-2026年预测），且快速增长')
P('2. 越障450mm是行业技术瓶颈，腾讯新闻2026年4月确认GE-02已突破')
P('3. 产品已实际投入使用：深圳/苏州/武汉地标建筑（腾讯新闻2026年4月）')
P('4. 渠道已验证：罗湖区政府+光明区街道政企合作')
P('5. 安全替代+劳动力短缺是核心驱动力，2026年报道验证')

doc.add_paragraph()
P('数据来源：', s=9)
P('1. 腾讯新闻2026年4月23日：' + QQ, s=9)
P('2. 博研咨询市场报告（豆丁网）2024年10月', s=9)
P('3. 搜狐行业报道2025年2月', s=9)
P('4. 搜狗微信搜索哈工鹏泽/埃欧珞/凌度智能（2026年）', s=9)
P('5. CCE上海清洁博览会/深安协公众号（2026年）', s=9)

path = os.path.join(OUT, '高空清洗机器人市场快速调研报告.docx')
doc.save(path)

t = ''
for p in doc.paragraphs:
    t += p.text
for tb in doc.tables:
    for r in tb.rows:
        for c in r.cells:
            t += c.text
hz = sum(1 for ch in t if '\u4e00' <= ch <= '\u9fff')
print(f'Hanzi: {hz}')
print(f'Total: {len(t)}')

img_c = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_c += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Images: {img_c}')

pdf_path = path.replace('.docx', '.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
