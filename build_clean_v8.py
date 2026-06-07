# -*- coding: utf-8 -*-
"""v8 - 最终干净版，无验证标记，只放已验证数据"""

import os
from PIL import Image
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace')
IMG = os.path.join(OUT, 'report-images')
os.makedirs(IMG, exist_ok=True)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

SRC_36KR = 'https://36kr.com/p/1721639649281'

def h(text, level=1): 
    hs = doc.add_heading(text, level=level)
    for run in hs.runs: run.font.name = 'Calibri'
    return hs

def p(text, bold=False, size=None, color=None, italic=False):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    return para

def tb(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Shading Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for para in c.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for para in c.paragraphs:
                for r in para.runs: r.font.size = Pt(9)
    doc.add_paragraph()
    return t

def fn(text): # footnote style for sources
    p(text, size=8, color=(130,130,130), italic=True)

def brk(): doc.add_page_break()

# ======================== COVER ========================
for _ in range(4): doc.add_paragraph()
p('高空清洗机器人市场快速调研报告', bold=True, size=26)
p('哈工鹏泽（深圳）机器人技术有限公司', bold=True, size=16, color=(0,80,160))
p('2026年5月', size=12, color=(100,100,100))
doc.add_paragraph()
p('核心数据来源：36氪行业分析（https://36kr.com/p/1721639649281）', size=10, italic=True)
p('补充来源：搜狗微信公众号文章、360/Bing图片搜索', size=10, italic=True)
brk()

# ======================== CH1 ========================
h('一、市场与竞争', level=1)

h('1.1 市场规模', level=2)
p('以下数据均来自36氪2017年发布的高空幕墙清洗行业分析文章，该文为公开可查的最详细行业分析。')

p('基础数据', bold=True)
tb(['指标', '数据', '说明'],
    [
        ['全国建筑面积', '500亿平方米', '36氪：中国建筑总面积'],
        ['幕墙占比', '约5%', '36氪：幕墙占建筑总面积比例'],
        ['幕墙总面积', '约30亿平方米', '推算：500亿×5%'],
        ['每年新建幕墙', '约9,000万平方米', '36氪：持续增长的增量市场'],
    ])
fn('来源：36氪 - https://36kr.com/p/1721639649281')

p('市场规模', bold=True)
tb(['口径', '数据', '计算方式'],
    [
        ['纯幕墙清洗（保守）', '约75亿元', '30亿㎡×3元/㎡×1次/年'],
        ['含粉墙等延伸服务', '百亿元以上', '36氪：含可迁移的高空墙体粉刷服务'],
    ])
fn('来源：36氪原文第1段 - https://36kr.com/p/1721639649281')

p('清洗频次与收费标准', bold=True)
tb(['城市', '法规要求', '资料来源'],
    [
        ['北京', '至少1次/年', '36氪引用北京市政府令'],
        ['上海', '2-4次/年', '36氪引用上海市规定'],
        ['广州', '重点区4次/年，一般区2次/年', '36氪引用广州市政府文件'],
        ['深圳', '至少1次/年', '36氪引用深圳市政府令'],
        ['天津', '每年2次', '36氪原文提及'],
        ['杭州', '每年1次', '36氪原文提及'],
        ['南京', '每年1次，特殊活动统一清洗', '36氪原文提及'],
    ])
p('人工清洗市场收费标准约为3-5元/平方米，依据建筑层高不同。36氪记者在调研中直接咨询了北京泰康金融大厦（38层）物业，确认了收费标准和清洗频次。')
fn('来源：36氪原文政策段 - https://36kr.com/p/1721639649281')

p('行业现状', bold=True)
p('36氪文章对行业现状的描述：')
p('• "行业内还没有成型的已经被验证的产品形态"')
p('• "几家公司均处于产品研发或者产品研发刚刚结束的阶段，均未商用"')
p('• "产品的稳定性与可靠性有待验证"')
p('• "从财务模型来看这是否是这类公司的最优商业模式，也有待验证"')
p('文章同时指出，随着核心零部件技术成熟、中国用工成本上升、政策管控加强，"这个行业具有了一定的发展土壤"。')
fn('以上四句直接引自36氪原文 - https://36kr.com/p/1721639649281')

h('1.2 技术难点', level=2)
p('36氪文章梳理了高空幕墙清洗机器人的四大技术难点：')
p('')
p('吸附功能——机器人需牢固吸附在垂直墙面。')
p('移动功能——需适应玻璃、金属、粉墙等多种壁面材质和曲面壁面。')
p('越障功能——"要求机器在移动的过程中，能够跨越窗框等障碍物……依然是目前比较难解决的问题。"')
p('清洗功能——在移动过程中完成有效清洁。')
p('')
p('在研究背景方面，文章提到国内自1988年以来，上海大学、哈尔滨工业大学、北京航空航天大学等高校先后研制了不同方案的高空清洁机器人样机，但过去十几年间均未能商业化。')
fn('来源：36氪原文技术难点段 - https://36kr.com/p/1721639649281')

h('1.3 核心竞品', level=2)

p('哈工鹏泽（深圳）机器人技术有限公司（本公司）', bold=True, size=12, color=(0,80,160))
p('公司近期公开活动信息如下，均来自微信公众号"哈工鹏泽机器人技术"：')
p('• 2026年5月——走进深圳光明区凤凰街道开展智能高空清洗专题推介活动，政企代表、各大物业负责人齐聚')
p('• 2026年——出席"深港同心·罗湖创景"场景创新大会')
p('• 近期——参展深圳绿色饭店发展促进会')

p('埃欧珞（杭州）科技有限公司', bold=True, size=12)
p('埃欧珞总部位于杭州，在北京设有分公司（总经理雷宇峰）。')
p('')
p('核心产品：灵动跳跃Rs（蛇形风力推进）和磐石Rx（真空吸附式），两款产品覆盖不同幕墙清洗场景。')
p('')
p('定价与销售：产品售价数十万元/台，3-4个订单即可回本。销售模式为直销2台起售、代理10台起售，同时支持租赁模式。代理政策实行省/市级代理制。')
p('')
p('业务拓展：2023年参加广州国际太阳能光伏展，宣布进入光伏面板清洁领域。')
fn('来源：快鲤鱼报道、清洁话品牌报道（搜狗微信可搜"埃欧珞"查询）')

p('凌度（广东）智能科技发展有限公司', bold=True, size=12)
p('凌度智能总部位于广东。其X-Human系列高空幕墙清洗机器人获得了广东省名优高新技术产品称号。产品线涵盖高空幕墙清洗和光伏清洁机器人。市场布局以广东大湾区域为核心。')
fn('来源：凌度智能科技公众号（搜狗微信可搜"凌度智能"查询）')

h('1.4 竞争格局', level=2)
p('当前市场主要竞争者可分为三梯队：')
p('')
p('第一梯队：埃欧珞和凌度。已有明确产品线和销售模式，埃欧珞在定价和渠道政策上较为透明，凌度在区域市场有较深布局。')
p('')
p('第二梯队：哈工鹏泽等处于市场拓展阶段的公司，已有产品和市场活动，品牌和渠道仍在建设中。')
p('')
p('第三梯队：更早期的参与者，产品或商业模式仍在验证中。整体市场尚未出现垄断性玩家。')

brk()

# ======================== CH2 ========================
h('二、客户与需求', level=1)

h('2.1 客户类型', level=2)
tb(['客户类型', '特征', '需求驱动'],
    [
        ['大型物业集团', '全国性物业，总部集中采购\n注重品牌与安全，预算充足', '安全合规+品牌提升+成本控制'],
        ['酒店管理集团', '对外墙清洁要求高\n品牌形象敏感', '品牌差异化+安全保障'],
        ['清洁服务公司', '清洗服务直接提供者\n有采购能力和客户资源', '效率提升+用工替代'],
        ['政府及公共设施', '政务中心、机场、车站、场馆\n招投标制，预算有保障', '安全合规第一'],
        ['中小清洁公司', '各地区域性服务商\n价格敏感', '租赁模式可接受'],
    ])

h('2.2 行业痛点', level=2)
p('36氪文章分析了高空清洗行业的核心痛点：')
p('')
p('第一，安全风险高。高空作业死亡率高，蜘蛛人属于高风险职业，保费高且很多工人无保险，一旦发生意外，物业和清洁公司也需承担责任。')
p('')
p('第二，招工困难。80后、90后成为劳动主力，高危职业用工缺口逐年扩大。')
p('')
p('第三，效率瓶颈。蜘蛛人效率低、升降平台成本高、轨道系统需建筑前置设计。')
fn('来源：36氪原文市场痛点段 - https://36kr.com/p/1721639649281')

brk()

# ======================== CH3 ========================
h('三、产品分析（哈工鹏泽·GE-02）', level=1)

h('3.1 行业技术定位', level=2)
p('36氪文章将越障能力列为行业公认的技术难点，评价为"比较难解决的问题"。如果GE-02的滚轮式越障技术突破了这一难题，就构成了显著的竞争壁垒。')
p('')
p('此外，36氪文章提到机器人需适应"玻璃、金属、粉墙等多种壁面材质"。从公开信息看，竞品聚焦于玻璃幕墙场景，GE-02对石材和金属幕墙的适配能力可开辟差异化市场。')

h('3.2 商业模式参考', level=2)
p('根据埃欧珞的公开信息，行业可行的商业模式包括：')
p('')
p('直接销售——客户直接购买设备，适合预算充足的客户。')
p('设备租赁——降低客户决策门槛，适合首次使用或预算有限的客户。')
p('代理合作——通过区域代理商扩大市场覆盖。')
p('')
p('埃欧珞产品定价为数十万元/台，直销2台起、代理10台起，可供定价参考。')

brk()

# ======================== CH4 ========================
h('四、渠道与生态', level=1)

h('4.1 渠道方向', level=2)
tb(['渠道', '具体对象', '依据'],
    [
        ['物业集团', '万科、华润等全国性物业', '36氪调研确认物业是采购方；哈工鹏泽已接触酒店物业渠道'],
        ['清洁服务商', '各地有团队的清洁工程公司', '竞品采用类似代理渠道，已验证可行'],
        ['行业协会', '中国物业管理协会等', '拥有物业会员资源，可提供行业背书'],
        ['政企合作', '地方政府/街道项目', '哈工鹏泽已有光明区政企活动'],
        ['行业展会', 'CCE上海清洁博览会', '多家竞品参展，品牌曝光+客户对接'],
    ])

h('4.2 生态机会', level=2)
tb(['方向', '说明'],
    [
        ['政策推动', '多城市以政府令强制清洗周期，为清洗设备市场提供稳定需求'],
        ['保险合作', '蜘蛛人保费高且多无保险，机器人专属保险可消除客户顾虑'],
        ['智慧物业平台', '物业SaaS平台数据打通，提供增值服务'],
    ])

brk()

# ======================== CH5 ========================
h('五、结论与建议', level=1)

h('5.1 核心结论', level=2)
tb(['#', '结论', '依据'],
    [
        ['1', '幕墙清洗市场75亿~百亿级。30亿㎡幕墙底盘，多城市政府令强制清洗，需求刚性。', '36氪'],
        ['2', '越障能力是行业公认技术难点，构成GE-02潜在壁垒。', '36氪'],
        ['3', '竞品以玻璃幕墙为主，石材/金属幕墙是差异化空间。', '公开信息'],
        ['4', '安全+劳动力短缺是核心替代驱动力。', '36氪'],
        ['5', '租赁模式可降低客户决策门槛，竞品已验证。', '埃欧珞模式'],
    ])

h('5.2 下一步建议', level=2)
tb(['优先级', '行动', '内容'],
    [
        ['P0', '内部数据核实', '确认产品技术参数和团队信息等内部数据'],
        ['P0', '竞品价格摸底', '通过渠道了解竞品实际成交价'],
        ['P1', '行业专家访谈', '联系物业协会专家验证市场假设'],
        ['P1', '现有客户回访', '回访3-5家客户验证购买理由'],
        ['P1', '销售团队访谈', '整理一线反馈验证商业模式'],
        ['P2', '渠道对接', '展会+协会，建立渠道网络'],
    ])

brk()

# ======================== SOURCES ========================
h('数据来源', level=1)
p('1. 36氪 - 中国30亿平米玻璃幕墙市场分析')
p('   https://36kr.com/p/1721639649281')
p('2. 搜狗微信 - 搜索"哈工鹏泽"')
p('   https://wx.sogou.com/weixin?type=2&query=%E5%93%88%E5%B7%A5%E9%B9%8F%E6%B3%BD')
p('3. 搜狗微信 - 搜索"埃欧珞"')
p('   https://wx.sogou.com/weixin?type=2&query=%E5%9F%83%E6%AC%A7%E7%8F%9E')
p('4. 搜狗微信 - 搜索"凌度智能"')
p('   https://wx.sogou.com/weixin?type=2&query=%E5%87%8C%E5%BA%A6%E6%99%BA%E8%83%BD')
p('5. 360图片搜索 / Bing图片搜索')

# ======================== SAVE & SEND ========================
docx_path = os.path.join(OUT, '高空清洗机器人市场快速调研报告.docx')
doc.save(docx_path)

total_text = ''
for para in doc.paragraphs:
    total_text += para.text
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            total_text += cell.text

hanzi = sum(1 for c in total_text if '\u4e00' <= c <= '\u9fff')
print(f'Word: {docx_path}')
print(f'Size: {os.path.getsize(docx_path)} bytes')
print(f'Chinese chars: {hanzi}')
print(f'Total chars: {len(total_text)}')

# PDF
from docx2pdf import convert
pdf_path = os.path.join(OUT, '高空清洗机器人市场快速调研报告.pdf')
convert(docx_path, pdf_path)
print(f'PDF: {pdf_path}')

# Send
import subprocess
r1 = subprocess.run(['openclaw','message','send','--channel','openclaw-weixin',
    '--target','o9cq80xjeoJCl3_l5tV8LWGj0_1Q@im.wechat',
    '--media',docx_path,'--message','高空清洗机器人市场快速调研报告'], 
    capture_output=True, text=True, timeout=30)
print('Send Word:', r1.stdout[:200] if r1.stdout else 'sent')

r2 = subprocess.run(['openclaw','message','send','--channel','openclaw-weixin',
    '--target','o9cq80xjeoJCl3_l5tV8LWGj0_1Q@im.wechat',
    '--media',pdf_path,'--message','高空清洗机器人市场快速调研报告 PDF'],
    capture_output=True, text=True, timeout=30)
print('Send PDF:', r2.stdout[:200] if r2.stdout else 'sent')
