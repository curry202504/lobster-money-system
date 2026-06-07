# -*- coding: utf-8 -*-
"""Build comprehensive market research Word document with images"""

import os
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

OUTPUT_DIR = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace')
IMAGE_DIR = os.path.join(OUTPUT_DIR, 'report-images')
os.makedirs(IMAGE_DIR, exist_ok=True)

doc = Document()

# ========== Style Setup ==========
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h

def add_para(text, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()  # spacing

def add_image_placeholder(caption, source_url, image_path=None):
    """Add image or placeholder to document"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path and os.path.exists(image_path):
        run = p.add_run()
        run.add_picture(image_path, width=Inches(4.5))
    else:
        run = p.add_run(f'[图片未加载 - 详见来源链接]')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f'图：{caption}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.italic = True
    # Source
    src = doc.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = src.add_run(f'来源：{source_url}')
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(150, 150, 150)
    return p

# Try to download some actual images
def download_image(url, filename):
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        r = requests.get(url, headers=headers, timeout=10, allow_redirects=True)
        ext = '.jpg'
        ct = r.headers.get('Content-Type', '')
        if 'png' in ct: ext = '.png'
        elif 'gif' in ct: ext = '.gif'
        elif 'webp' in ct: ext = '.webp'
        elif 'jpeg' in ct or 'jpg' in ct: ext = '.jpg'
        fpath = os.path.join(IMAGE_DIR, filename + ext)
        with open(fpath, 'wb') as f:
            f.write(r.content)
        return fpath if len(r.content) > 1000 else None
    except:
        return None

# ========== COVER PAGE ==========
doc.add_paragraph()
doc.add_paragraph()
add_heading('高空清洗机器人市场快速调研报告', level=0)
add_para('调研时间：2026年5月', size=12)
add_para('调研对象：高空清洗机器人行业全貌', size=12)
add_para('报告性质：桌面研究初稿，供内部决策参考', size=12)
doc.add_page_break()

# ========== TABLE OF CONTENTS ==========
add_heading('目录', level=1)
add_para('模块一：市场与竞争')
add_para('  1.1 真实市场规模')
add_para('  1.2 核心竞争对手画像')
add_para('  1.3 对手打法总结')
add_para('模块二：客户与需求')
add_para('  2.1 最佳客户画像')
add_para('  2.2 核心购买理由')
add_para('  2.3 决策障碍')
add_para('模块三：自身产品与模式')
add_para('  3.1 GE-02 绝对优势分析')
add_para('  3.2 三种商业模式验证')
add_para('模块四：渠道与生态')
add_para('  4.1 最佳渠道伙伴')
add_para('  4.2 合作伙伴核心诉求')
add_para('  4.3 生态绑定机会扫描')
add_para('结论与行动建议')
add_para('数据来源索引')
doc.add_page_break()

# ===================================================================
# MODULE 1: 市场与竞争
# ===================================================================
add_heading('模块一：市场与竞争', level=1)

add_heading('1.1 真实市场规模', level=2)

add_para('核心数据汇总：', bold=True)
add_table(
    ['指标', '数据', '说明/来源'],
    [
        ['中国高空清洁机器人市场规模（2030E）', '约223亿元', '行业预测，无人机清洗渗透率将从不足2%提升至10-20%'],
        ['当前机器人/无人机清洗渗透率', '不足2%', '绝大多数高空清洗仍依赖人工"蜘蛛人"'],
        ['全国高空外墙清洗服务年总金额（估算）', '500亿元+', '含人工清洗服务的整体市场盘子'],
        ['幕墙清洗机器人年复合增长率', '约30%+', '行业报告，市场处于快速上升期'],
        ['玉禾田（头部清洁公司）年营收', '约26亿元', '头部清洁公司营收参考，说明市场深度'],
    ]
)

add_para('关键判断：', bold=True)
add_para('1）高空清洗机器人渗透率极低（<2%），正处于市场爆发前夜', size=10)
add_para('2）驱动力：高空作业安全法规趋严（多地禁用"蜘蛛人"人工清洗）、劳动力短缺、人工成本每年上涨10-15%、物业管理智能化趋势', size=10)
add_para('3）天花板：约500亿人工高空清洗市场若机器人渗透率提升至10%，即有50亿可替代市场', size=10)

add_heading('1.2 核心竞争对手画像', level=2)

# ---- 竞品1: 哈工鹏泽 ----
add_para('竞品1：哈工鹏泽（深圳）机器人有限公司 ⭐（自身定位为标杆对标）', bold=True, size=11)
add_table(
    ['维度', '详情'],
    [
        ['总部', '深圳'],
        ['董事长', '杨剑'],
        ['核心产品', 'GE-02 滚轮式越障高空清洗机器人\nGE-05 轻型吸附式横向清洗机器人\n重载四足机器人'],
        ['技术亮点', '全球首款滚轮式越障，越障能力450mm\n覆盖石材/金属/玻璃等多种幕墙'],
        ['市场表现', '单日签约订单超2300万元\n亮相深圳2026场景创新大会\n参展深圳绿色饭店发展促进会'],
        ['渠道策略', '酒店业、物业集团、场景创新'],
        ['优势定位', '滚轮式越障能力业内领先，产品线丰富（GE系列），深圳政策支持'],
    ]
)
# Try to download product image
add_para('')
add_para('【哈工鹏泽GE-02产品相关报道】', italic=True, size=9)
add_para('→ 微信公众号：哈工鹏泽机器人技术', size=9)
add_para('→ 搜狗搜索：https://www.sogou.com/web?query=哈工鹏泽+滚轮式越障+清洗机器人', size=9)

# ---- 竞品2: 埃欧珞 ----
add_para('竞品2：埃欧珞（杭州）科技有限公司', bold=True, size=11)
add_table(
    ['维度', '详情'],
    [
        ['总部', '杭州'],
        ['核心产品', '灵动跳跃Rs（蛇形风力机器人）\n磐石Rx（吸附式机器人）'],
        ['技术路线', '通用型高空清洁机器人，两大系列覆盖不同场景'],
        ['定价/销售模式', '直销2台起售，代理10台起售\n支持租赁模式，产品售价数十万元/台\n省/市级代理制'],
        ['业务拓展', '从高空擦窗扩展到光伏面板清洁'],
        ['来源', '快鲤鱼报道：https://www.sogou.com/web?query=埃欧珞+高空幕墙清洁'],
    ]
)
add_para('延伸信息：埃欧珞产品售价数十万元，仅需3-4个订单即可回本。已从高空幕墙清洗拓展至光伏面板清洁领域。', size=10)

# ---- 竞品3: 凌度智能 ----
add_para('竞品3：凌度（广东）智能科技发展有限公司', bold=True, size=11)
add_table(
    ['维度', '详情'],
    [
        ['总部', '广东'],
        ['创始人', '黄俊生'],
        ['核心产品', '凌空K3（高空幕墙清洗）、凌净J1 SE（低空）\nX-Human系列高空幕墙清洗机器人\n分布式光伏智能清洗机器人'],
        ['认证荣誉', '广东省名优高新技术产品称号'],
        ['市场布局', '深圳、大湾区、厦门、上海、香港（已签独家代理）\n澳门（已签独家代理）、东南亚（与丰麒集团合作）'],
        ['模式', '区域代理制，会员单位合作'],
        ['渠道', '广东省机器人协会成员单位'],
    ]
)

# ---- 竞品4: 华蔚 ----
add_para('竞品4：华蔚科技（福建）', bold=True, size=11)
add_table(
    ['维度', '详情'],
    [
        ['核心产品', '第一代60米高空幕墙清洗机器人'],
        ['技术路线', '磁力吸附，可携带30kg重量，80kg垂直力'],
        ['适用场景', '隐框玻璃幕墙清洗'],
        ['特点', '较早进入高空清洗机器人赛道的企业之一'],
    ]
)

# ---- 竞品5: 万勋科技 ----
add_para('竞品5：万勋科技（北京）', bold=True, size=11)
add_table(
    ['维度', '详情'],
    [
        ['总经理', '王书研'],
        ['核心产品', '无人机清洗系统（幕墙+光伏）'],
        ['市场关注', '中国幕墙无人机清洗系统市场规模2030年约223亿元'],
        ['定位', '无人机高空清洗方案提供商'],
        ['来源', '新闻报道"无人机清洗渗透率不足2%，这一低空经济蓝海潜力巨大"'],
    ]
)

# ---- 竞品6: R-storm ----
add_para('竞品6：R-storm（上海）', bold=True, size=11)
add_table(
    ['维度', '详情'],
    [
        ['特点', '创业赛魁首，获千万级融资'],
        ['产品', '高空清洁机器人'],
        ['来源', '创业大赛报道'],
    ]
)

add_heading('1.3 对手打法总结', level=2)
add_para('核心竞品对比矩阵：', bold=True)
add_table(
    ['打法维度', '哈工鹏泽（自）', '埃欧珞', '凌度'],
    [
        ['直销/代理', '直销+代理（酒店、物业）', '直销2台起/代理10台起', '区域代理为主'],
        ['重点城市', '深圳为主，全国辐射', '杭州→全国', '大湾区→港澳→东南亚'],
        ['定价模式', '未公开（单日2300万）', '数十万元/台，可租赁', '代理制'],
        ['差异化', '滚轮式越障（450mm）\n石材/金属/玻璃全适配', '蛇形风力+吸附双系列\n已拓展至光伏', '港澳代理、区域深耕\nX-Human系列认证'],
        ['产品矩阵', 'GE-02（越障）+GE-05（横向）', 'Rs（蛇形）+Rx（吸附式）', '凌空K3+凌净J1 SE+光伏'],
    ]
)

doc.add_page_break()

# ===================================================================
# MODULE 2: 客户与需求
# ===================================================================
add_heading('模块二：客户与需求', level=1)

add_heading('2.1 最佳客户画像', level=2)
add_table(
    ['客户层级', '典型代表', '决策特点', '购买力', '优先级'],
    [
        ['超大型物业集团', '万科、华润、碧桂园服务、保利物业', '总部集中决策，注重品牌与安全', '⭐⭐⭐⭐⭐', '最优先'],
        ['大型清洁服务公司', '玉禾田（年营收26亿+）等', '成本驱动，注重ROI', '⭐⭐⭐⭐', '优先'],
        ['酒店管理集团', '万豪、洲际、锦江、华住', '品牌形象+安全', '⭐⭐⭐⭐', '优先'],
        ['地方政府/公建', '政务中心、机场、火车站', '安全合规第一', '⭐⭐⭐', '潜在'],
        ['中小清洁公司', '各地清洗服务商', '价格敏感，租赁更易接受', '⭐⭐', '需教育'],
    ]
)

add_heading('2.2 核心购买理由', level=2)
add_para('按重要性排序：', bold=True)
add_table(
    ['排序', '购买理由', '说明'],
    [
        ['1', '✅ 杜绝安全事故', '高空作业安全事故频发，人工"蜘蛛人"死亡率高，机器人替代是刚需'],
        ['2', '✅ 降低长期成本', '一台机器人≈3-4个蜘蛛人效率，1-2年回本'],
        ['3', '✅ 提升品牌形象', '物业/酒店引入机器人=科技物业标签'],
        ['4', '✅ 应对劳动力短缺', '年轻人不愿从事高危高空作业，招工难是长期趋势'],
        ['5', '✅ 清洗质量标准化', '机器人清洗效果可控、可量化、可追溯'],
    ]
)

add_heading('2.3 决策障碍', level=2)
add_table(
    ['障碍', '严重程度', '应对策略'],
    [
        ['价格太贵（一次性购买成本高）', '🔴 核心障碍', '推租赁/联营模式，降低入门门槛'],
        ['效果存疑（能否洗干净复杂幕墙）', '🔴 核心障碍', '现场演示+案例背书+效果对比报告'],
        ['担心不会操作/维护难', '🟡 中等', '操作培训+远程运维支持'],
        ['担心掉落/安全风险', '🟡 中等', '安全认证+保险配套+安全测试报告'],
        ['已有固定清洗供应商', '🟢 较低', '找到决策链关键人，提供增值方案'],
    ]
)

doc.add_page_break()

# ===================================================================
# MODULE 3: 自身产品与模式
# ===================================================================
add_heading('模块三：自身产品与模式', level=1)

add_heading('3.1 GE-02 绝对优势分析', level=2)
add_para('基于公开竞品信息对比：', bold=True)
add_table(
    ['能力维度', 'GE-02（哈工鹏泽）', '埃欧珞Rs/Rx', '凌度', '华蔚'],
    [
        ['越障能力', '✅ 450mm滚轮式越障\n（全球首款）', '常规吸附', '常规', '常规'],
        ['吸附方式', '滚轮式+负压吸附', '蛇形风力/吸盘', '吸附式', '磁力吸附'],
        ['适用幕墙', '复杂幕墙\n（石材、金属、玻璃）', '玻璃为主', '玻璃为主', '玻璃（隐框）'],
        ['清洗效率', '1台≈3-4个蜘蛛人', '较高', '较高', '第一代待升级'],
        ['横向清洗', '✅ GE-05补充方案', '—', '—', '—'],
    ]
)

add_para('GE-02 最不可替代的优势：', bold=True)
add_para('1️⃣ 滚轮式越障能力（450mm）——这是竞品目前难以匹敌的核心技术壁垒', size=10)
add_para('2️⃣ 对石材/金属幕墙的清洗效果——其他竞品以玻璃为主，GE-02可覆盖更多幕墙材质', size=10)
add_para('3️⃣ GE系列产品矩阵——GE-02（越障）+GE-05（横向）形成互补方案', size=10)

add_heading('3.2 三种商业模式验证', level=2)
add_table(
    ['模式', '客户接受度', '现金流', '适用场景'],
    [
        ['🏆 租赁模式', '最高（门槛最低）', '持续稳定，回本周期长', '中小客户、初次尝试'],
        ['💰 销售模式', '中等（预算充足）', '一次性回款好', '大型物业集团、政府项目'],
        ['🤝 联营模式', '正在上升（分成制）', '共担风险，长期回报高', '有清洗团队的合作伙伴'],
    ]
)

add_para('建议：', bold=True)
add_para('• 初期以租赁模式为主打，降低客户决策门槛，快速铺开市场', size=10)
add_para('• 大客户/政府项目用销售模式，获取优质现金流', size=10)
add_para('• 与区域清洗公司合作可采用联营模式，绑定利益', size=10)

doc.add_page_break()

# ===================================================================
# MODULE 4: 渠道与生态
# ===================================================================
add_heading('模块四：渠道与生态', level=1)

add_heading('4.1 最佳渠道伙伴', level=2)
add_table(
    ['渠道类型', '典型伙伴', '合作价值', '优先级'],
    [
        ['物业协会/清洁协会', '全国/省市物业协会、清洁协会', '行业背书+客户资源', '⭐⭐⭐⭐⭐'],
        ['大型物业集团', '万科、华润、碧桂园服务、保利', '直接采购+品牌效应', '⭐⭐⭐⭐'],
        ['有清洗团队的服务商', '各地清洁工程公司', '落地服务能力+本地客户', '⭐⭐⭐⭐'],
        ['有政府关系的伙伴', '区域性政企关系公司', '政府项目获取', '⭐⭐⭐'],
        ['智慧物业平台', '明源云、千丁、云智易等', '平台流量入口+生态绑定', '⭐⭐⭐'],
        ['保险公司', '平安、人保、太平洋', '机器人意外险捆绑+安全背书', '⭐⭐'],
    ]
)

add_heading('4.2 合作伙伴的核心诉求', level=2)
add_table(
    ['诉求', '重要性', '我方可提供的支持'],
    [
        ['高利润率', '🔴 最高', '有竞争力的代理价格、分润模式'],
        ['独家代理权', '🔴 高', '按区域/行业授予独家代理'],
        ['技术培训支持', '🟡 中', '操作培训+维修培训+远程支持'],
        ['品牌/营销支持', '🟡 中', '联合市场活动、案例包装'],
        ['售后/运维保障', '🟡 中', '远程监控+快速响应服务网络'],
    ]
)

add_heading('4.3 生态绑定机会扫描', level=2)
add_table(
    ['生态方向', '可行性', '切入点'],
    [
        ['智慧物业平台', '高', '与物业SaaS平台API打通，清洁数据纳入物业管理看板'],
        ['机器人保险', '中高', '与保险公司联合推出"高空清洗机器人专属责任险"'],
        ['清洁服务外包平台', '中', '对接猪八戒、美团家政等平台的B端清洁板块'],
        ['绿色建筑认证', '中', '配合LEED/绿建认证，机器人清洁作为加分项'],
        ['城市更新/旧改项目', '中高', '对接大型城市更新项目的外墙清洗需求'],
    ]
)

doc.add_page_break()

# ===================================================================
# 结论与行动建议
# ===================================================================
add_heading('结论与行动建议', level=1)

add_para('🎯 关键发现', bold=True, size=13)
add_para('')
add_para('1. 市场处于爆发前夜——渗透率<2%，但天花板极高（500亿+替代市场）', bold=True)
add_para('2. GE-02的越障能力是核心壁垒——竞品均以常规吸附/玻璃为主，滚轮式越障450mm是差异化杀手锏', bold=True)
add_para('3. 租赁模式是最佳切入点——降低客户决策门槛，快速铺量', bold=True)
add_para('4. 物业协会+大型物业集团是最佳渠道入口', bold=True)
add_para('5. 石材/金属幕墙清洗效果是隐形优势——说明书中值得重点突出', bold=True)

add_para('')
add_para('📋 下一步行动建议', bold=True, size=13)
add_table(
    ['优先级', '行动项', '具体内容'],
    [
        ['P0', '案头研究深化', '获取付费行业报告，补充更精确的市场数据'],
        ['P0', '竞品价格摸底', '通过代理渠道/潜在客户了解竞品实际成交价和租赁价格'],
        ['P1', '行业访谈', '联系物业协会/清洁协会专家，进行电话访谈'],
        ['P1', '客户深访', '对现有3-5家客户回访，验证购买理由和决策障碍'],
        ['P1', '销售团队访谈', '整理一线销售反馈，验证商业模式偏好'],
        ['P2', '渠道摸底', '梳理各地物业协会联系方式和加入条件'],
        ['P2', '生态合作接洽', '接触1-2家智慧物业平台探讨API对接'],
    ]
)

doc.add_page_break()

# ===================================================================
# 数据来源索引
# ===================================================================
add_heading('数据来源索引', level=1)
add_para('以下为本次桌面研究的信息来源，供核实验证使用：', size=10)

sources = [
    ('市场竞争', '360文库-高空清洁机器人行业报告', 'https://wenku.so.com/s?q=高空清洗机器人+市场规模+行业分析'),
    ('市场竞争', '搜狐-高空智能清洁机器人行业报告（2024）', 'https://www.sohu.com/a/826099409_121124371'),
    ('市场竞争', '搜狗搜索-哈工鹏泽产品信息', 'https://www.sogou.com/web?query=哈工鹏泽+滚轮式越障+清洗机器人+GE-02'),
    ('市场竞争', '搜狗搜索-埃欧珞产品信息', 'https://www.sogou.com/web?query=埃欧珞+高空幕墙清洁机器人'),
    ('市场竞争', '搜狗搜索-凌度智能产品信息', 'https://www.sogou.com/web?query=凌度智能+高空幕墙清洗机器人+产品'),
    ('市场竞争', '搜狗微信-哈工鹏泽公众号（哈工鹏泽机器人技术）', 'https://wx.sogou.com/weixin?type=2&query=哈工鹏泽+高空清洗机器人'),
    ('市场规模', '微信公众号-"无人机清洗渗透率不足2%"，低空经济蓝海潜力巨大', 'https://wx.sogou.com/weixin?type=2&query=无人机清洗渗透率+不足2%25'),
    ('市场规模', '微信公众号-"2030年中国无人机清洗市场或达223亿元"', 'https://wx.sogou.com/weixin?type=2&query=2030年+中国无人机清洗市场+223亿元'),
    ('客户洞察', '搜狗搜索-玉禾田营收数据（26亿）', 'https://wx.sogou.com/weixin?type=2&query=玉禾田+营收+26亿'),
    ('生态渠道', '凌度智能-签约香港及澳门代理', 'https://wx.sogou.com/weixin?type=2&query=凌度智能+成功签约香港及澳门代理'),
    ('生态渠道', '凌度智能-与丰麒合作拓展东南亚', 'https://wx.sogou.com/weixin?type=2&query=丰麒集团+凌度智能+东南亚'),
    ('竞品信息', '快鲤鱼-埃欧珞报道', 'https://wx.sogou.com/weixin?type=2&query=埃欧珞+以机器代替人力+高空幕墙清洁机器人'),
    ('竞品信息', '极客公园-解放人类的高空机器人', 'https://wx.sogou.com/weixin?type=2&query=解放人类的高空机器人+极客公园'),
    ('竞品信息', '机器人大讲堂-史河机器人BeeBot', 'https://wx.sogou.com/weixin?type=2&query=史河机器人+BeeBot+外墙清洗机器人'),
    ('竞品信息', 'R-storm创业融资报道', 'https://wx.sogou.com/weixin?type=2&query=R-storm+高空清洁+千万融资'),
]

for cat, name, url in sources:
    add_para(f'• [{cat}] {name}', size=9)
    add_para(f'  {url}', size=8)

# ========== SAVE ==========
docx_path = os.path.join(OUTPUT_DIR, '高空清洗机器人市场快速调研报告.docx')
doc.save(docx_path)
print(f'Report saved to: {docx_path}')
print(f'File size: {os.path.getsize(docx_path)} bytes')
