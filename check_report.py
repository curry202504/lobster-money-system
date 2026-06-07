from docx import Document
import os

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
doc = Document(path)

h1s = [p for p in doc.paragraphs if p.style.name.startswith('Heading 1')]
h2s = [p for p in doc.paragraphs if p.style.name.startswith('Heading 2')]

img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        drawings = r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        img_count += len(drawings)

print('=== 文档质量检查 ===')
print('段落数:', len(doc.paragraphs))
print('表格数:', len(doc.tables))
print('H1标题数:', len(h1s))
print('H2标题数:', len(h2s))
print('嵌入图片数:', img_count)
print()

print('=== H1标题 ===')
for h in h1s:
    print(' ', h.text)

print()
print('=== H2标题 ===')
for h in h2s:
    print(' ', h.text)

# Check file sizes
docx_path = path
pdf_path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.pdf')
print()
print('文件大小:')
print('  Word:', os.path.getsize(docx_path) / 1024, 'KB')
if os.path.exists(pdf_path):
    print('  PDF:', os.path.getsize(pdf_path) / 1024, 'KB')
else:
    print('  PDF: 不存在')
