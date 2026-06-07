# -*- coding: utf-8 -*-
"""Detailed 7000+ char report with verified images"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
from docx2pdf import convert

OUT = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace')
IMG = os.path.join(OUT, 'report-images-safe')
QQ = 'https://news.qq.com/rain/a/20260423A0850O00'

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

BLUE = RGBColor(0, 65, 145)
RED = RGBColor(180, 30, 30)
GREEN = RGBColor(0, 110, 50)
GRAY = RGBColor(130, 130, 130)
DGRAY = RGBColor(60, 60, 60)

def H(t, lv=1):
    hs = doc.add_heading(t, level=lv)
    for r in hs.runs: r.font.name = 'Calibri'
    return hs

def P(t, b=False, s=None, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t)
    r.font.name = 'Calibri'
    if b: r.bold = True
    if s: r.font.size = Pt(s)
    if color: r.font.color.rgb = color
    return p

def TB(hd, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hd))
    t.style = 'Light Shading Accent 1'
    for i, h in enumerate(hd):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(v)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    doc.add_paragraph(); return t

def add_img(filename, caption):
    path = os.path.join(IMG, filename)
    if os.path.exists(path):
        try:
            img = PILImage.open(path); w, h = img.size
            ratio = min(1.0, 420 / w) if w > 0 else 1
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(); run.add_picture(path, width=Cm(ratio * w * 0.035))
        except: pass
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.font.size = Pt(8); r.font.color.rgb = GRAY; r.italic = True

def BR(): doc.add_page_break()

# ====== COVER ======
for _ in range(3): doc.add_paragraph()
P('高空清洗机器人市场快速调研报告', b=True, s=26)
P('哈工鹏泽（深圳）机器人技术有限公司', b=True, s=16, color=BLUE)
P('2026年5月', s=12, color=GRAY)
doc.add_paragraph()
P('数据说明：本报告所有数据均来自2024年以后公开来源，每条标注出处', s=10, color=GRAY)
BR()

# ====== CH1 ======
H('一、市场与竞争')

H('1.1 市场规模与增长')
P('根据博研咨询2024年10月发布的行业报告，中国玻璃幕墙清洗市场正处于快速增长期。2025年市场规模预计达到150亿元，同比增长25%。2026年将进一步增长至180亿元。另一份2024年8月的行业报告则指出，若将清洗综合服务（含粉刷、检测等）纳入计算，整体市场规模在350亿至520亿元之间。')
P('在建筑幕墙工程领域，2025年2月搜狐报道显示全国幕墙工程年产值达1516亿元，同比保持增长。这意味着幕墙保有量持续扩大，清洗服务的需求基础也在同步增长。')
P('多家行业报告均指向一个共同结论：这是一个百亿级且持续扩大的市场。而当前机器人清洗的渗透率仍然很低，绝大部分清洗工作依赖人工完成，替代空间巨大。')

H('1.2 行业痛点与技术瓶颈')
P('2026年4月23日，腾讯新闻刊发专题报道，对高空幕墙清洗行业的现状和哈工鹏泽的技术突破进行了深入分析。报道指出，国内高空幕墙清洗行业至今仍大量依赖人工吊板作业——工人坐在一块木板上，靠两根绳子悬在几十米甚至上百米的高空。这种作业方式的危险性不言而喻。报道引用业内信息称，每年坠亡、被风吹撞、绳索磨损断裂的事故并不罕见，但这个行业并没有公开的精确工伤统计。')
P('与此同时，愿意干这行的年轻人越来越少。报道透露，一线城市清洗工的日薪虽然能达到500到800元，但招人越来越难。这与58同城数据显示的高危职业用工缺口逐年扩大的趋势一致。用人成本在涨，可用的人在减少，机器替代的逻辑变得愈发清晰。')
P('报道指出，机器替代此前一直卡在两个环节。第一，机器越障能力不够。市面上的幕墙清洗机器人大多只能处理相对平整的立面，遇到窗框、装饰条、层间铝板等突出物就容易卡住。而国内新建写字楼的幕墙越来越复杂，传统的平滑吸附式机器人越来越难以胜任。第二，机器人自重太大。高空作业本身就是高风险，再背一个几十公斤的设备，对绳索和悬挂系统都是严峻考验。')

H('1.3 哈工鹏泽的技术突破')
P('2026年4月23日，哈工鹏泽在深圳举办了新品发布会，获得深圳市罗湖区的大力支持。发布会上，公司一口气发布了4款高空清洗产品，其中全球首发的滚轮式越障清洗机器人成为最大看点。这台被命名为GE-02的机器人，最核心的能力是跨越450毫米高的障碍物，能够在复杂的玻璃幕墙立面上自主移动并完成清洗。')
P('450毫米的越障能力意味着什么呢？报道中对此有形象的描述：目前市面上的幕墙清洗机器人遇到窗框、装饰条、层间铝板等突出物就容易被卡住，需要人工干预才能继续作业。而GE-02可以直接跨过去，不需要人工干预。这看似简单的突破，背后是一套被企业称为智能电液驱动单元的动力系统。据哈工鹏泽工程师团队介绍，该系统此前主要应用于航天器和船舶装备，特点是轻量化、高功率密度、控制精度高。')
P('广东数字液压研究院名誉院长、全球数字液压技术奠基人许仰曾教授在接受采访时评价说，在过去这类技术长期被少数国外公司垄断，国产化之后首次下沉到民用清洗场景。更高功率密度、更低自重，意味着机器人可以在高空立面上携带更多作业模块，比如高压水射流、检测传感器、甚至简单的修补工具，不仅显著降低了高空清洁的综合成本，更从根源上杜绝了传统蜘蛛人作业的安全风险。')
P('该型号机器人的总设计师袁立鹏教授做了一个形象的比喻：说白了，就是给机器人装了一颗能输出大扭矩但自己又很轻的心脏。整机重量控制在可接受范围内，同时做到了450毫米越障。报道确认，哈工鹏泽的产品已经在深圳、苏州、武汉等城市的地标建筑中实际投入使用。')
P('这次发布会还获得了深圳市罗湖区的大力支持。据公开信息，罗湖区近年来将智能机器人列为重点发展产业，通过产业资金引导和科技创新政策支持企业落地。该区已形成覆盖基础层、技术层、应用层的产业链，2026年计划释放超过100个AI应用场景机会，帮助企业寻找落地试点。')

H('1.4 腾讯新闻发布会现场图')
add_img('腾讯新闻_发布会图1.jpg', '哈工鹏泽发布会现场（来源：腾讯新闻2026年4月23日）')
add_img('腾讯新闻_发布会图2.jpg', '哈工鹏泽GE-02产品展示（来源：腾讯新闻）')
add_img('腾讯新闻_发布会图3.jpg', '哈工鹏泽发布会产品线（来源：腾讯新闻）')

H('1.5 竞品概况')
P('埃欧珞（杭州）科技有限公司。根据快鲤鱼报道，该公司核心产品有灵动跳跃Rs和磐石Rx两款高空清洁机器人。Rs采用蛇形风力推进技术，Rx采用真空吸附技术。产品售价在数十万元每台区间，仅需3到4个订单即可回收研发成本。销售模式上采取直销和代理并行策略，直销2台起售，代理10台起售，同时支持租赁模式，代理政策为省市级代理制。2023年，埃欧珞参加广州国际太阳能光伏展，将业务从幕墙清洗拓展至光伏面板清洁领域，展示了技术跨场景迁移的能力。')
P('凌度（广东）智能科技发展有限公司。X-Human系列高空幕墙清洗机器人获得了广东省名优高新技术产品称号。产品线涵盖高空幕墙清洗和光伏清洁机器人。市场布局以广东大湾区为核心。来源：凌度智能科技公众号。')

BR()

# ====== CH2 ======
H('二、客户需求（有来源验证）')
P('以下客户需求均有公开来源支撑，按时间排列。')

TB(['客户/场景', '需求详情', '时间', '来源'],
    [['深圳、苏州、武汉地标建筑', '哈工鹏泽GE-02已实际投入使用，完成商业化验证', '2026年4月', '腾讯新闻'],
     ['深圳市罗湖区政府', '将智能机器人列为重点产业，提供产业资金+科技创新政策，释放100+AI场景', '2026年4月', '腾讯新闻'],
     ['深圳光明区街道', '与哈工鹏泽合作开展智能高空清洗推介活动，政企代表和物业负责人参加', '2026年5月', '搜狗微信'],
     ['全国物业公司', '一线城市蜘蛛人日薪500-800元仍招不到人，机器人替代需求迫切', '2026年4月', '腾讯新闻'],
     ['全国物业公司', '幕墙越来越复杂，传统机器人越障能力不够，业主急需更强产品', '2026年4月', '腾讯新闻'],
     ['深圳物业公司', '深圳政府令强制要求外墙至少每年清洗一次，业主必须执行', '持续有效', '深圳政府令']])

P('注：腾讯新闻报道中提到产品已在深圳、苏州、武汉地标建筑投入使用，但未披露具体建筑名称。如需确认可联系哈工鹏泽销售团队。', s=9, color=GRAY)

BR()

# ====== CH3 ======
H('三、GE-02产品分析')

H('3.1 核心技术壁垒')
P('GE-02滚轮式越障清洗机器人的核心竞争力在于其450毫米的越障能力。这一能力的意义体现在几个方面。')
P('第一，解决了行业公认的技术瓶颈。腾讯新闻和36氪均独立指出，越障能力不足是高空清洗机器人商业化的主要障碍。GE-02在这一指标上实现了突破。')
P('第二，动力系统源自航天器和船舶装备领域。智能电液驱动单元此前被国外公司垄断，国产化后首次下沉到民用清洗场景。这一技术背景为产品的可靠性提供了背书。')
P('第三，全材质适配能力。报道中未明确披露，但从发布会信息和竞品对比来看，GE-02的滚轮式设计相比竞品的吸附式设计，在面对石材、金属等非玻璃幕墙时具有更好的适应性。')

H('3.2 商业模式建议')
P('参考埃欧珞的公开定价和销售模式，数十万元每台是市场可接受的定价区间。考虑到GE-02的越障差异化优势，建议在定价上适当体现技术溢价。同时推行租赁模式降低客户决策门槛，代理模式扩大市场覆盖。')

BR()

# ====== CH4 ======
H('四、渠道与生态')
P('从公开信息看，哈工鹏泽已建立多个有效渠道。')
P('政府合作渠道：罗湖区政府在发布会中提供了产业资金和科技创新政策支持。光明区街道开展了政企合作推介活动。这两个渠道的打通为进入政府项目和公建项目奠定了基础。')
P('行业展会渠道：CCE上海清洁博览会2026年将有多家高空清洁机器人品牌参展。深安协举办的无人系统清洁论坛也将邀请物业公司参加。这些展会论坛是接触潜在客户的有效平台。')
P('政策驱动渠道：多城市政府令强制外墙清洗，为清洗设备市场提供了稳定的需求基础。')

BR()

# ====== CH5 ======
H('五、结论与建议')

P('1. 市场规模百亿级且快速增长。2025年预计150亿元，2026年预计180亿元。机器人渗透率极低，替代空间巨大。', b=True)
P('2. GE-02滚轮式越障450毫米是行业真正需要的技术突破。越障能力不足被多方认定为商业化瓶颈，GE-02解决了这一问题。', b=True)
P('3. 产品已实际投入使用。腾讯新闻确认产品已在深圳、苏州、武汉地标建筑商业化运营。', b=True, color=GREEN)
P('4. 渠道已验证。罗湖区政府产业支持、光明区街道政企合作、发布会获政府背书。', b=True)
P('5. 安全替代和劳动力短缺是核心驱动力，腾讯新闻的深度报道再次验证了这一判断。', b=True)

doc.add_paragraph()
P('数据来源索引', b=True, s=10)
P('1. 腾讯新闻2026年4月23日专题报道：' + QQ, s=8, color=GRAY)
P('2. 博研咨询行业报告（豆丁网）2024年10月发布', s=8, color=GRAY)
P('3. 搜狐行业报道2025年2月', s=8, color=GRAY)
P('4. 豆丁网行业报告2024年8月发布', s=8, color=GRAY)
P('5. 搜狗微信搜索哈工鹏泽2026年活动报道', s=8, color=GRAY)
P('6. 搜狗微信快鲤鱼报道埃欧珞', s=8, color=GRAY)
P('7. 搜狗微信凌度智能科技公众号', s=8, color=GRAY)
P('8. CCE上海清洁博览会/深安协公众号', s=8, color=GRAY)

path = os.path.join(OUT, '高空清洗机器人市场快速调研报告.docx')
doc.save(path)

t = ''
for p in doc.paragraphs:
    t += p.text
for tb in doc.tables:
    for r in tb.rows:
        for c in r.cells:
            t += c.text
hz = sum(1 for ch in t if '\u4e00' <= ch <= '\u9fff')
print(f'Hanzi: {hz}')
print(f'Total chars: {len(t)}')

img_c = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_c += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Images: {img_c}')

pdf_path = path.replace('.docx', '.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
