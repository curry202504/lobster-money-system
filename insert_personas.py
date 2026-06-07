# -*- coding: utf-8 -*-
"""Insert persona images into existing report"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
from docx2pdf import convert

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
safe = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', 'report-images-safe')

doc = Document(path)

persona_files = [
    '用户画像_S级超大型物业集团.png',
    '用户画像_A级酒店管理集团.png',
    '用户画像_A级大型清洁公司.png',
    '用户画像_B级政府及公建.png',
    '用户画像_C级中小清洁公司.png',
]

# Find the customer section and insert personas before the table
# We'll insert after the 2.1 最佳客户画像 heading
insert_after = None
for i, para in enumerate(doc.paragraphs):
    if '最佳客户画像' in para.text and para.style.name.startswith('Heading'):
        insert_after = i + 1
        break

if insert_after:
    # Insert persona images in reverse order
    for pf in reversed(persona_files):
        img_path = os.path.join(safe, pf)
        if os.path.exists(img_path):
            try:
                img = PILImage.open(img_path)
                w, h = img.size
                ratio = min(1.0, 450 / w) if w > 0 else 1
                
                # Create paragraph with image
                para_img = doc.add_paragraph()
                para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para_img.add_run()
                run.add_picture(img_path, width=Cm(ratio * w * 0.035))
                
                # Move after heading
                ref = doc.paragraphs[insert_after]._element
                ref.addnext(para_img._element)
                
                # Space
                space = doc.add_paragraph()
                ref.addnext(space._element)
                
                print(f'  Inserted: {pf}')
            except Exception as e:
                print(f'  Error: {e}')

doc.save(path)

# Count images
img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_count += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Total images in document: {img_count}')

# Count chars
text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text
print(f'Total chars: {len(text)}')

pdf_path = path.replace('.docx', '.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
