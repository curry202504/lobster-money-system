import os
from docx import Document
from docx.shared import Pt
from docx2pdf import convert

path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.docx')
doc = Document(path)

t = '''关于埃欧珞商业模式的参考价值。埃欧珞的产品售价在数十万元每台区间，3至4个订单即可收回研发成本，这说明几个问题。第一，高空清洗机器人产品的利润率较高，是一个有吸引力的商业市场。第二，客户有支付意愿，愿意为自动化清洗方案付费。第三，只要产品质量过硬，市场是愿意买单的。埃欧珞的销售模式也值得参考。直销2台起售意味着中小客户也有采购可能。代理10台起售说明代理商需要有一定的资金实力和市场开拓能力。支持租赁模式说明租赁是一种有效的市场教育方式。'''

for line in t.strip().split('\n'):
    line = line.strip()
    if line:
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(11)

doc.save(path)

text = ''
for p in doc.paragraphs:
    text += p.text
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            text += c.text
hz = sum(1 for ch in text if '\u4e00' <= ch <= '\u9fff')
t = len(text)
print(f'Hanzi: {hz}')
print(f'Total: {t}')

img_count = 0
for p in doc.paragraphs:
    for r in p.runs:
        img_count += len(r._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
print(f'Images: {img_count}')

pdf_path = os.path.join(os.environ['USERPROFILE'], '.openclaw', 'workspace', '高空清洗机器人市场快速调研报告.pdf')
convert(path, pdf_path)
print(f'PDF: {os.path.getsize(pdf_path)} bytes')
